"""
AI 数据分析工具 v3.0 — 问题驱动的五步分析流程
运行方式: streamlit run app.py
"""

import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from sklearn.linear_model import LinearRegression
from sklearn.metrics import r2_score, mean_squared_error
from scipy import stats
from scipy.stats import gaussian_kde
import warnings
import json
import requests
import os
import tempfile as _tempfile

# 跨平台临时目录（Windows 兼容）
_TMPDIR = _tempfile.gettempdir()

warnings.filterwarnings("ignore")

# ── 中文字体配置（本地 + 云端双兼容）────────────────────────────────────────
import matplotlib
import matplotlib.font_manager as fm

def _set_chinese_font():
    """
    优先读取项目目录中的字体文件，其次查找系统字体。
    兼容 Windows 本地 + Streamlit Cloud (Linux) 部署。
    支持 .ttf 和 .otf 格式。
    """
    import matplotlib.font_manager as fm

    # 第一步：在项目根目录查找字体文件（支持多种文件名和格式）
    _base_dir = os.path.dirname(os.path.abspath(__file__))
    _font_candidates = [
        "NotoSansSC-Regular.ttf",
        "NotoSansSC-Regular.otf",
        "NotoSansSC-Regular.ttf.otf",
        "NotoSansCJKsc-Regular.otf",
        "NotoSansCJKsc-Regular.ttf",
    ]
    for fname in _font_candidates:
        fpath = os.path.join(_base_dir, fname)
        if os.path.exists(fpath):
            try:
                # 用 matplotlib 的 ft2font 直接加载，跳过格式检测
                import matplotlib.ft2font as ft2
                ft2.FT2Font(fpath)  # 测试能否正常打开
                fm.fontManager.addfont(fpath)
                prop = fm.FontProperties(fname=fpath)
                font_name = prop.get_name()
                matplotlib.rcParams["font.family"] = font_name
                matplotlib.rcParams["axes.unicode_minus"] = False
                return font_name, fpath, f"成功加载字体文件 {fname}"
            except Exception as e:
                return None, fpath, f"字体加载失败：{e}"

    # 第二步：查找系统已有中文字体（本地 Windows/Mac 备用）
    candidates = [
        "Microsoft YaHei", "微软雅黑",
        "SimHei", "黑体",
        "SimSun", "宋体",
        "Noto Sans CJK SC", "NotoSansSC",
        "WenQuanYi Micro Hei",
        "Arial Unicode MS",
    ]
    available = {f.name for f in fm.fontManager.ttflist}
    for font in candidates:
        if font in available:
            matplotlib.rcParams["font.family"] = font
            matplotlib.rcParams["axes.unicode_minus"] = False
            return font, "系统字体", f"使用系统字体 {font}"

    matplotlib.rcParams["axes.unicode_minus"] = False
    try:
        files = [f for f in os.listdir(_base_dir) if f.endswith(('.ttf','.otf'))]
    except Exception:
        files = []
    return None, _base_dir, f"未找到任何中文字体，目录中字体文件：{files}"

_font_name, _font_path, _font_msg = _set_chinese_font()

# ══════════════════════════════════════════════════════════════════════════════
# 配置 & 样式
# ══════════════════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="AI 数据分析工具",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="collapsed",
)

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Noto+Sans+SC:wght@300;400;500;700&family=JetBrains+Mono:wght@400;600&display=swap');

html, body, [class*="css"] { font-family: 'Noto Sans SC', sans-serif; }
.stApp { background: #0b0e17; color: #dde2f0; }

/* ── 步骤导航栏 ── */
.step-bar {
    display: flex;
    gap: 0;
    margin-bottom: 2rem;
    border-radius: 10px;
    overflow: hidden;
    border: 1px solid #1e2840;
}
.step-item {
    flex: 1;
    padding: 0.65rem 0.5rem;
    text-align: center;
    font-size: 0.72rem;
    font-weight: 600;
    letter-spacing: 0.04em;
    color: #3d4a6a;
    background: #111520;
    border-right: 1px solid #1e2840;
    transition: background 0.2s;
}
.step-item:last-child { border-right: none; }
.step-item.active  { background: #162040; color: #5b8dee; }
.step-item.done    { background: #0d1f14; color: #3dd68c; }
.step-num { font-size: 0.65rem; display: block; margin-bottom: 0.15rem; opacity: 0.6; }

/* ── 卡片 ── */
.card {
    background: #141828;
    border: 1px solid #1e2840;
    border-radius: 12px;
    padding: 1.5rem 1.8rem;
    margin-bottom: 1.2rem;
}
.card-accent {
    background: #0d1520;
    border: 1px solid #1e3a5f;
    border-radius: 12px;
    padding: 1.5rem 1.8rem;
    margin-bottom: 1.2rem;
}

/* ── AI 解读块 ── */
.ai-block {
    background: #0d1520;
    border-left: 3px solid #5b8dee;
    border-radius: 0 10px 10px 0;
    padding: 1rem 1.4rem;
    margin: 0.8rem 0 1.4rem 0;
    font-size: 0.87rem;
    line-height: 1.9;
    color: #a8b4d0;
}
.ai-label {
    font-size: 0.68rem;
    font-weight: 700;
    letter-spacing: 0.14em;
    text-transform: uppercase;
    color: #5b8dee;
    margin-bottom: 0.6rem;
}

/* ── 推荐方案卡片 ── */
.rec-card {
    background: #111a2e;
    border: 1px solid #243050;
    border-radius: 10px;
    padding: 1.1rem 1.4rem;
    margin-bottom: 0.8rem;
}
.rec-title {
    font-size: 0.7rem;
    font-weight: 700;
    letter-spacing: 0.1em;
    text-transform: uppercase;
    color: #5b8dee;
    margin-bottom: 0.4rem;
}
.rec-value {
    font-size: 0.95rem;
    font-weight: 600;
    color: #dde2f0;
    margin-bottom: 0.3rem;
}
.rec-reason {
    font-size: 0.8rem;
    color: #6b7a9a;
    line-height: 1.6;
}

/* ── 步骤标题 ── */
.phase-header {
    display: flex;
    align-items: baseline;
    gap: 0.8rem;
    margin-bottom: 1.4rem;
}
.phase-num {
    font-size: 0.68rem;
    font-weight: 700;
    letter-spacing: 0.15em;
    text-transform: uppercase;
    color: #5b8dee;
    white-space: nowrap;
}
.phase-title {
    font-size: 1.35rem;
    font-weight: 700;
    color: #dde2f0;
    letter-spacing: -0.01em;
}

/* ── 统计数字 ── */
.stat-box {
    background: #161d30;
    border: 1px solid #1e2840;
    border-radius: 10px;
    padding: 1rem 1.2rem;
    text-align: center;
}
.stat-number {
    font-family: 'JetBrains Mono', monospace;
    font-size: 1.7rem;
    font-weight: 600;
    color: #5b8dee;
    line-height: 1;
}
.stat-label {
    font-size: 0.72rem;
    color: #5a6480;
    margin-top: 0.4rem;
}

/* ── 回归表格 ── */
.reg-table { width: 100%; border-collapse: collapse; }
.reg-table th {
    background: #161d30; color: #5b8dee;
    font-size: 0.72rem; letter-spacing: 0.08em; text-transform: uppercase;
    padding: 0.55rem 0.9rem; text-align: left; border-bottom: 1px solid #1e2840;
}
.reg-table td {
    padding: 0.5rem 0.9rem;
    font-family: 'JetBrains Mono', monospace;
    font-size: 0.8rem; color: #b8c2d8; border-bottom: 1px solid #161d30;
}
.reg-table tr:hover td { background: #151c2e; }
.sig { color: #3dd68c; }

/* ── 徽章 ── */
.badge { display:inline-block; padding:0.18rem 0.55rem; border-radius:20px;
         font-size:0.7rem; font-weight:600; font-family:'JetBrains Mono',monospace; }
.badge-blue  { background:#1a2d5a; color:#5b8dee; border:1px solid #2a4080; }
.badge-green { background:#0f2a1a; color:#3dd68c; border:1px solid #1a4a30; }
.badge-red   { background:#2a1a1a; color:#f06b6b; border:1px solid #4a2a2a; }
.badge-gray  { background:#1a2030; color:#7880a0; border:1px solid #252d40; }

/* ── 结论卡片 ── */
.conclusion-card {
    background: linear-gradient(135deg, #0d1a2e 0%, #0f2018 100%);
    border: 1px solid #1e4060;
    border-radius: 14px;
    padding: 1.8rem 2rem;
    margin-top: 1rem;
}
.conclusion-title {
    font-size: 1.1rem; font-weight: 700; color: #5b8dee;
    margin-bottom: 1rem; letter-spacing: -0.01em;
}

/* ── 问题展示 ── */
.question-display {
    background: #111828;
    border: 1px solid #243050;
    border-left: 3px solid #3dd68c;
    border-radius: 0 10px 10px 0;
    padding: 0.9rem 1.3rem;
    font-size: 0.92rem;
    color: #c8d4e8;
    margin-bottom: 1.5rem;
}

/* ── 分隔线 ── */
.divider { border:none; border-top:1px solid #1e2840; margin:1.8rem 0; }

/* ── 按钮 ── */
.stButton > button {
    background: linear-gradient(135deg, #2a4aaa 0%, #1a3490 100%);
    color: white; border: none; border-radius: 8px;
    font-family: 'Noto Sans SC', sans-serif; font-weight: 600;
    padding: 0.55rem 1.6rem; letter-spacing: 0.03em; transition: opacity 0.2s;
}
.stButton > button:hover { opacity: 0.85; }

/* ── 主标题 ── */
.main-title { font-size:1.9rem; font-weight:700; color:#dde2f0;
              letter-spacing:-0.02em; line-height:1.2; }
.sub-title  { font-size:0.88rem; color:#5a6480; margin-top:0.3rem; }
</style>
""", unsafe_allow_html=True)

# ── Matplotlib 主题 ──────────────────────────────────────────────────────────
PLOT_BG   = "#141828"
PLOT_FG   = "#dde2f0"
PLOT_GRID = "#1e2840"
ACCENT    = "#5b8dee"
ACCENT2   = "#3dd68c"

def apply_plot_style(fig, ax):
    fig.patch.set_facecolor(PLOT_BG)
    ax.set_facecolor(PLOT_BG)
    ax.tick_params(colors=PLOT_FG, labelsize=9)
    ax.xaxis.label.set_color(PLOT_FG)
    ax.yaxis.label.set_color(PLOT_FG)
    ax.title.set_color(PLOT_FG)
    for spine in ax.spines.values():
        spine.set_visible(False)
    ax.grid(False)
    fig.subplots_adjust(left=0.08, right=0.97, top=0.92, bottom=0.12)


# ══════════════════════════════════════════════════════════════════════════════
# Session State 初始化
# ══════════════════════════════════════════════════════════════════════════════
defaults = {
    "step":                      1,
    "lang":                      "zh",  # zh=中文 en=English
    "question":                  "",
    "df":                        None,
    "col_info":                  {},
    "plan":                      None,
    # 通用确认字段
    "confirmed_analysis_types":  [],
    # 描述性统计
    "confirmed_desc_vars":       [],
    # 对比 / 趋势
    "confirmed_target_vars":     [],
    "confirmed_group_var":       None,
    "confirmed_time_var":        None,
    # 相关性
    "confirmed_corr_vars":       [],
    # 回归
    "confirmed_y":               None,
    "confirmed_x":               [],
    "confirmed_scatter":         None,
    # 图表选择
    "chart_types":               [],   # 用户选择的图表类型列表
    "chart_config":              {},   # 每种图表的变量配置 dict
    # 分析缓存
    "analysis_done":             False,
    "reg_results":               None,
    "desc_summary":              "",
    "sig_vars_list":             [],
    "insig_vars_list":           [],
    "chart_images":              [],   # 本次分析生成的图表路径列表
    "report_text":               "",   # 生成的报告文本缓存
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v


# ══════════════════════════════════════════════════════════════════════════════
# 工具函数
# ══════════════════════════════════════════════════════════════════════════════
@st.cache_data
def load_data(file) -> pd.DataFrame:
    name = file.name.lower()
    if name.endswith(".csv"):
        for enc in ["utf-8", "gbk", "latin-1"]:
            try:
                file.seek(0)
                return pd.read_csv(file, encoding=enc)
            except Exception:
                continue
    elif name.endswith((".xlsx", ".xls")):
        return pd.read_excel(file)
    raise ValueError("不支持的文件格式")


def build_col_info(df: pd.DataFrame) -> dict:
    """生成列摘要用于发送给 AI"""
    info = {}
    for col in df.columns:
        s = df[col]
        if pd.api.types.is_numeric_dtype(s):
            info[col] = {
                "type": "numeric",
                "mean": round(float(s.mean()), 4) if s.notna().any() else None,
                "std":  round(float(s.std()),  4) if s.notna().any() else None,
                "missing_pct": round(s.isna().mean() * 100, 1),
            }
        else:
            info[col] = {
                "type": "categorical",
                "unique": int(s.nunique()),
                "missing_pct": round(s.isna().mean() * 100, 1),
            }
    return info


def dtype_badge(dtype) -> str:
    d = str(dtype)
    if "int" in d or "float" in d:
        return '<span class="badge badge-blue">数值</span>'
    elif "object" in d or "string" in d:
        return '<span class="badge badge-gray">文本</span>'
    elif "datetime" in d:
        return '<span class="badge badge-green">日期</span>'
    return f'<span class="badge badge-gray">{d}</span>'


def miss_badge(pct) -> str:
    if pct == 0:   return '<span class="badge badge-green">无缺失</span>'
    elif pct < 5:  return f'<span class="badge badge-blue">{pct:.1f}%</span>'
    elif pct < 20: return f'<span class="badge badge-gray">{pct:.1f}%</span>'
    return f'<span class="badge badge-red">{pct:.1f}%</span>'


def render_ai(label: str, text: str, is_ai: bool = True):
    _T = _get_T()
    source = _T.get("ai_label","✦ AI 解读") if is_ai else _T.get("rule_label","✦ 规则解读")
    lines = text.split("\n")
    formatted_lines = []
    for line in lines:
        line = line.strip()
        if not line:
            continue
        # 【核心结论】【关键发现】【建议】 三级标题
        if line.startswith("【") and "】" in line:
            formatted_lines.append(
                f'<div style="color:#5b8dee;font-weight:700;font-size:0.85rem;'
                f'margin-top:1rem;margin-bottom:0.4rem;padding-left:0;">'
                f'{line}</div>'
            )
        # 数字编号行：1. 2. 3. 等
        elif len(line) >= 2 and line[0].isdigit() and line[1] == '.':
            num = line[0]
            rest = line[2:].strip()
            formatted_lines.append(
                f'<div style="display:flex;gap:0.6rem;margin-bottom:0.5rem;'
                f'align-items:flex-start;padding-left:0.2rem;">'
                f'<span style="color:#5b8dee;font-weight:700;font-size:0.85rem;'
                f'min-width:1.2rem;flex-shrink:0;margin-top:0.05rem;">{num}.</span>'
                f'<span style="color:#c0cce0;font-size:0.87rem;line-height:1.7;">{rest}</span>'
                f'</div>'
            )
        else:
            formatted_lines.append(
                f'<div style="color:#a8b4d0;font-size:0.87rem;line-height:1.7;'
                f'margin-bottom:0.3rem;">{line}</div>'
            )
    html_text = "\n".join(formatted_lines)
    st.markdown(
        f'<div class="ai-block">'
        f'<div class="ai-label">{source} · {label}</div>'
        f'{html_text}'
        f'</div>',
        unsafe_allow_html=True,
    )


def call_claude(system: str, user: str, max_tokens: int = 1200) -> str:
    try:
        resp = requests.post(
    "https://api.anthropic.com/v1/messages",
    headers={
        "Content-Type": "application/json",
        "x-api-key": os.getenv("ANTHROPIC_API_KEY"),
        "anthropic-version": "2023-06-01"
    },
            json={
                "model": "claude-sonnet-4-5",
                "max_tokens": max_tokens,
                "system": system,
                "messages": [{"role": "user", "content": user}],
            },
            timeout=45,
        )

        # 1. 检查 HTTP 状态码
        if resp.status_code != 200:
            return f"（API 请求失败：HTTP {resp.status_code}，详情：{resp.json()}）"

        data = resp.json()

        # 2. 检查是否包含 content 字段
        if "content" not in data:
            return f"（API 返回异常，缺少 content 字段，完整响应：{data}）"

        # 3. 检查 content 格式是否正确
        content = data["content"]
        if not isinstance(content, list) or len(content) == 0:
            return f"（API 返回的 content 格式异常：{content}）"

        first = content[0]
        if not isinstance(first, dict) or "text" not in first:
            return f"（API 返回的 content[0] 缺少 text 字段：{first}）"

        return first["text"].strip()

    except requests.exceptions.Timeout:
        return "（AI 请求超时，请稍后重试）"
    except requests.exceptions.ConnectionError:
        return "（网络连接失败，请检查网络后重试）"
    except Exception as e:
        return f"（AI 调用出现未知错误：{type(e).__name__}: {e}）"


# ══════════════════════════════════════════════════════════════════════════════
# 本地 Fallback 规则文字生成
# ══════════════════════════════════════════════════════════════════════════════
def _is_api_error(text: str) -> bool:
    """判断 call_claude 返回的是否是错误信息而非正常文本"""
    return text.startswith("（") and ("不可用" in text or "失败" in text
                                      or "超时" in text or "异常" in text
                                      or "错误" in text or "连接" in text)


def fallback_desc(stats_dict: dict, focus_var: str, lang: str = "zh") -> str:
    lines = []
    if focus_var in stats_dict:
        s = stats_dict[focus_var]
        mean, median = s.get("均值", 0), s.get("中位数", 0)
        std  = s.get("标准差", 0)
        skew = s.get("偏度", 0)
        kurt = s.get("峰度", 0)
        diff_pct = abs(mean - median) / (abs(mean) + 1e-9) * 100
        cv = std / (abs(mean) + 1e-9)
        if lang == "en":
            t1 = f"1. [Central Tendency] Mean={mean:.3f}, Median={median:.3f} — {'Symmetric distribution.' if diff_pct < 5 else ('Right-skewed, mean pulled up by high values.' if mean > median else 'Left-skewed, mean pulled down by low values.')}"
            t2 = f"2. [Dispersion] Std={std:.3f}, CV={cv:.2f} — {'Data is concentrated, low variation.' if cv < 0.15 else ('Moderate dispersion.' if cv < 0.5 else 'High dispersion, significant individual differences.')}"
            t3 = f"3. [Distribution] Skewness={skew:.3f}, Kurtosis={kurt:.3f} — {'Approximately symmetric, near-normal.' if abs(skew) < 0.5 else ('Right-tailed, check for outliers.' if skew > 0 else 'Left-tailed, check for outliers.')}"
            t4 = f"4. [Implication] These features suggest monitoring skewness effects on {focus_var}; use both mean and median for robust interpretation."
        else:
            t1 = f"1. [集中趋势] 均值 {mean:.3f}，中位数 {median:.3f}，{'二者接近 — 数据分布较对称。' if diff_pct < 5 else ('均值>中位数 — 分布右偏，少数高值拉高均值。' if mean > median else '均值<中位数 — 分布左偏，少数低值拉低均值。')}"
            t2 = f"2. [离散程度] 标准差 {std:.3f}，变异系数 {cv:.2f} — {'数据集中，个体差异小。' if cv < 0.15 else ('存在一定离散程度。' if cv < 0.5 else '离散程度高，需关注异常值。')}"
            t3 = f"3. [分布形态] 偏度 {skew:.3f}，峰度 {kurt:.3f} — {'分布接近对称，基本符合正态假设。' if abs(skew) < 0.5 else ('正偏，分布右拖尾。' if skew > 0 else '负偏，分布左拖尾。')}"
            t4 = f"4. [研究意义] 分析 {focus_var} 时需关注偏态对均值的影响，建议结合中位数综合判断。"
        lines = [t1, t2, t3, t4]
    else:
        if lang == "en":
            lines.append(f"1. [Overview] Descriptive statistics computed for {focus_var} — see table above for mean, std, and distribution.")
        else:
            lines.append(f"1. [数据概况] 已完成 {focus_var} 的描述性统计计算 — 请参考上方表格。")
    return "\n".join(lines)


def fallback_corr(var_pairs: list, lang: str = "zh") -> str:
    if not var_pairs:
        if lang == "en":
            return "1. [Overview] Correlation matrix computed — see heatmap above; darker colors indicate stronger correlation."
        return "1. [数据概况] 已完成相关性矩阵计算 — 请参考上方热力图，颜色越深表示相关性越强。"
    strong, moderate, weak = [], [], []
    for v1, v2, r in var_pairs:
        ar = abs(r)
        direction = ("positive" if r > 0 else "negative") if lang == "en" else ("正相关" if r > 0 else "负相关")
        if ar >= 0.7:
            strong.append(f"{v1} & {v2} (r={r:.3f})")
        elif ar >= 0.4:
            moderate.append(f"{v1} & {v2} (r={r:.3f}, {direction})")
        elif ar >= 0.2:
            weak.append(f"{v1} & {v2} (r={r:.3f})")
    lines = []
    if lang == "en":
        if strong:   lines.append(f"1. [Strong |r|≥0.7] {'; '.join(strong)} — Strong linear relationship; watch for multicollinearity.")
        if moderate: lines.append(f"2. [Moderate 0.4≤|r|<0.7] {'; '.join(moderate)} — Moderate association; consider including in regression.")
        if weak:     lines.append(f"3. [Weak |r|<0.4] {'; '.join(weak)} — Weak linear relationship; may have non-linear patterns.")
        if not lines: lines.append("1. [Overview] All |r| < 0.2 — generally weak linear correlations; consider non-linear analysis.")
        lines.append("4. [Implication] Strongest correlated pairs are priority candidates for deeper analysis.")
    else:
        if strong:   lines.append(f"1. [强相关 |r|≥0.7] {'；'.join(strong)} — 变量间线性关系显著，需警惕多重共线性。")
        if moderate: lines.append(f"2. [中等相关 0.4≤|r|<0.7] {'；'.join(moderate)} — 存在一定线性关联，可纳入回归模型验证。")
        if weak:     lines.append(f"3. [弱相关 |r|<0.4] {'；'.join(weak)} — 线性关系不明显，可能存在非线性关系。")
        if not lines: lines.append("1. [相关性概况] 所有变量对 |r| 均低于 0.2 — 线性相关性普遍较弱，建议考虑非线性分析。")
        lines.append("4. [研究意义] 以上最强相关变量对是优先值得深入分析的方向，建议结合业务逻辑判断因果关系。")
    return "\n".join(lines)


def fallback_chart(y_var: str, mean: float, median: float, skew: float,
                   scatter_x: str, corr: float, lang: str = "zh") -> str:
    if lang == "en":
        skew_desc = "right-skewed" if skew > 0.5 else ("left-skewed" if skew < -0.5 else "approximately symmetric")
        strength  = "strong" if abs(corr) >= 0.7 else ("moderate" if abs(corr) >= 0.4 else "weak")
        direction = "positive" if corr > 0 else "negative"
        return (
            f"Distribution of {y_var} is {skew_desc} (skew={skew:.2f}), "
            f"mean={mean:.3f} {'above' if mean > median else 'below'} median={median:.3f}. "
            f"Scatter shows {strength} {direction} correlation with {scatter_x} (r={corr:.3f}). "
            f"{'Positive slope — variables move together.' if corr > 0 else 'Negative slope — variables move inversely.'} "
            f"{'Strong correlation — regression recommended.' if abs(corr) >= 0.5 else 'Weak correlation — interpret trend line cautiously.'}"
        )
    direction = "正相关" if corr > 0 else "负相关"
    strength  = "强" if abs(corr) >= 0.7 else ("中等" if abs(corr) >= 0.4 else "弱")
    skew_desc = "右偏" if skew > 0.5 else ("左偏" if skew < -0.5 else "近似对称")
    return (
        f"从分布直方图来看，{y_var} 的分布呈{skew_desc}（偏度={skew:.2f}），"
        f"均值 {mean:.3f} {'高于' if mean > median else '低于'}中位数 {median:.3f}。"
        f"散点图显示 {y_var} 与 {scatter_x} 之间存在{strength}{direction}（r={corr:.3f}），"
        f"{'趋势线斜率为正，表明两变量同向变动。' if corr > 0 else '趋势线斜率为负，表明两变量反向变动。'}"
        f"{'相关性较强，可进一步用回归分析量化影响。' if abs(corr) >= 0.5 else '相关性较弱，需谨慎解读趋势线。'}"
    )


def fallback_regression(y_var: str, x_vars_str: str, r2: float, adj_r2: float,
                        rmse: float, sig_vars: str, insig_vars: str, equation: str,
                        lang: str = "zh") -> str:
    if lang == "en":
        fit_level = "good (>0.6)" if r2 >= 0.6 else ("moderate (0.3–0.6)" if r2 >= 0.3 else "weak (<0.3)")
        lines = [
            f"1. [Model Fit] R²={r2:.4f}, Adj R²={adj_r2:.4f} — Model fit is {fit_level}, explaining {r2*100:.1f}% of {y_var} variance.",
            f"2. [Prediction] RMSE={rmse:.4f} — {'High accuracy.' if rmse < 0.1 else 'Some prediction error exists.'}",
        ]
        if sig_vars:
            lines.append(f"3. [Significant] Significant variables (p<0.05): {sig_vars} — statistically reliable effects on {y_var}.")
        else:
            lines.append("3. [Significant] No significant variables (p≥0.05) — possible causes: small sample, poor variable selection, or multicollinearity.")
        if insig_vars:
            lines.append(f"4. [Non-significant] {insig_vars} — consider removing to simplify model or increasing sample size.")
        lines.append(f"5. [Conclusion] Regression equation quantifies each variable's effect on {y_var}, providing empirical evidence for the research question.")
    else:
        fit_level = "较好（>0.6）" if r2 >= 0.6 else ("一般（0.3~0.6）" if r2 >= 0.3 else "较差（<0.3）")
        lines = [
            f"1. [模型拟合] R²={r2:.4f}，调整R²={adj_r2:.4f} — 模型拟合{fit_level}，解释了 {y_var} 约 {r2*100:.1f}% 的变异。",
            f"2. [预测精度] RMSE={rmse:.4f} — 预测值平均偏差约为 {rmse:.4f}，{'精度较高' if rmse < 0.1 else '存在一定预测误差'}。",
        ]
        if sig_vars:
            lines.append(f"3. [显著影响] 显著自变量（p<0.05）：{sig_vars} — 这些变量对 {y_var} 有统计上可靠的影响。")
        else:
            lines.append("3. [显著影响] 无显著自变量（p≥0.05）— 可能原因：样本量不足、变量选择不当或存在多重共线性。")
        if insig_vars:
            lines.append(f"4. [不显著变量] {insig_vars} — 建议考虑移除以简化模型，或扩大样本量后重新检验。")
        lines.append(f"5. [研究结论] 基于回归方程，可量化各自变量对 {y_var} 的影响大小，为研究问题提供了实证依据。")
    return "\n".join(lines)


def fallback_conclusion(question: str, y_var: str, x_vars_str: str,
                        analysis_types: list, r2: float, sig_vars: str,
                        miss_pct: float, n: int, desc_summary: str,
                        lang: str = "zh") -> str:
    has_reg = "回归分析" in analysis_types
    fit_level = "较强" if r2 >= 0.6 else ("一般" if r2 >= 0.3 else "有限")
    lines = ["【核心结论】"]
    if has_reg and r2 > 0:
        lines.append(f"1. 针对“{question}”，回归模型（R²={r2:.4f}）表明模型对 {y_var} 的解释力{fit_level}；{'显著影响因素：' + sig_vars if sig_vars else '暂未发现显著自变量'}。")
    else:
        lines.append(f"1. 针对“{question}”，本次通过 {'、'.join(analysis_types)} 分析，样本量 {n}，完成了对核心变量的系统性分析。")
    lines.append("\n【关键发现】")
    lines.append(f"1. 样本量 {n}，缺失率 {miss_pct:.1f}% — {'数据质量良好' if miss_pct < 5 else '缺失率偏高，结论需谨慎解读'}。")
    if has_reg and sig_vars:
        lines.append(f"2. 显著自变量 {sig_vars} — 对 {y_var} 有统计上可靠的影响，是关键解释变量。")
    if miss_pct > 5:
        lines.append(f"3. 数据缺失率 {miss_pct:.1f}% — 建议在正式分析前处理缺失值，以减少偏差。")
    lines.append("\n【建议】")
    lines.append("1. 结合专业知识验证上述统计结论，避免纯数据驱动的误判。")
    lines.append("2. 考虑增加控制变量或扩大样本量，以提升模型解释力和结论稳健性。")
    lines.append("3. 如存在非线性关系，可尝试多项式回归或机器学习方法进一步探索。")
    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════════════════
# AI 专项调用（带缓存 + fallback）
# ══════════════════════════════════════════════════════════════════════════════
SYSTEM_ANALYST = (
    "你是一位严谨友好的数据分析助教，帮助学生和研究人员完成统计分析任务。"
    "使用中文，语言清晰简洁，像分析报告一样专业。\n"
    "【输出格式要求（必须严格遵守）】\n"
    "1. 禁止输出任何长段落，所有内容必须分点表达。\n"
    "2. 每一点使用数字编号：1. 2. 3. 4. 每点单独一行。\n"
    "3. 每点格式固定为：编号. [标签] 数据描述（含具体数值）— 分析解释（说明意义）\n"
    "   例如：1. [集中趋势] 均值=3.45，中位数=3.20 — 均值高于中位数，分布右偏，存在高值样本。\n"
    "4. 结论部分必须用三个标题分组：【核心结论】【关键发现】【建议】，每组下各自编号 1. 2. 3.\n"
    "5. 不使用 ** 加粗、markdown 标题（# ## ###）或任何长段落。\n"
    "6. 每点不超过两行，简洁直接。"
)


def _call_with_fallback(prompt: str, fallback_text: str) -> tuple[str, bool]:
    """调用 API，失败时返回 fallback；返回 (text, is_ai)"""
    result = call_claude(_get_analyst_system(), prompt)
    if _is_api_error(result):
        return fallback_text, False
    return result, True

@st.cache_data(show_spinner=False)
def ai_design_plan(question: str, col_info_json: str, lang: str = "zh") -> dict:
    if lang == "en":
        user_prompt = (
            f"User research question: {question}\n\n"
            f"Dataset variable info (JSON):\n{col_info_json}\n\n"
            "As a data analysis plan designer, select appropriate analysis types for the question "
            "(do not select all by default, do not default to regression):\n"
            "1. Descriptive Stats\n2. Comparison\n3. Trend Analysis\n4. Correlation\n5. Regression\n\n"
            "Output ONLY the following JSON, no other text:\n"
            '{{\n'
            '  "analysis_types": ["Type1", "Type2"],\n'
            '  "analysis_methods": {{\n'
            '    "Type1": "specific method description (1-2 sentences)",\n'
            '    "Type2": "specific method description"\n'
            '  }},\n'
            '  "core_vars": {{\n'
            '    "dependent": "dependent variable name (or null)",\n'
            '    "independent": ["var1", "var2"],\n'
            '    "vars_reason": "why these variables (1-2 sentences)"\n'
            '  }},\n'
            '  "group_dims": ["grouping variable name"],\n'
            '  "group_dims_reason": "why group by this dimension, or null",\n'
            '  "charts": [\n'
            '    {{"name": "chart name", "vars": "variables used", "purpose": "what question this chart answers"}}\n'
            '  ],\n'
            '  "use_regression": true or false,\n'
            '  "regression_reason": "rationale for regression (1-2 sentences)",\n'
            '  "logic": "overall analysis logic: what to do first, then what, and what question to answer (3-5 sentences)"\n'
            '}}'
        )
    else:
        user_prompt = (
            f"用户的研究问题：{question}\n\n"
            f"数据集变量信息（JSON）：\n{col_info_json}\n\n"
            "请你作为数据分析方案设计师，围绕用户的研究问题，"
            "从以下五大分析类型中选择合适的组合（不要默认全选，也不要默认使用回归）：\n"
            "1. 描述性统计\n2. 对比分析\n3. 趋势分析\n4. 相关性分析\n5. 回归分析\n\n"
            "请严格按以下 JSON 格式输出，不要输出任何其他内容：\n"
            '{{\n'
            '  "analysis_types": ["类型1", "类型2"],\n'
            '  "analysis_methods": {{\n'
            '    "类型1": "针对该类型的具体分析方法说明（1-2句，结合数据变量）",\n'
            '    "类型2": "针对该类型的具体分析方法说明"\n'
            '  }},\n'
            '  "core_vars": {{\n'
            '    "dependent": "因变量名（若适用，否则填 null）",\n'
            '    "independent": ["自变量1", "自变量2"],\n'
            '    "vars_reason": "为什么选这些变量（1-2句）"\n'
            '  }},\n'
            '  "group_dims": ["分组维度变量名"],\n'
            '  "group_dims_reason": "为什么按此维度分组，若无分组填 null",\n'
            '  "charts": [\n'
            '    {{"name": "图表名称", "vars": "涉及变量", "purpose": "这张图回答什么问题"}}\n'
            '  ],\n'
            '  "use_regression": true或false,\n'
            '  "regression_reason": "是否建议回归的理由（1-2句）",\n'
            '  "logic": "整体分析逻辑说明：先做什么、再做什么、最终回答什么问题（3-5句）"\n'
            '}}'
        )
    # ai_design_plan 使用独立的 system prompt，不受 SYSTEM_ANALYST 分点格式干扰
    _PLAN_SYSTEM = (
        "你是一位数据分析方案设计师。"
        "你的任务是根据用户问题和数据集，输出一个 JSON 格式的分析方案。"
        "重要：只输出纯 JSON 文本，第一个字符必须是 {，最后一个字符必须是 }。"
        "绝对不要输出 ```json 或 ``` 或任何 markdown 标记。"
        "绝对不要在 JSON 前后加任何文字说明。"
    )
    raw = call_claude(_PLAN_SYSTEM, user_prompt, max_tokens=1200)
    try:
        import re as _re
        clean = raw.strip()
        # 强力提取：找到第一个 { 和最后一个 } 之间的内容
        start = clean.find("{")
        end   = clean.rfind("}")
        if start != -1 and end != -1 and end > start:
            clean = clean[start:end+1]
        return json.loads(clean)
    except Exception as e:
        return {"_raw": raw, "_error": f"JSON 解析失败：{e}"}


@st.cache_data(show_spinner=False)
def ai_desc_interp(question: str, stats_json: str, focus_var: str,
                   fallback_text: str, lang: str = "zh") -> tuple:
    if lang == "en":
        prompt = (
            f"Research question: {question}\n"
            f"Key variable: {focus_var}\n"
            f"Descriptive statistics (JSON):\n{stats_json}\n\n"
            "Strictly follow this format — no long paragraphs:\n\n"
            "1. [Central Tendency] Mean/median values — explain distribution symmetry\n"
            "2. [Dispersion] Std dev value — explain individual variation\n"
            "3. [Distribution Shape] Skewness/kurtosis — explain shape and research impact\n"
            "4. [Research Implication] What these features mean for the research question\n\n"
            "Format: N. [Tag] Data (specific values) — Explanation. One line per point."
        )
    else:
        prompt = (
            f"研究问题：{question}\n"
            f"核心变量：{focus_var}\n"
            f"描述性统计摘要（JSON）：\n{stats_json}\n\n"
            "请结合研究问题，严格按以下格式输出描述性统计解读，禁止写长段落：\n\n"
            "1. [集中趋势] 均值/中位数具体数值 — 说明数据分布是否对称及意义\n"
            "2. [离散程度] 标准差具体数值 — 说明个体差异大小及其含义\n"
            "3. [分布形态] 偏度/峰度具体数值 — 说明分布特征及对研究问题的影响\n"
            "4. [研究意义] 上述特征结合研究问题意味着什么\n\n"
            "严格格式：数字编号. [标签] 数据描述（具体数值）— 分析解释，每点单独一行，不超过两行。"
        )
    return _call_with_fallback(prompt, fallback_text)


@st.cache_data(show_spinner=False)
def ai_corr_interp(question: str, corr_summary: str, fallback_text: str, lang: str = "zh") -> tuple:
    if lang == "en":
        prompt = (
            f"Research question: {question}\n"
            f"Correlation analysis summary:\n{corr_summary}\n\n"
            "Strictly follow this format — no long paragraphs:\n\n"
            "1. [Strong |r|≥0.7] List variable pairs and r values — explain strength\n"
            "2. [Moderate 0.4≤|r|<0.7] List pairs and r values — explain association\n"
            "3. [Weak |r|<0.4] List pairs and r values — explain weak linear relationship\n"
            "4. [Research Implication] Key insight from correlation for the research question\n\n"
            "Format: N. [Tag] Data (specific values) — Explanation. Skip empty categories."
        )
    else:
        prompt = (
            f"研究问题：{question}\n"
            f"相关性分析摘要：\n{corr_summary}\n\n"
            "请严格按以下格式输出相关性解读，禁止写长段落：\n\n"
            "1. [强相关 |r|≥0.7] 列出变量对和具体 r 值 — 说明关系强度意义\n"
            "2. [中等相关 0.4≤|r|<0.7] 列出变量对和具体 r 值 — 说明关联程度\n"
            "3. [弱/无相关 |r|<0.4] 列出变量对和具体 r 值 — 说明线性关系不明显\n"
            "4. [研究意义] 相关性对回答研究问题最关键的启示\n\n"
            "严格格式：数字编号. [标签] 数据描述（具体数值）— 分析解释，每点单独一行。\n"
            "若某类别无变量对，该点跳过，编号重排。"
        )
    return _call_with_fallback(prompt, fallback_text)


@st.cache_data(show_spinner=False)
def ai_chart_interp(question: str, y_var: str, mean: float, median: float,
                    skew: float, scatter_x: str, corr: float,
                    fallback_text: str, lang: str = "zh") -> tuple:
    if lang == "en":
        prompt = (
            f"Research question: {question}\n"
            f"Variable {y_var}: mean={mean:.3f}, median={median:.3f}, skewness={skew:.3f}.\n"
            f"Scatter plot: Pearson r between {y_var} and {scatter_x} = {corr:.3f}.\n\n"
            "Strictly follow this format — no long paragraphs:\n\n"
            f"1. [Distribution] {y_var} skewness={skew:.3f}, mean vs median — explain symmetry\n"
            f"2. [Correlation] {y_var} vs {scatter_x} r={corr:.3f} — explain strength and direction\n"
            "3. [Preliminary Conclusion] Initial conclusion for the research question from chart\n\n"
            "Format: N. [Tag] Data (specific values) — Explanation. One line per point."
        )
    else:
        prompt = (
            f"研究问题：{question}\n"
            f"变量 {y_var}：均值={mean:.3f}，中位数={median:.3f}，偏度={skew:.3f}。\n"
            f"散点图：{y_var} 与 {scatter_x} 的 Pearson 相关系数 = {corr:.3f}。\n\n"
            "请严格按以下格式输出图表解读，禁止写长段落：\n\n"
            f"1. [分布形态] {y_var} 偏度={skew:.3f}，均值与中位数对比 — 说明分布是否对称及含义\n"
            f"2. [相关关系] {y_var} 与 {scatter_x} 的 r={corr:.3f} — 说明关联强度、方向及趋势含义\n"
            "3. [初步判断] 基于图形，对研究问题能得出的初步结论\n\n"
            "严格格式：数字编号. [标签] 数据描述（具体数值）— 分析解释，每点单独一行。"
        )
    return _call_with_fallback(prompt, fallback_text)


@st.cache_data(show_spinner=False)
def ai_reg_interp(question: str, y_var: str, x_vars_str: str,
                  r2: float, adj_r2: float, rmse: float,
                  sig_vars: str, insig_vars: str, equation: str,
                  fallback_text: str, lang: str = "zh") -> tuple:
    if lang == "en":
        prompt = (
            f"Research question: {question}\n"
            f"Dependent: {y_var}, Independent: {x_vars_str}\n"
            f"R²={r2:.4f}, Adj R²={adj_r2:.4f}, RMSE={rmse:.4f}\n"
            f"Significant (p<0.05): {sig_vars or 'None'}\n"
            f"Not significant: {insig_vars or 'None'}\n"
            f"Equation: {equation}\n\n"
            "Strictly follow this format — no long paragraphs:\n\n"
            f"1. [Model Fit] R²={r2:.4f}, Adj R²={adj_r2:.4f} — explain variance explained and fit quality\n"
            f"2. [Prediction Accuracy] RMSE={rmse:.4f} — explain prediction error magnitude\n"
            "3. [Significant Effects] List significant vars and coefficient direction — explain impact\n"
            "4. [Non-significant Vars] List non-significant vars — discuss possible reasons\n"
            "5. [Research Conclusion] What regression results mean for the research question\n\n"
            "Format: N. [Tag] Data (specific values) — Explanation. Skip empty categories."
        )
    else:
        prompt = (
            f"研究问题：{question}\n"
            f"因变量：{y_var}，自变量：{x_vars_str}\n"
            f"R²={r2:.4f}，调整 R²={adj_r2:.4f}，RMSE={rmse:.4f}\n"
            f"显著变量（p<0.05）：{sig_vars or '无'}\n"
            f"不显著变量：{insig_vars or '无'}\n"
            f"回归方程：{equation}\n\n"
            "请严格按以下格式输出回归解读，禁止写长段落：\n\n"
            f"1. [模型拟合] R²={r2:.4f}，调整R²={adj_r2:.4f} — 说明模型解释了多少变异，拟合程度评价\n"
            f"2. [预测精度] RMSE={rmse:.4f} — 说明预测误差大小及实际意义\n"
            "3. [显著影响] 列出显著变量（p<0.05）及系数方向 — 说明各变量对因变量的具体影响\n"
            "4. [不显著变量] 列出不显著变量（p≥0.05）— 分析可能原因\n"
            "5. [研究结论] 回归结果对回答研究问题的直接意义\n\n"
            "严格格式：数字编号. [标签] 数据描述（具体数值）— 分析解释，每点单独一行。\n"
            "若某类别为空，该点跳过，编号重排。"
        )
    return _call_with_fallback(prompt, fallback_text)


@st.cache_data(show_spinner=False)
def ai_conclusion(question: str, y_var: str, x_vars_str: str,
                  analysis_types_str: str, r2: float, sig_vars: str,
                  miss_pct: float, n: int, desc_summary: str,
                  fallback_text: str, lang: str = "zh") -> tuple:
    if lang == "en":
        prompt = (
            f"Original research question: {question}\n\n"
            f"Analysis summary:\n"
            f"- Sample size {n}, missing rate {miss_pct:.1f}%\n"
            f"- Analysis types: {analysis_types_str}\n"
            f"- Core variables: {y_var} (independent: {x_vars_str})\n"
            f"- Regression R²={r2:.4f}\n"
            f"- Significant variables: {sig_vars or 'None'}\n"
            f"- Descriptive summary: {desc_summary}\n\n"
            "Output the overall conclusion in three sections, no long paragraphs:\n\n"
            "[Core Conclusion]\n"
            "1. Directly answer the research question (one sentence with key data)\n\n"
            "[Key Findings]\n"
            "1. Most important finding (with values) — significance\n"
            "2. Second finding (with values) — significance\n"
            "3. Data quality or limitation reminder (missing rate, sample size)\n\n"
            "[Recommendations]\n"
            "1. First actionable recommendation for the research question\n"
            "2. Second actionable recommendation\n"
            "3. Future analysis directions\n\n"
            "Format: N. Content (one line per point, max 2 lines, include specific values)"
        )
    else:
        prompt = (
            f"用户最初的研究问题：{question}\n\n"
            f"分析摘要：\n"
            f"- 样本量 {n}，缺失率 {miss_pct:.1f}%\n"
            f"- 执行的分析类型：{analysis_types_str}\n"
            f"- 核心变量：{y_var}（自变量：{x_vars_str}）\n"
            f"- 回归模型 R²={r2:.4f}\n"
            f"- 显著自变量：{sig_vars or '无'}\n"
            f"- 描述性统计摘要：{desc_summary}\n\n"
            "请严格按以下三层结构输出总体结论，每层用标题行分隔，禁止写长段落：\n\n"
            "【核心结论】\n"
            "1. 直接回答研究问题（一句话，结合关键数据和具体数值）\n\n"
            "【关键发现】\n"
            "1. 最重要的分析发现（含具体数值）— 说明意义\n"
            "2. 第二重要发现（含具体数值）— 说明意义\n"
            "3. 数据质量或局限性提醒（缺失率、样本量等具体数值）\n\n"
            "【建议】\n"
            "1. 针对研究问题的第一条具体可执行建议\n"
            "2. 针对研究问题的第二条具体可执行建议\n"
            "3. 后续分析方向建议（如增加变量、非线性模型等）\n\n"
            "严格格式：数字编号. 内容（每点单独一行，不超过两行，含具体数值）"
        )
    result = call_claude(_get_analyst_system(), prompt, max_tokens=1400)
    if _is_api_error(result):
        return fallback_text, False
    return result, True



# ══════════════════════════════════════════════════════════════════════════════
# 多语言翻译字典 (zh / en)
# ══════════════════════════════════════════════════════════════════════════════
_TRANSLATIONS = {
    "zh": {
        "app_title":     "AI 数据分析工具",
        "app_subtitle":  "问题驱动 · AI 协作 · 数据与解读同步输出",
        "lang_btn":      "🌐 English",
        "step_labels":   [("01","问题导入"),("02","方案设计与确认"),
                          ("03","分析结果与解释"),("04","总体结论")],
        "s1_title":      "问题导入",
        "s1_desc":       "在开始分析之前，请先描述你想研究的问题或目标。<br>明确的问题能帮助 AI 推荐更合适的变量和分析方法，让后续结论更有针对性。",
        "s1_q_label":    "你的研究问题 / 分析目标",
        "s1_placeholder":"例如：哪些因素影响学生的期末成绩？\n例如：房屋面积和地段对房价有多大影响？\n例如：员工工作年限与薪资之间是否存在显著关系？",
        "s1_upload":     "📂 上传数据文件",
        "s1_upload_hint":"支持 CSV（UTF-8 / GBK）和 Excel（.xlsx / .xls）",
        "s1_next":       "下一步：AI 方案设计 →",
        "s1_err_q":      "请先输入你的研究问题。",
        "s1_err_file":   "请上传数据文件。",
        "s1_err_num":    "数据中没有可分析的数值列，请检查文件内容。",
        "s2_title":      "分析方案设计与确认",
        "s2_q_prefix":   "🔍 研究问题：",
        "s2_overview":   "📋 数据概览（点击展开）",
        "s2_ov_n":       "样本量", "s2_ov_cols":"变量数", "s2_ov_miss":"缺失率",
        "s2_plan_hdr":   "AI 方案推荐",
        "s2_plan_err":   "AI 方案生成遇到解析问题，请在下方手动配置变量后继续。",
        "s2_plan_raw":   "查看原始 AI 输出",
        "s2_adjust_hdr": "调整与确认",
        "s2_adjust_tip": "AI 已根据你的问题预填了以下配置，你可以直接确认，也可以修改后再继续。",
        "s2_types_label":"分析类型（可多选，选择后下方变量配置随之变化）",
        "s2_types_help": "所选类型直接决定执行哪种分析，无需额外勾选",
        "s2_all_types":  ["描述性统计","对比分析","趋势分析","相关性分析","回归分析"],
        "s2_back":       "← 返回修改问题",
        "s2_regen":      "↺ 重新生成方案",
        "s2_confirm":    "✓ 确认方案，开始分析 →",
        "s2_err_type":   "请至少选择一种分析类型。",
        "s3_title":      "分析结果与解释",
        "s3_q_prefix":   "🔍 研究问题：",
        "s3_sec_a":      "A · 描述性统计",
        "s3_sec_b":      "B · 对比分析", "s3_sec_b2":"B · 趋势分析", "s3_sec_b3":"B · 对比分析 / 趋势分析",
        "s3_sec_c":      "C · 相关性分析",
        "s3_sec_d":      "D · 回归分析",
        "s3_no_target":  "未配置目标变量，请返回方案步骤选择。",
        "s3_no_corr":    "相关性分析需要至少 2 个变量，请返回方案步骤配置。",
        "s3_no_reg":     "回归分析需要配置因变量和自变量，请返回方案步骤。",
        "s3_spin_desc":  "✦ 正在生成描述性统计解读…",
        "s3_spin_corr":  "✦ 正在生成相关性解读…",
        "s3_spin_chart": "✦ 正在生成图表趋势解读…",
        "s3_spin_reg":   "✦ 正在生成回归结果解读…",
        "s3_lbl_desc":   "描述性统计解读",
        "s3_lbl_corr":   "相关性分析解读",
        "s3_lbl_chart":  "分布形态 · 趋势判断",
        "s3_lbl_reg":    "回归结果解读",
        "s3_model_sum":  "模型摘要",
        "s3_reg_coef":   "回归系数",
        "s3_desc_cols":  ["样本量","均值","标准差","最小值","25%分位","中位数","75%分位","最大值","偏度","峰度"],
        "s3_reg_cols":   ["变量","系数","标准误","t 值","p 值","95% 置信区间"],
        "s3_resid_ttl":  "残差 vs 拟合值", "s3_resid_x":"拟合值", "s3_resid_y":"残差",
        "s3_qq_ttl":     "Q-Q 图", "s3_qq_x":"理论分位数", "s3_qq_y":"样本分位数",
        "s3_back":       "← 返回调整方案",
        "s3_next":       "下一步：总体结论与建议 →",
        "s3_r2":"R²","s3_adj_r2":"调整 R²","s3_rmse":"RMSE","s3_n_eff":"有效样本量",
        "s4_title":      "总体结论与建议",
        "s4_stat_types": "已执行分析","s4_stat_n":"有效样本量",
        "s4_stat_r2":    "回归 R²","s4_stat_sig":"显著变量",
        "s4_no_reg":     "未执行回归",
        "s4_spinner":    "✦ 正在生成总体结论与建议…",
        "s4_conc_title": "📝 总体结论与建议",
        "s4_ai_gen":     "（AI 生成）","s4_local_gen":"（本地规则生成）",
        "s4_rpt_title":  "📄 分析报告生成与下载",
        "s4_rpt_hint":   "系统将汇总本次所有分析结果，生成一份结构化分析报告，支持 Word (.docx) 格式下载。",
        "s4_gen_btn":    "✦ 生成分析报告",
        "s4_gen_spin":   "✦ 正在生成结构化报告…",
        "s4_gen_done":   "报告生成完成，请点击下方按钮下载。",
        "s4_preview":    "📋 报告预览（点击展开）",
        "s4_dl_word":    "⬇ 下载 Word 报告 (.docx)",
        "s4_dl_txt":     "⬇ 下载纯文本报告 (.txt)",
        "s4_no_docx":    "需要安装 python-docx：`pip install python-docx`",
        "s4_back":       "← 返回查看分析详情",
        "s4_restart":    "↺ 重新开始新分析",
        "s4_rpt_cover":  "数据分析报告",
        "s4_rpt_q":      "研究问题：",
        "s4_rpt_date":   "生成时间：",
        "s4_rpt_annex":  "附录：分析图表",
        "s4_rpt_fig":    "图 ",
        "corr_heatmap":  "Pearson 相关系数热力图",
        "ai_label":      "✦ AI 解读",
        "rule_label":    "✦ 规则解读",
        "card_types":    "分析类型",
        "card_vars":     "核心变量",
        "card_dep":      "因变量：",
        "card_indep":    "自变量：",
        "card_group":    "分组维度",
        "card_no_group": "无",
        "card_no_group_reason": "无需分组",
        "card_charts":   "推荐图表",
        "card_reg":      "是否建议回归",
        "card_reg_yes":  "建议回归",
        "card_reg_no":   "暂不需要",
        "card_logic":    "分析逻辑说明",
        "s2_desc_vars":  "📋 描述性统计 · 分析变量",
        "s2_desc_sel":   "选择参与描述性统计的变量",
        "s2_corr_vars":  "🔗 相关性分析 · 参与变量",
        "s2_corr_sel":   "参与相关性分析的变量（建议 2-8 个数值变量）",
        "s2_corr_help":  "将计算所选变量之间的两两 Pearson 相关系数并绘制热力图",
        "s2_reg_vars":   "📈 回归分析 · 变量配置",
        "s2_dep_var":    "因变量（Y）",
        "s2_indep_vars": "自变量（X）",
        "s2_scatter_sel":"散点图：选择与因变量对比展示的自变量",
        "card_ai_plan":  "AI 方案设计",
        "s2_spinner":    "✦ AI 正在设计分析方案，请稍候…",
        # ── 图表配置 UI 标签 ──
        "chart_x_cat":       "X 轴（分类变量）",
        "chart_y_num":       "Y 轴（数值变量）",
        "chart_x_num":       "X 轴（数值变量，条形长度）",
        "chart_y_cat":       "Y 轴（分类变量，横轴类别）",
        "chart_group_opt":   "分组变量（可选）",
        "chart_no_group":    "（不分组）",
        "chart_x_time":      "X 轴（时间/序列）",
        "chart_y_multi":     "Y 变量（数值，可多选≤4）",
        "chart_warn_nocat":  "柱状图需要分类型 X 轴，当前数据集无明显分类变量。",
        "chart_warn_nocat2": "⚠️ 无纯分类列，已候选低唯一值数值列",
        "chart_warn_notime": "⚠️ 未检测到时间变量，请确认 X 轴是否为有序序列",
        "chart_warn_line4":  "折线图最多 4 条线，已截取前 4 个。",
        "chart_warn_sc2":    "散点图需要至少 2 个数值变量。",
        "chart_warn_bub3":   "气泡图需要至少 3 个数值变量（X / Y / 气泡大小）。",
        "chart_warn_min1":   "折线图至少需要 1 个 Y 变量。",
        "chart_x_num2":      "X 变量（数值）",
        "chart_y_num2":      "Y 变量（数值，须与 X 不同）",
        "chart_bubble_x":    "X 变量（数值）",
        "chart_bubble_y":    "Y 变量（数值）",
        "chart_bubble_size": "气泡大小变量（数值）",
        "chart_pie_label":   "类别变量（各扇区名称）",
        "chart_pie_value":   "数值变量（各扇区大小）",
        "chart_donut_label": "类别变量",
        "chart_donut_value": "数值变量",
        "chart_box_num":     "数值变量（必选）",
        "chart_box_grp":     "分组变量（可选，仅分类/低基数列）",
        "chart_violin_num":  "数值变量（必选）",
        "chart_violin_grp":  "分组变量（可选）",
        "chart_radar_vars":  "维度变量（数值，选 3-8 个）",
        "chart_radar_grp":   "分组变量（可选，用于多组对比）",
        "chart_rank_var":    "排序变量（数值变量）",
        "chart_combo_x":     "X 轴（分类/时间）",
        "chart_combo_bar":   "柱状 Y（数值）",
        "chart_combo_line":  "折线 Y（数值，双轴）",
        "chart_sel_label":   "选择图表类型（可多选，已根据分析类型和问题自动推荐）",
        "chart_sel_help":    "选择后下方自动显示该图表所需的变量配置",
        "s2_chart_title":    "图表选择（Visualization Configuration）",
        "s2_chart_hint":     "先选图表类型，系统将自动约束可选变量，避免无效组合。",
        # ── Step 2 变量组 ──
        "s2_group_detect":   "系统从 {n} 个数值变量中检测到 <b style='color:#dde2f0;'>{g} 个变量组</b>，选择一组后将自动纳入该组全部变量：",
        "s2_group_sel":      "选择变量组（可多选，每组所有变量自动纳入对比）",
        "s2_group_extra":    "额外追加单个变量（可选）",
        "s2_group_added":    "✓ 已纳入 {n} 个变量：{vars}",
        "s2_pair_hint":      "未检测到明显分组结构，请手动选择要对比的变量（至少 2 个）：",
        "s2_pair_a":         "对比变量 A",
        "s2_pair_b":         "对比变量 B",
        "s2_pair_extra":     "额外追加更多对比变量（可选）",
        "s2_group_var":      "分组维度（Group by，可选）",
        "s2_no_group":       "（不分组）",
        "s2_time_var":       "时间变量（可选，用于趋势图）",
        "s2_no_time":        "（不使用）",
        # ── Step 3 图表副标题 ──
        "ct_bar":       "柱状图",
        "ct_hbar":      "条形图",
        "ct_gbar":      "分组柱状图",
        "ct_line":      "折线图",
        "ct_area":      "面积图",
        "ct_scatter":   "散点图",
        "ct_bubble":    "气泡图",
        "ct_pie":       "饼图",
        "ct_donut":     "圆环图",
        "ct_box":       "箱线图",
        "ct_violin":    "小提琴图",
        "ct_radar":     "雷达图",
        "ct_heatmap":   "热力图",
        "ct_rank":      "变化排序图",
        "ct_combo":     "组合图",
        "ct_by":        "按",
        "ct_grouped":   "（分组：",
        "ct_trend":     "趋势",
        "ct_share":     "占比",
        "ct_vs":        "vs",
        "ct_unified":   "统一",
        "ct_fail":      "绘制失败：",
        "ct_mean_cmp":  "变量对比",
        "ct_mean_y":    "均值（±标准差）",
        "ct_mean_title":"各变量均值对比",
        "ct_min2":      "请至少选择 2 个变量以生成对比图。",
        "s3_model_r2":  ["R²","调整 R²","RMSE","有效样本量"],
        # ── Step 4 ──
        "s4_word_fail": "Word 生成失败：",
        "sig_label":     "个显著",
        "ov_col_name":   "变量名",
        "ov_col_type":   "类型",
        "ov_col_uniq":   "唯一值",
        "ov_col_miss":   "缺失",
        "chart_violin_sub": "小提琴图",
        "chart_radar_sub":  "雷达图",
        "chart_heat_sub":   "热力图 · 相关矩阵",
        "chart_combo_sub":  "组合图",
        "chart_scatter_sub":"散点图",
        "chart_dist_sub":   "分布",
        "s3_no_sample":     "有效样本量不足，无法运行回归。",
        "ct_scatter_vs":    "散点图",
        "chart_warn_radar3":  "雷达图至少需要 3 个维度变量。",
        "chart_heat_caption": "热力图将自动使用已选数值变量的 Pearson 相关矩阵，或指定行/列变量。",
        "chart_heat_sel":     "参与热力图的数值变量（留空则自动）",
        "chart_warn_sc_xy":   "散点图需同时选择 X 和 Y 变量。",
        "chart_warn_radar_e": "雷达图需要至少 3 个维度变量。",
        "chart_rank_x":  "样本排序（从小到大）",
        "chart_mean":    "均值",
        "chart_freq":    "频数",
        "chart_dist":    "分布直方图",
        "chart_trend":   "趋势线",
        "chart_median":  "中位数",
        "chart_rules": {
            "柱状图":"X: 分类变量  |  Y: 数值变量",
            "条形图":"Y: 分类变量  |  X: 数值变量（横向）",
            "分组柱状图":"X: 分类  |  Y: 数值  |  分组: 可选分类",
            "折线图":"X: 时间/序列（优先）  |  Y: 数值（可多选）",
            "面积图":"X: 时间/序列  |  Y: 数值（可多选）",
            "散点图":"X: 数值  |  Y: 数值（须不同列）",
            "气泡图":"X: 数值  |  Y: 数值  |  气泡大小: 第三数值",
            "饼图":"分类变量: 类别  |  数值变量: 各类别值",
            "圆环图":"分类变量: 类别  |  数值变量: 各类别值",
            "箱线图":"数值变量: 必选  |  分组变量: 可选分类",
            "小提琴图":"数值变量: 必选  |  分组变量: 可选分类",
            "雷达图":"维度变量: 3-8个数值  |  分组变量: 可选",
            "热力图":"数值变量（自动生成相关矩阵）",
            "变化排序图":"排序变量: 数值变量",
            "组合图（柱+线）":"X: 分类/时间  |  柱Y: 数值  |  线Y: 数值",
        },
        "prompt_lang":   "中文",
    },
    "en": {
        "app_title":     "AI Data Analysis Tool",
        "app_subtitle":  "Question-Driven · AI-Assisted · Data & Insights Unified",
        "lang_btn":      "🌐 中文",
        "step_labels":   [("01","Problem Input"),("02","Plan Design"),
                          ("03","Analysis Results"),("04","Conclusion")],
        "s1_title":      "Problem Input",
        "s1_desc":       "Describe the question or goal you want to explore.<br>A clear question helps AI recommend better variables and analysis methods.",
        "s1_q_label":    "Your Research Question / Analysis Goal",
        "s1_placeholder":"e.g. What factors affect students' exam scores?\ne.g. How do area and location influence house prices?\ne.g. Is there a significant relationship between tenure and salary?",
        "s1_upload":     "📂 Upload Data File",
        "s1_upload_hint":"Supports CSV (UTF-8 / GBK) and Excel (.xlsx / .xls)",
        "s1_next":       "Next: AI Plan Design →",
        "s1_err_q":      "Please enter your research question first.",
        "s1_err_file":   "Please upload a data file.",
        "s1_err_num":    "No numeric columns found. Please check your file.",
        "s2_title":      "Plan Design & Confirmation",
        "s2_q_prefix":   "🔍 Research Question: ",
        "s2_overview":   "📋 Data Overview (click to expand)",
        "s2_ov_n":       "Sample Size","s2_ov_cols":"Variables","s2_ov_miss":"Missing Rate",
        "s2_plan_hdr":   "AI Recommended Plan",
        "s2_plan_err":   "AI plan parsing failed. Please configure variables manually below.",
        "s2_plan_raw":   "View Raw AI Output",
        "s2_adjust_hdr": "Adjust & Confirm",
        "s2_adjust_tip": "AI has pre-filled the configuration below. Confirm or modify before continuing.",
        "s2_types_label":"Analysis Types (multi-select, variable config updates accordingly)",
        "s2_types_help": "Selected types directly drive what analysis is executed",
        "s2_all_types":  ["Descriptive Stats","Comparison","Trend Analysis","Correlation","Regression"],
        "s2_back":       "← Back to Edit Question",
        "s2_regen":      "↺ Regenerate Plan",
        "s2_confirm":    "✓ Confirm Plan & Start Analysis →",
        "s2_err_type":   "Please select at least one analysis type.",
        "s3_title":      "Analysis Results & Interpretation",
        "s3_q_prefix":   "🔍 Research Question: ",
        "s3_sec_a":      "A · Descriptive Statistics",
        "s3_sec_b":      "B · Comparison Analysis","s3_sec_b2":"B · Trend Analysis","s3_sec_b3":"B · Comparison / Trend Analysis",
        "s3_sec_c":      "C · Correlation Analysis",
        "s3_sec_d":      "D · Regression Analysis",
        "s3_no_target":  "No target variables configured. Please go back to plan step.",
        "s3_no_corr":    "Correlation requires at least 2 variables. Please go back.",
        "s3_no_reg":     "Regression requires dependent and independent variables. Please go back.",
        "s3_spin_desc":  "✦ Generating descriptive statistics interpretation…",
        "s3_spin_corr":  "✦ Generating correlation interpretation…",
        "s3_spin_chart": "✦ Generating chart trend interpretation…",
        "s3_spin_reg":   "✦ Generating regression interpretation…",
        "s3_lbl_desc":   "Descriptive Statistics",
        "s3_lbl_corr":   "Correlation Analysis",
        "s3_lbl_chart":  "Distribution & Trend",
        "s3_lbl_reg":    "Regression Results",
        "s3_model_sum":  "Model Summary",
        "s3_reg_coef":   "Regression Coefficients",
        "s3_desc_cols":  ["Count","Mean","Std","Min","25%","Median","75%","Max","Skew","Kurt"],
        "s3_reg_cols":   ["Variable","Coef","Std Err","t","p-value","95% CI"],
        "s3_resid_ttl":  "Residuals vs Fitted","s3_resid_x":"Fitted Values","s3_resid_y":"Residuals",
        "s3_qq_ttl":     "Q-Q Plot","s3_qq_x":"Theoretical Quantiles","s3_qq_y":"Sample Quantiles",
        "s3_back":       "← Back to Adjust Plan",
        "s3_next":       "Next: Overall Conclusion →",
        "s3_r2":"R²","s3_adj_r2":"Adj R²","s3_rmse":"RMSE","s3_n_eff":"Sample Size",
        "s4_title":      "Overall Conclusion & Recommendations",
        "s4_stat_types": "Analysis Executed","s4_stat_n":"Sample Size",
        "s4_stat_r2":    "Model R²","s4_stat_sig":"Significant Vars",
        "s4_no_reg":     "No regression",
        "s4_spinner":    "✦ Generating overall conclusion…",
        "s4_conc_title": "📝 Overall Conclusion & Recommendations",
        "s4_ai_gen":     "(AI Generated)","s4_local_gen":"(Rule-Based)",
        "s4_rpt_title":  "📄 Report Generation & Download",
        "s4_rpt_hint":   "Compile all analysis results into a structured report, downloadable as Word (.docx).",
        "s4_gen_btn":    "✦ Generate Analysis Report",
        "s4_gen_spin":   "✦ Generating structured report…",
        "s4_gen_done":   "Report generated. Click below to download.",
        "s4_preview":    "📋 Report Preview (click to expand)",
        "s4_dl_word":    "⬇ Download Word Report (.docx)",
        "s4_dl_txt":     "⬇ Download Plain Text Report (.txt)",
        "s4_no_docx":    "python-docx required: `pip install python-docx`",
        "s4_back":       "← Back to Analysis Details",
        "s4_restart":    "↺ Start New Analysis",
        "s4_rpt_cover":  "Data Analysis Report",
        "s4_rpt_q":      "Research Question: ",
        "s4_rpt_date":   "Generated: ",
        "s4_rpt_annex":  "Appendix: Analysis Charts",
        "s4_rpt_fig":    "Figure ",
        "corr_heatmap":  "Pearson Correlation Heatmap",
        "ai_label":      "✦ AI Insight",
        "rule_label":    "✦ Rule-Based Insight",
        "card_types":    "Analysis Types",
        "card_vars":     "Core Variables",
        "card_dep":      "Dependent: ",
        "card_indep":    "Independent: ",
        "card_group":    "Group Dimension",
        "card_no_group": "None",
        "card_no_group_reason": "No grouping needed",
        "card_charts":   "Recommended Charts",
        "card_reg":      "Regression Recommended",
        "card_reg_yes":  "Recommend",
        "card_reg_no":   "Not needed",
        "card_logic":    "Analysis Logic",
        "s2_desc_vars":  "📋 Descriptive Stats · Variables",
        "s2_desc_sel":   "Select variables for descriptive statistics",
        "s2_corr_vars":  "🔗 Correlation · Variables",
        "s2_corr_sel":   "Variables for correlation (2–8 numeric recommended)",
        "s2_corr_help":  "Pairwise Pearson coefficients and heatmap",
        "s2_reg_vars":   "📈 Regression · Variable Config",
        "s2_dep_var":    "Dependent Variable (Y)",
        "s2_indep_vars": "Independent Variables (X)",
        "s2_scatter_sel":"Scatter Plot: select X variable to compare with Y",
        "card_ai_plan":  "AI Plan Design",
        "s2_spinner":    "✦ Generating analysis plan, please wait…",
        # ── Chart config UI labels ──
        "chart_x_cat":       "X Axis (Category)",
        "chart_y_num":       "Y Axis (Numeric)",
        "chart_x_num":       "X Axis (Numeric, bar length)",
        "chart_y_cat":       "Y Axis (Category)",
        "chart_group_opt":   "Group Variable (optional)",
        "chart_no_group":    "(No grouping)",
        "chart_x_time":      "X Axis (Time/Sequence)",
        "chart_y_multi":     "Y Variables (numeric, multi-select ≤4)",
        "chart_warn_nocat":  "Bar chart requires a categorical X axis. No clear categorical column found.",
        "chart_warn_nocat2": "⚠️ No pure categorical columns, using low-cardinality numeric columns",
        "chart_warn_notime": "⚠️ No time variable detected. Confirm X axis is an ordered sequence.",
        "chart_warn_line4":  "Line chart max 4 lines. Truncated to first 4.",
        "chart_warn_sc2":    "Scatter plot requires at least 2 numeric variables.",
        "chart_warn_bub3":   "Bubble chart requires at least 3 numeric variables.",
        "chart_warn_min1":   "Line chart requires at least 1 Y variable.",
        "chart_x_num2":      "X Variable (numeric)",
        "chart_y_num2":      "Y Variable (numeric, must differ from X)",
        "chart_bubble_x":    "X Variable (numeric)",
        "chart_bubble_y":    "Y Variable (numeric)",
        "chart_bubble_size": "Bubble Size Variable (numeric)",
        "chart_pie_label":   "Category Variable (sector labels)",
        "chart_pie_value":   "Numeric Variable (sector sizes)",
        "chart_donut_label": "Category Variable",
        "chart_donut_value": "Numeric Variable",
        "chart_box_num":     "Numeric Variable (required)",
        "chart_box_grp":     "Group Variable (optional, categorical only)",
        "chart_violin_num":  "Numeric Variable (required)",
        "chart_violin_grp":  "Group Variable (optional)",
        "chart_radar_vars":  "Dimension Variables (numeric, select 3–8)",
        "chart_radar_grp":   "Group Variable (optional, for multi-group comparison)",
        "chart_rank_var":    "Sort Variable (numeric)",
        "chart_combo_x":     "X Axis (category/time)",
        "chart_combo_bar":   "Bar Y (numeric)",
        "chart_combo_line":  "Line Y (numeric, dual axis)",
        "chart_sel_label":   "Select chart types (multi-select, auto-recommended)",
        "chart_sel_help":    "Variable config for each chart appears below after selection",
        "s2_chart_title":    "Chart Selection (Visualization Configuration)",
        "s2_chart_hint":     "Select chart type first — variables will be constrained automatically.",
        # ── Step 2 variable groups ──
        "s2_group_detect":   "System detected <b style='color:#dde2f0;'>{g} variable groups</b> from {n} numeric columns:",
        "s2_group_sel":      "Select Variable Groups (all variables auto-included)",
        "s2_group_extra":    "Add Individual Variables (optional)",
        "s2_group_added":    "✓ {n} variables included: {vars}",
        "s2_pair_hint":      "No clear group structure. Manually select at least 2 variables:",
        "s2_pair_a":         "Variable A",
        "s2_pair_b":         "Variable B",
        "s2_pair_extra":     "Add more comparison variables (optional)",
        "s2_group_var":      "Group By (optional)",
        "s2_no_group":       "(No grouping)",
        "s2_time_var":       "Time Variable (optional, for trend chart)",
        "s2_no_time":        "(Not used)",
        # ── Step 3 chart subtitles ──
        "ct_bar":       "Bar Chart",
        "ct_hbar":      "Horizontal Bar",
        "ct_gbar":      "Grouped Bar",
        "ct_line":      "Line Chart",
        "ct_area":      "Area Chart",
        "ct_scatter":   "Scatter Plot",
        "ct_bubble":    "Bubble Chart",
        "ct_pie":       "Pie Chart",
        "ct_donut":     "Donut Chart",
        "ct_box":       "Box Plot",
        "ct_violin":    "Violin Plot",
        "ct_radar":     "Radar Chart",
        "ct_heatmap":   "Heatmap",
        "ct_rank":      "Ranked Change Chart",
        "ct_combo":     "Combined Chart",
        "ct_by":        "by",
        "ct_grouped":   "(grouped: ",
        "ct_trend":     "trend",
        "ct_share":     "share",
        "ct_vs":        "vs",
        "ct_unified":   "uniform",
        "ct_fail":      "Chart failed: ",
        "ct_mean_cmp":  "Variable Comparison",
        "ct_mean_y":    "Mean (±Std Dev)",
        "ct_mean_title":"Mean Comparison",
        "ct_min2":      "Please select at least 2 variables to generate a comparison chart.",
        "s3_model_r2":  ["R²","Adj R²","RMSE","Sample Size"],
        # ── Step 4 ──
        "s4_word_fail": "Word generation failed: ",
        "sig_label":     "significant",
        "ov_col_name":   "Variable",
        "ov_col_type":   "Type",
        "ov_col_uniq":   "Unique",
        "ov_col_miss":   "Missing",
        "chart_violin_sub": "Violin Plot",
        "chart_radar_sub":  "Radar Chart",
        "chart_heat_sub":   "Heatmap · Correlation Matrix",
        "chart_combo_sub":  "Combined Chart",
        "chart_scatter_sub":"Scatter Plot",
        "chart_dist_sub":   "Distribution",
        "s3_no_sample":     "Insufficient sample size for regression.",
        "ct_scatter_vs":    "Scatter Plot",
        "chart_warn_radar3":  "Radar chart requires at least 3 dimension variables.",
        "chart_heat_caption": "Heatmap will use the Pearson correlation matrix of selected numeric variables.",
        "chart_heat_sel":     "Numeric variables for heatmap (leave empty for auto)",
        "chart_warn_sc_xy":   "Scatter plot requires both X and Y variables.",
        "chart_warn_radar_e": "Radar chart requires at least 3 dimension variables.",
        "chart_rank_x":  "Sorted Index (ascending)",
        "chart_mean":    "Mean",
        "chart_freq":    "Frequency",
        "chart_dist":    "Distribution Histogram",
        "chart_trend":   "Trend Line",
        "chart_median":  "Median",
        "chart_rules": {
            "Bar Chart":"X: Category  |  Y: Numeric",
            "Horizontal Bar":"Y: Category  |  X: Numeric (bar length)",
            "Grouped Bar":"X: Category  |  Y: Numeric  |  Group: optional",
            "Line Chart":"X: Time/Sequence (preferred)  |  Y: Numeric (multi-select)",
            "Area Chart":"X: Time/Sequence  |  Y: Numeric (multi-select)",
            "Scatter Plot":"X: Numeric  |  Y: Numeric (must differ from X)",
            "Bubble Chart":"X: Numeric  |  Y: Numeric  |  Size: 3rd numeric",
            "Pie Chart":"Category: labels  |  Numeric: sector sizes",
            "Donut Chart":"Category: labels  |  Numeric: sector sizes",
            "Box Plot":"Numeric: required  |  Group: optional category",
            "Violin Plot":"Numeric: required  |  Group: optional category",
            "Radar Chart":"Dimensions: 3–8 numeric  |  Group: optional",
            "Heatmap":"Numeric variables (auto correlation matrix)",
            "Ranked Change Chart":"Sort variable: numeric",
            "Combined Chart":"X: category/time  |  Bar Y: numeric  |  Line Y: numeric",
        },
        "prompt_lang":   "English",
    }
}


def _get_T():
    """获取当前语言的翻译字典"""
    return _TRANSLATIONS[st.session_state.get("lang","zh")]


def _get_analyst_system():
    """返回当前语言的 AI system prompt"""
    lang = st.session_state.get("lang","zh")
    if lang == "zh":
        return (
            "你是一位严谨友好的数据分析助教，帮助学生和研究人员完成统计分析任务。"
            "使用中文，语言清晰简洁，像分析报告一样专业。\n"
            "【输出格式要求（必须严格遵守）】\n"
            "1. 禁止输出任何长段落，所有内容必须分点表达。\n"
            "2. 每一点使用数字编号：1. 2. 3. 4. 每点单独一行。\n"
            "3. 每点格式：编号. [标签] 数据描述（含具体数值）— 分析解释（说明意义）\n"
            "4. 结论用三个标题：【核心结论】【关键发现】【建议】，每组下各自编号。\n"
            "5. 不使用 ** 加粗、# 标题或长段落。每点不超过两行。"
        )
    return (
        "You are a rigorous and friendly data analysis assistant."
        " Use English. Be clear, concise, and professional like an analysis report.\n"
        "[OUTPUT FORMAT - STRICTLY FOLLOW]\n"
        "1. No long paragraphs. All content must be in numbered points.\n"
        "2. Use numeric numbering: 1. 2. 3. Each point on its own line.\n"
        "3. Each point: N. [Tag] Data description (with values) — Analysis explanation\n"
        "4. Conclusion uses three headings: [Core Conclusion] [Key Findings] [Recommendations]\n"
        "5. No ** bold, no # headers. Max 2 lines per point."
    )

# ══════════════════════════════════════════════════════════════════════════════
# 步骤导航条
# ══════════════════════════════════════════════════════════════════════════════
step_labels = [
    ("01", "问题导入"),
    ("02", "方案设计与确认"),
    ("03", "分析结果与解释"),
    ("04", "总体结论"),
]

def render_step_bar(current: int):
    T = _get_T()
    items = ""
    for i, (num, label) in enumerate(T["step_labels"]):
        s = i + 1
        css = "active" if s == current else ("done" if s < current else "")
        icon = "✓ " if s < current else ""
        items += (
            f'<div class="step-item {css}">'
            f'<span class="step-num">STEP {num}</span>'
            f'{icon}{label}</div>'
        )
    st.markdown(f'<div class="step-bar">{items}</div>', unsafe_allow_html=True)


def phase_header(num: str, title: str):
    st.markdown(
        f'<div class="phase-header">'
        f'<span class="phase-num">STEP {num}</span>'
        f'<span class="phase-title">{title}</span>'
        f'</div>',
        unsafe_allow_html=True,
    )


# ══════════════════════════════════════════════════════════════════════════════
# 语言切换 + 顶部标题
# ══════════════════════════════════════════════════════════════════════════════
T = _get_T()  # 当前语言翻译

_col_title, _col_lang = st.columns([9, 1])
with _col_title:
    st.markdown(
        f'<div class="main-title">{T["app_title"]}</div>'
        f'<div class="sub-title">{T["app_subtitle"]}</div>',
        unsafe_allow_html=True,
    )
with _col_lang:
    _cur_lang = st.session_state.get("lang","zh")
    _lang_options = ["🇨🇳 中文", "🇺🇸 English"]
    _lang_idx = 0 if _cur_lang == "zh" else 1
    _lang_sel = st.selectbox("🌐", _lang_options, index=_lang_idx,
                             key="_lang_sel", label_visibility="collapsed")
    _new_lang = "zh" if _lang_sel == "🇨🇳 中文" else "en"
    if _new_lang != _cur_lang:
        st.session_state["lang"] = _new_lang
        st.session_state["plan"] = None   # 清空方案缓存，切换语言后重新生成
        T = _get_T()
        st.rerun()

st.markdown("<div style='height:0.4rem'></div>", unsafe_allow_html=True)
render_step_bar(st.session_state.step)
st.markdown("<hr class='divider'>", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
# STEP 1 · 问题导入 + 数据上传
# ══════════════════════════════════════════════════════════════════════════════
if st.session_state.step == 1:
    T = _get_T()
    phase_header("01", T["s1_title"])

    st.markdown(
        f'<div class="card">'
        f'<div style="font-size:0.82rem;color:#6b7a9a;line-height:1.8;margin-bottom:1.2rem;">'
        f'{T["s1_desc"]}</div>',
        unsafe_allow_html=True,
    )

    question_input = st.text_area(
        T["s1_q_label"],
        placeholder=T["s1_placeholder"],
        height=120,
        value=st.session_state.question,
    )
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown(
        '<div class="card">',
        unsafe_allow_html=True,
    )
    st.markdown(
        '<div style="font-size:0.72rem;font-weight:700;letter-spacing:0.12em;'
        'text-transform:uppercase;color:#5b8dee;margin-bottom:0.8rem;">📂 上传数据文件</div>',
        unsafe_allow_html=True,
    )
    uploaded = st.file_uploader(
        T["s1_upload_hint"],
        type=["csv", "xlsx", "xls"],
    )
    st.markdown("</div>", unsafe_allow_html=True)

    col_btn, col_tip = st.columns([1, 3])
    with col_btn:
        go = st.button(T["s1_next"], width='stretch')

    if go:
        if not question_input.strip():
            st.error(T["s1_err_q"])
        elif uploaded is None:
            st.error(T["s1_err_file"])
        else:
            try:
                df = load_data(uploaded)
                st.session_state.df        = df
                st.session_state.question  = question_input.strip()
                st.session_state.col_info  = build_col_info(df)
                st.session_state.plan      = None   # 清除旧方案
                st.session_state.step      = 2
                st.rerun()
            except Exception as e:
                st.error(T.get("s1_err_load","文件读取失败：") + str(e))


# ══════════════════════════════════════════════════════════════════════════════
# STEP 2 · AI 分析方案设计 + 用户确认与调整（合并步骤）
# ══════════════════════════════════════════════════════════════════════════════
elif st.session_state.step == 2:
    T = _get_T()
    phase_header("02", T["s2_title"])

    st.markdown(
        f'<div class="question-display">{T["s2_q_prefix"]}{st.session_state.question}</div>',
        unsafe_allow_html=True,
    )

    df       = st.session_state.df
    num_cols = df.select_dtypes(include=np.number).columns.tolist()
    all_cols = df.columns.tolist()

    # ── 数据概览（折叠） ──────────────────────────────────────────────────────
    with st.expander(T["s2_overview"], expanded=False):
        c1, c2, c3 = st.columns(3)
        miss_pct_ov = df.isnull().sum().sum() / df.size * 100
        for col, val, lbl in zip(
            [c1, c2, c3],
            [df.shape[0], df.shape[1], f"{miss_pct_ov:.1f}%"],
            [T["s2_ov_n"], T["s2_ov_cols"], T["s2_ov_miss"]],
        ):
            col.markdown(
                f'<div class="stat-box"><div class="stat-number">{val}</div>'
                f'<div class="stat-label">{lbl}</div></div>',
                unsafe_allow_html=True,
            )
        st.markdown("<div style='height:0.8rem'></div>", unsafe_allow_html=True)
        rows_html = ""
        for c in df.columns:
            mp = df[c].isnull().mean() * 100
            rows_html += (
                f"<tr><td><b>{c}</b></td><td>{dtype_badge(df[c].dtype)}</td>"
                f"<td>{df[c].nunique():,}</td><td>{miss_badge(mp)}</td></tr>"
            )
        st.markdown(
            f"<table class='reg-table'><thead><tr>"
            f"<th>{T.get('ov_col_name','变量名')}</th><th>{T.get('ov_col_type','类型')}</th><th>{T.get('ov_col_uniq','唯一值')}</th><th>{T.get('ov_col_miss','缺失')}</th>"
            f"</tr></thead><tbody>{rows_html}</tbody></table>",
            unsafe_allow_html=True,
        )

    # ── AI 生成分析方案 ────────────────────────────────────────────────────────
    _cur_plan_lang = st.session_state.get("plan_lang","")
    _cur_lang_now  = st.session_state.get("lang","zh")
    if st.session_state.plan is None or _cur_plan_lang != _cur_lang_now:
        _spin_txt = "✦ Generating analysis plan, please wait…" if _cur_lang_now == "en" else "✦ AI 正在设计分析方案，请稍候…"
        with st.spinner(_spin_txt):
            plan = ai_design_plan(
                st.session_state.question,
                json.dumps(st.session_state.col_info, ensure_ascii=False),
                lang=_cur_lang_now,
            )
            st.session_state.plan = plan
            st.session_state.plan_lang = _cur_lang_now

    plan = st.session_state.plan

    # ── 方案展示区 ────────────────────────────────────────────────────────────
    if "_error" in plan:
        st.warning(T["s2_plan_err"])
        with st.expander(T["s2_plan_raw"], expanded=False):
            st.code(plan.get("_raw", ""))
    else:
        st.markdown(
            '<div style="font-size:0.72rem;font-weight:700;letter-spacing:0.13em;'
            'text-transform:uppercase;color:#5b8dee;margin:1.2rem 0 0.8rem 0;">'
            f'{T["card_ai_plan"]}</div>',
            unsafe_allow_html=True,
        )

        # 六张推荐卡并排展示（2列）
        col_l, col_r = st.columns(2)

        # 分析类型
        a_types = plan.get("analysis_types", [])
        a_methods = plan.get("analysis_methods", {})
        methods_text = "  ·  ".join([
            f"{t}：{a_methods.get(t, '')}" for t in a_types
        ])
        with col_l:
            st.markdown(
                f'<div class="rec-card">'
                f'<div class="rec-title">{T["card_types"]}</div>'
                f'<div class="rec-value">{"  ·  ".join(a_types) if a_types else "—"}</div>'
                f'<div class="rec-reason">{methods_text}</div>'
                f'</div>',
                unsafe_allow_html=True,
            )

        # 核心变量
        core = plan.get("core_vars", {})
        dep  = core.get("dependent") or "—"
        indep = core.get("independent", [])
        with col_r:
            st.markdown(
                f'<div class="rec-card">'
                f'<div class="rec-title">{T["card_vars"]}</div>'
                f'<div class="rec-value">{T["card_dep"]}{dep}</div>'
                f'<div class="rec-reason">'
                f'{T["card_indep"]}{" · ".join(indep) if indep else "—"}<br>'
                f'{core.get("vars_reason", "")}'
                f'</div>'
                f'</div>',
                unsafe_allow_html=True,
            )

        # 分组维度
        g_dims = plan.get("group_dims", [])
        g_reason = plan.get("group_dims_reason") or T.get("card_no_group_reason","无需分组")
        with col_l:
            st.markdown(
                f'<div class="rec-card">'
                f'<div class="rec-title">{T["card_group"]}</div>'
                f'<div class="rec-value">{" · ".join(g_dims) if g_dims else T["card_no_group"]}</div>'
                f'<div class="rec-reason">{g_reason}</div>'
                f'</div>',
                unsafe_allow_html=True,
            )

        # 推荐图表
        charts = plan.get("charts", [])
        charts_text = "  ·  ".join([
            f'{c.get("name","")}（{c.get("purpose","")}）' for c in charts
        ]) if charts else "—"
        with col_r:
            st.markdown(
                f'<div class="rec-card">'
                f'<div class="rec-title">{T["card_charts"]}</div>'
                f'<div class="rec-value">{" · ".join([c.get("name","") for c in charts]) if charts else "—"}</div>'
                f'<div class="rec-reason">{charts_text}</div>'
                f'</div>',
                unsafe_allow_html=True,
            )

        # 是否建议回归
        use_reg    = plan.get("use_regression", False)
        reg_reason = plan.get("regression_reason", "")
        reg_badge  = (
            f'<span class="badge badge-green">{T["card_reg_yes"]}</span>'
            if use_reg else
            f'<span class="badge badge-gray">{T["card_reg_no"]}</span>'
        )
        with col_l:
            st.markdown(
                f'<div class="rec-card">'
                f'<div class="rec-title">{T["card_reg"]}</div>'
                f'<div class="rec-value">{reg_badge}</div>'
                f'<div class="rec-reason">{reg_reason}</div>'
                f'</div>',
                unsafe_allow_html=True,
            )

        # 分析逻辑说明
        with col_r:
            st.markdown(
                f'<div class="rec-card">'
                f'<div class="rec-title">{T["card_logic"]}</div>'
                f'<div class="rec-reason" style="color:#c0cce0;line-height:1.75;">'
                f'{plan.get("logic", "—")}'
                f'</div>'
                f'</div>',
                unsafe_allow_html=True,
            )

    # ── 用户调整区 ────────────────────────────────────────────────────────────
    st.markdown("<hr class='divider'>", unsafe_allow_html=True)
    st.markdown(
        '<div style="font-size:0.72rem;font-weight:700;letter-spacing:0.13em;'
        'text-transform:uppercase;color:#5b8dee;margin-bottom:0.8rem;">'
        f'{T["s2_adjust_hdr"]}</div>',
        unsafe_allow_html=True,
    )
    st.markdown(
        f'<div style="font-size:0.82rem;color:#6b7a9a;margin-bottom:1.2rem;">{T["s2_adjust_tip"]}</div>',
        unsafe_allow_html=True,
    )

    st.markdown('<div class="card">', unsafe_allow_html=True)

    # ── 第一行：分析类型（多选）────────────────────────────────────────────
    all_analysis_types = T["s2_all_types"]
    ai_types_raw = plan.get("analysis_types", []) if "_error" not in plan else ["描述性统计"]
    # 双向映射：无论 plan 返回中文还是英文类型名，都能匹配当前语言的选项
    _zh2en = {"描述性统计":"Descriptive Stats","对比分析":"Comparison",
              "趋势分析":"Trend Analysis","相关性分析":"Correlation","回归分析":"Regression"}
    _en2zh = {v:k for k,v in _zh2en.items()}
    _lang_now = st.session_state.get("lang","zh")
    if _lang_now == "en":
        ai_types = [_zh2en.get(t, t) for t in ai_types_raw]
    else:
        ai_types = [_en2zh.get(t, t) for t in ai_types_raw]
    # fallback：若默认值为空（匹配失败）则选全部
    _default_types = [t for t in ai_types if t in all_analysis_types]
    if not _default_types and ai_types_raw:
        _default_types = all_analysis_types[:min(len(ai_types_raw), len(all_analysis_types))]
    chosen_types = st.multiselect(
        T["s2_types_label"],
        all_analysis_types,
        default=_default_types,
        help=T["s2_types_help"],
    )

    st.markdown("<div style='height:0.6rem'></div>", unsafe_allow_html=True)

    # ── 预读 AI 推荐变量 ────────────────────────────────────────────────────
    core_vars = plan.get("core_vars", {}) if "_error" not in plan else {}
    ai_dep    = core_vars.get("dependent") or ""
    ai_indep  = core_vars.get("independent", [])
    ai_groups = plan.get("group_dims", []) if "_error" not in plan else []

    # 收集变量配置（按分析类型动态显示）
    cfg = {}   # 存放每种类型的变量配置

    # ── 描述性统计变量 ──────────────────────────────────────────────────────
    # 将 chosen_types 映射为中文用于变量配置的判断
    _chosen_zh = [_en2zh.get(t,t) for t in chosen_types]
    if "描述性统计" in _chosen_zh:
        st.markdown(
            f'<div style="font-size:0.75rem;color:#5b8dee;font-weight:600;margin:0.8rem 0 0.4rem 0;">{T["s2_desc_vars"]}</div>',
            unsafe_allow_html=True,
        )
        default_desc = [v for v in (ai_indep + ([ai_dep] if ai_dep else [])) if v in num_cols]
        if not default_desc:
            default_desc = num_cols[:min(4, len(num_cols))]
        cfg["desc_vars"] = st.multiselect(
            T["s2_desc_sel"],
            num_cols,
            default=default_desc,
            key="sel_desc_vars",
        )

    # ── 对比分析 / 趋势分析变量 ─────────────────────────────────────────────
    if "对比分析" in _chosen_zh or "趋势分析" in _chosen_zh:
        active_label = "对比分析 / 趋势分析"
        if "对比分析" in _chosen_zh and "趋势分析" not in _chosen_zh:
            active_label = "对比分析"
        elif "趋势分析" in _chosen_zh and "对比分析" not in _chosen_zh:
            active_label = "趋势分析"
        st.markdown(
            f'<div style="font-size:0.75rem;color:#5b8dee;font-weight:600;'
            f'margin:0.8rem 0 0.4rem 0;">📊 {active_label} · 变量配置</div>',
            unsafe_allow_html=True,
        )

        # ══════════════════════════════════════════════════════════════════════
        # 通用变量组识别引擎
        # 策略1：公共前缀  "Sales_Q1 / Sales_Q2 / Sales_Q3"
        # 策略2：公共后缀  "Revenue_2020 / Cost_2020 / Profit_2020"
        # 策略3：去除变化类词 "GDP_Change / GDP_Growth / GDP_Level"
        # 策略4：去除数字/年份 "Score_1989 / Score_2009"
        # 策略5：分隔符结构  "A_before / A_after / B_before / B_after"
        # ══════════════════════════════════════════════════════════════════════
        import re as _re
        from collections import defaultdict as _dd

        # ── 辅助：标准化列名 → token 列表 ──────────────────────────────────
        _NOISE = {
            # 时间类
            'jan','feb','mar','apr','may','jun','jul','aug','sep','oct','nov','dec',
            'q1','q2','q3','q4','h1','h2',
            # 变化类
            'change','growth','rate','pct','percent','delta','diff',
            'increase','decrease','variation','shift',
            # 统计类
            'mean','avg','average','sum','total','count','max','min',
            'index','score','value','level','amount',
            # 时态类
            'before','after','pre','post','base','current','prev','next',
            'start','end','initial','final',
        }

        def _tokenize(col: str) -> list:
            """把列名拆成 token，去除分隔符"""
            s = col.strip()
            tokens = _re.split(r'[\s_\-\.\/\\]+', s)
            return [t for t in tokens if t]

        def _is_noise(token: str) -> bool:
            """是否是噪声词（年份、数字、变化类词）"""
            t = token.lower()
            if _re.match(r'^\d+$', t):          # 纯数字
                return True
            if _re.match(r'^(19|20)\d{2}$', t): # 年份
                return True
            return t in _NOISE

        def _theme_tokens(col: str) -> tuple:
            """返回去噪后的 token tuple，作为分组 key"""
            tokens = _tokenize(col)
            cleaned = tuple(t.lower() for t in tokens if not _is_noise(t))
            return cleaned if cleaned else (col.lower(),)

        def _longest_common_prefix(cols: list) -> str:
            """获取一组列名的最长公共字符前缀（按空格/分隔符对齐）"""
            if not cols:
                return ""
            token_lists = [_tokenize(c) for c in cols]
            prefix = []
            for tokens in zip(*token_lists):
                if len(set(t.lower() for t in tokens)) == 1:
                    prefix.append(tokens[0])
                else:
                    break
            return " ".join(prefix)

        # ── 策略1+4：去噪后的主题词分组（覆盖年份/数字/变化词）──────────────
        theme_raw: dict = _dd(list)
        for col in num_cols:
            key = _theme_tokens(col)
            theme_raw[key].append(col)

        # ── 策略2：公共后缀分组 ─────────────────────────────────────────────
        # 把 token list 反转后做前缀分组，找到有公共后缀的列
        suffix_raw: dict = _dd(list)
        for col in num_cols:
            tokens = _tokenize(col)
            # 取最后1个非噪声 token 作为后缀 key
            noise_free = [t for t in tokens if not _is_noise(t)]
            if noise_free:
                suffix_raw[noise_free[-1].lower()].append(col)

        # ── 策略3：分隔符结构（A_before/A_after, B_before/B_after）───────────
        # 找到 token 数 ≥2 且第一个 token 相同的列组
        first_token_raw: dict = _dd(list)
        for col in num_cols:
            tokens = _tokenize(col)
            if len(tokens) >= 2:
                first_token_raw[tokens[0].lower()].append(col)

        # ── 合并所有策略结果 ────────────────────────────────────────────────
        # 收集所有"≥2个变量"的候选组，用 frozenset 去重
        seen_sets: set = set()
        candidate_groups: list = []   # list of (label, [cols])

        def _add_group(label: str, cols: list):
            key = frozenset(cols)
            if key not in seen_sets and len(cols) >= 2:
                seen_sets.add(key)
                candidate_groups.append((label, sorted(cols)))

        # 策略1+4 结果
        for key, cols in theme_raw.items():
            if len(cols) >= 2:
                # 用公共前缀作为 label；若全是单词则用 key
                prefix = _longest_common_prefix(cols)
                label  = prefix if prefix else " ".join(key).title()
                _add_group(label, cols)

        # 策略2 结果（公共后缀，仅在策略1没覆盖时补充）
        for suffix, cols in suffix_raw.items():
            if len(cols) >= 2:
                prefix = _longest_common_prefix(cols)
                label  = prefix if prefix else f"[{suffix.title()} 系列]"
                _add_group(label, cols)

        # 策略3 结果（首词相同，仅在前两种没覆盖时补充）
        for first, cols in first_token_raw.items():
            if len(cols) >= 2:
                _add_group(first.title(), cols)

        # 按组内变量数降序排列（越大的组越靠前）
        candidate_groups.sort(key=lambda x: -len(x[1]))

        # 构建最终 group_labels_map（label → cols）
        group_labels_map = {}
        for label, cols in candidate_groups:
            # 避免 label 重复
            final_label = label
            i = 2
            while final_label in group_labels_map:
                final_label = f"{label} ({i})"
                i += 1
            group_labels_map[final_label] = cols

        # ══════════════════════════════════════════════════════════════════════
        # UI 渲染
        # ══════════════════════════════════════════════════════════════════════
        if group_labels_map:
            st.markdown(
                f'<div style="font-size:0.78rem;color:#6b7a9a;margin-bottom:0.6rem;">'
                + T["s2_group_detect"].format(n=len(num_cols), g=len(group_labels_map))
                + '</div>',
                unsafe_allow_html=True,
            )

            # 组成员预览卡片
            for gname, gcols in group_labels_map.items():
                st.markdown(
                    f'<div style="font-size:0.74rem;color:#3d4a6a;margin-bottom:0.2rem;'
                    f'padding:0.15rem 0;">'
                    f'<span style="color:#5b8dee;font-weight:600;">{gname}</span>'
                    f'<span style="color:#3d4a6a;margin:0 0.3rem;">→</span>'
                    f'{" · ".join(gcols)}</div>',
                    unsafe_allow_html=True,
                )

            group_options_list = list(group_labels_map.keys())

            # AI 推荐默认组（包含 ai_dep / ai_indep 的优先）
            _ai_vars = set(([ai_dep] if ai_dep else []) + ai_indep)
            default_groups = [
                g for g, cols in group_labels_map.items()
                if any(v in cols for v in _ai_vars)
            ] or group_options_list[:1]

            selected_groups = st.multiselect(
                T["s2_group_sel"],
                group_options_list,
                default=default_groups,
                key="sel_var_groups",
                help=T.get("s2_group_sel_help","选中一个组 = 自动选中该组的全部变量"),
            )

            # 展开 → target_vars
            expanded_vars = list(dict.fromkeys(
                v for g in selected_groups for v in group_labels_map.get(g, [])
            ))

            # 额外追加
            extra = st.multiselect(
                T["s2_group_extra"],
                [v for v in num_cols if v not in expanded_vars],
                default=[],
                key="sel_extra_vars",
            )
            cfg["target_vars"] = expanded_vars + extra

            if cfg["target_vars"]:
                st.markdown(
                    '<div style="font-size:0.74rem;color:#3dd68c;margin-top:0.4rem;">'
                    + T["s2_group_added"].format(n=len(cfg["target_vars"]), vars=" · ".join(cfg["target_vars"]))
                    + '</div>',
                    unsafe_allow_html=True,
                )
            else:
                cfg["target_vars"] = num_cols[:min(2, len(num_cols))]

        else:
            # 完全无法识别分组 → 成对选择模式
            st.markdown(
                '<div style="font-size:0.78rem;color:#6b7a9a;margin-bottom:0.5rem;">'
                + T["s2_pair_hint"]
                + '</div>',
                unsafe_allow_html=True,
            )
            c_a, c_b = st.columns(2)
            with c_a:
                var_a = st.selectbox(T["s2_pair_a"], num_cols, index=0, key="sel_pair_a")
            with c_b:
                remaining = [v for v in num_cols if v != var_a]
                var_b = st.selectbox(T["s2_pair_b"], remaining, index=0, key="sel_pair_b")
            extra_pair = st.multiselect(
                T["s2_pair_extra"],
                [v for v in num_cols if v not in [var_a, var_b]],
                default=[], key="sel_pair_extra",
            )
            cfg["target_vars"] = [var_a, var_b] + extra_pair

        # ── 分组维度 ─────────────────────────────────────────────────────────
        cat_cols = df.select_dtypes(exclude=np.number).columns.tolist()
        group_options = cat_cols if cat_cols else all_cols
        ai_group = ai_groups[0] if ai_groups and ai_groups[0] in group_options else None
        cfg["group_var"] = st.selectbox(
            T["s2_group_var"],
            [T["s2_no_group"]] + group_options,
            index=(group_options.index(ai_group) + 1) if ai_group else 0,
            key="sel_group_var",
        )
        cfg["group_var"] = None if cfg["group_var"] == T["s2_no_group"] else cfg["group_var"]

        # ── 时间变量 ─────────────────────────────────────────────────────────
        time_candidates = [c for c in all_cols if any(
            kw in c.lower() for kw in
            ["time", "date", "year", "month", "week", "quarter",
             "日期", "时间", "年", "月", "季度"]
        )]
        if time_candidates:
            cfg["time_var"] = st.selectbox(
                T["s2_time_var"],
                [T["s2_no_time"]] + time_candidates,
                key="sel_time_var",
            )
            cfg["time_var"] = None if cfg["time_var"] == T["s2_no_time"] else cfg["time_var"]
        else:
            cfg["time_var"] = None

    # ── 相关性分析变量 ──────────────────────────────────────────────────────
    if "相关性分析" in _chosen_zh:
        st.markdown(
            f'<div style="font-size:0.75rem;color:#5b8dee;font-weight:600;margin:0.8rem 0 0.4rem 0;">{T["s2_corr_vars"]}</div>',
            unsafe_allow_html=True,
        )
        default_corr = [v for v in (([ai_dep] if ai_dep else []) + ai_indep) if v in num_cols]
        if not default_corr:
            default_corr = num_cols[:min(5, len(num_cols))]
        cfg["corr_vars"] = st.multiselect(
            "参与相关性分析的变量（建议 2-8 个数值变量）",
            num_cols,
            default=default_corr,
            key="sel_corr_vars",
            help=T["s2_corr_help"],
        )

    # ── 回归分析变量 ────────────────────────────────────────────────────────
    if "回归分析" in _chosen_zh:
        st.markdown(
            f'<div style="font-size:0.75rem;color:#5b8dee;font-weight:600;margin:0.8rem 0 0.4rem 0;">{T["s2_reg_vars"]}</div>',
            unsafe_allow_html=True,
        )
        y_default = num_cols.index(ai_dep) if ai_dep in num_cols else 0
        cfg["reg_y"] = st.selectbox(
            T["s2_dep_var"],
            num_cols,
            index=y_default,
            key="sel_reg_y",
            help="AI 推荐：" + (ai_dep or "未指定"),
        )
        x_options = [c for c in num_cols if c != cfg["reg_y"]]
        ai_x_valid = [v for v in ai_indep if v in x_options]
        if not ai_x_valid:
            ai_x_valid = x_options[:min(3, len(x_options))]
        cfg["reg_x"] = st.multiselect(
            T["s2_indep_vars"],
            x_options,
            default=ai_x_valid,
            key="sel_reg_x",
        )
        if cfg["reg_x"]:
            cfg["scatter_x"] = st.selectbox(
                T["s2_scatter_sel"],
                cfg["reg_x"],
                key="sel_scatter_x",
            )
        else:
            cfg["scatter_x"] = None

    # ══════════════════════════════════════════════════════════════════════════
    # ── 图表选择模块（Visualization Configuration）────────────────────────────
    # ══════════════════════════════════════════════════════════════════════════
    st.markdown("<hr class='divider'>", unsafe_allow_html=True)
    st.markdown(
        f'<div style="font-size:0.75rem;color:#5b8dee;font-weight:600;margin:0.8rem 0 0.4rem 0;">📊 {T.get("s2_chart_title","图表选择")}</div>',
        unsafe_allow_html=True,
    )
    st.markdown(
        f'<div style="font-size:0.8rem;color:#6b7a9a;margin-bottom:0.8rem;">{T.get("s2_chart_hint","先选图表类型，系统将自动约束可选变量")}</div>',
        unsafe_allow_html=True,
    )

    # ── 变量分类（约束引擎基础）──────────────────────────────────────────────
    nc   = df.select_dtypes(include=np.number).columns.tolist()
    ac   = df.columns.tolist()
    cat_c = df.select_dtypes(exclude=np.number).columns.tolist()
    time_c = [c for c in ac if any(
        kw in c.lower() for kw in
        ["time","date","year","month","week","quarter","day",
         "日期","时间","年","月","季度","周"]
    )]
    low_card_num = [c for c in nc if df[c].nunique() <= 20]
    cat_or_low   = list(dict.fromkeys(cat_c + low_card_num))
    x_cat_pool   = cat_or_low if cat_or_low else ac

    _default_y = (cfg.get("reg_y")
                  or (cfg.get("target_vars") or [None])[0]
                  or (nc[0] if nc else None))
    _default_x_cat = (cat_c[0] if cat_c else (low_card_num[0] if low_card_num else
                      (ac[0] if ac else None)))
    _default_x_num = (nc[1] if len(nc) > 1 else (nc[0] if nc else None))

    # ── 全量图表类型 ──────────────────────────────────────────────────────────
    # 图表类型列表：始终用中文内部名 → 显示时用翻译
    _all_chart_zh = [
        "柱状图", "条形图", "分组柱状图",
        "折线图", "面积图",
        "散点图", "气泡图",
        "饼图", "圆环图",
        "箱线图", "小提琴图",
        "雷达图",
        "热力图",
        "变化排序图", "组合图（柱+线）",
    ]
    _ct_zh2label = {
        "柱状图":T["ct_bar"],"条形图":T["ct_hbar"],"分组柱状图":T["ct_gbar"],
        "折线图":T["ct_line"],"面积图":T["ct_area"],"散点图":T["ct_scatter"],
        "气泡图":T["ct_bubble"],"饼图":T["ct_pie"],"圆环图":T["ct_donut"],
        "箱线图":T["ct_box"],"小提琴图":T["ct_violin"],"雷达图":T["ct_radar"],
        "热力图":T["ct_heatmap"],"变化排序图":T["ct_rank"],"组合图（柱+线）":T["ct_combo"],
    }
    _ct_label2zh = {v:k for k,v in _ct_zh2label.items()}
    all_chart_types = [_ct_zh2label[c] for c in _all_chart_zh]

    # ── 动态推荐逻辑：结合分析类型 + 数据特征 + 研究问题关键词 ────────────────
    _q = st.session_state.get("question", "").lower()

    def _smart_defaults(chosen: list, question_lower: str, has_time: bool,
                        has_cat: bool, n_num: int) -> list:
        recs = []
        # 按分析类型推荐
        if "描述性统计" in chosen:
            recs += ["柱状图", "箱线图"]
        if "对比分析" in chosen:
            recs += ["分组柱状图", "条形图", "箱线图"]
            if has_cat:
                recs += ["小提琴图"]
        if "趋势分析" in chosen:
            recs += ["折线图", "面积图"]
        if "相关性分析" in chosen:
            recs += ["散点图", "热力图"]
            if n_num >= 3:
                recs += ["气泡图"]
        if "回归分析" in chosen:
            recs += ["散点图", "变化排序图"]
        # 按问题关键词追加推荐
        if any(kw in question_lower for kw in ["占比","比例","构成","份额","percent","share","proportion"]):
            recs += ["饼图", "圆环图"]
        if any(kw in question_lower for kw in ["趋势","变化","增长","下降","trend","growth","change"]):
            recs += ["折线图", "面积图"]
        if any(kw in question_lower for kw in ["对比","比较","差异","compare","versus","vs"]):
            recs += ["分组柱状图", "条形图"]
        if any(kw in question_lower for kw in ["分布","distribution","spread","range"]):
            recs += ["箱线图", "小提琴图"]
        if any(kw in question_lower for kw in ["多维","综合","整体","overview","综合对比"]):
            recs += ["雷达图"]
        if any(kw in question_lower for kw in ["叠加","组合","双轴","combined"]):
            recs += ["组合图（柱+线）"]
        # 去重保序，过滤不在列表的
        return list(dict.fromkeys([c for c in recs if c in all_chart_types])) or ["柱状图"]

    # _smart_defaults 内部用中文类型名判断，返回结果映射到当前语言
    _zh2ct = {"柱状图":"Bar Chart","条形图":"Horizontal Bar","分组柱状图":"Grouped Bar",
              "折线图":"Line Chart","面积图":"Area Chart","散点图":"Scatter Plot",
              "气泡图":"Bubble Chart","饼图":"Pie Chart","圆环图":"Donut Chart",
              "箱线图":"Box Plot","小提琴图":"Violin Plot","雷达图":"Radar Chart",
              "热力图":"Heatmap","变化排序图":"Ranked Change Chart","组合图（柱+线）":"Combined Chart"}
    _default_charts_zh = _smart_defaults(
        _chosen_zh, _q,
        has_time=bool(time_c),
        has_cat=bool(cat_c),
        n_num=len(nc),
    )
    default_charts = [_ct_zh2label.get(c,c) for c in _default_charts_zh if _ct_zh2label.get(c,c) in all_chart_types]
    if not default_charts:
        default_charts = all_chart_types[:1]

    chart_types_sel = st.multiselect(
        T["chart_sel_label"],
        all_chart_types,
        default=default_charts,
        key="sel_chart_types",
        help="选择后下方自动显示该图表所需的变量配置",
    )
    st.markdown("<div style='height:0.3rem'></div>", unsafe_allow_html=True)

    # ── 约束规则说明 ─────────────────────────────────────────────────────────
    # chart_rules key 需要用当前语言的图表显示名
    _rules_raw = T.get("chart_rules", {})
    _lang_chart = st.session_state.get("lang", "zh")
    _CHART_RULES = {_ct_zh2label.get(k, k): v for k, v in _rules_raw.items()} if _lang_chart != "en" else _rules_raw
    for ct in chart_types_sel:
        st.markdown(
            f'<div style="font-size:0.7rem;color:#3d4a6a;margin-bottom:0.15rem;">'
            f'<span style="color:#5b8dee;font-weight:600;">{ct}</span>'
            f' &nbsp;—&nbsp; {_CHART_RULES.get(ct,"")}</div>',
            unsafe_allow_html=True,
        )
    if chart_types_sel:
        st.markdown("<div style='height:0.4rem'></div>", unsafe_allow_html=True)

    ccfg = {}

    # ── 1. 柱状图 ─────────────────────────────────────────────────────────────
    if "柱状图" in chart_types_sel:
        st.markdown(f'<div style="font-size:0.72rem;color:#a8b4d0;font-weight:600;margin:0.8rem 0 0.3rem 0;">🔵 {T["ct_bar"]}</div>', unsafe_allow_html=True)
        if not x_cat_pool:
            st.warning(T["chart_warn_nocat"])
        else:
            c1, c2 = st.columns(2)
            with c1:
                if not cat_c and low_card_num:
                    st.caption(T["chart_warn_nocat2"])
                ccfg["bar_x"] = st.selectbox(T["chart_x_cat"], x_cat_pool,
                    index=x_cat_pool.index(_default_x_cat) if _default_x_cat in x_cat_pool else 0,
                    key="chart_bar_x")
            with c2:
                y_pool = [v for v in nc if v != ccfg.get("bar_x")]
                ccfg["bar_y"] = st.selectbox(T["chart_y_num"], y_pool or nc,
                    index=y_pool.index(_default_y) if _default_y in y_pool else 0,
                    key="chart_bar_y")

    # ── 2. 条形图（横向柱状图）───────────────────────────────────────────────
    if "条形图" in chart_types_sel:
        st.markdown(f'<div style="font-size:0.72rem;color:#a8b4d0;font-weight:600;margin:0.8rem 0 0.3rem 0;">↔ {T["ct_hbar"]}</div>', unsafe_allow_html=True)
        c1, c2 = st.columns(2)
        with c1:
            ccfg["hbar_y"] = st.selectbox(T["chart_y_cat"], x_cat_pool or ac,
                index=0, key="chart_hbar_y")
        with c2:
            y_pool_hb = [v for v in nc if v != ccfg.get("hbar_y")]
            ccfg["hbar_x"] = st.selectbox(T["chart_x_num"], y_pool_hb or nc,
                index=y_pool_hb.index(_default_y) if _default_y in y_pool_hb else 0,
                key="chart_hbar_x")

    # ── 3. 分组柱状图 ─────────────────────────────────────────────────────────
    if "分组柱状图" in chart_types_sel:
        st.markdown(f'<div style="font-size:0.72rem;color:#a8b4d0;font-weight:600;margin:0.8rem 0 0.3rem 0;">🟦 {T["ct_gbar"]}</div>', unsafe_allow_html=True)
        c1, c2, c3 = st.columns(3)
        with c1:
            ccfg["gbar_x"] = st.selectbox(T["chart_x_cat"], x_cat_pool or ac,
                index=x_cat_pool.index(_default_x_cat) if _default_x_cat in x_cat_pool else 0,
                key="chart_gbar_x")
        with c2:
            y_pool_gb = [v for v in nc if v != ccfg.get("gbar_x")]
            ccfg["gbar_y"] = st.selectbox(T["chart_y_num"], y_pool_gb or nc,
                index=y_pool_gb.index(_default_y) if _default_y in y_pool_gb else 0,
                key="chart_gbar_y")
        with c3:
            grp_pool = [v for v in cat_or_low if v != ccfg.get("gbar_x")]
            ccfg["gbar_group"] = st.selectbox(T["chart_group_opt"], [T["s2_no_group"]] + grp_pool, key="chart_gbar_group")
            ccfg["gbar_group"] = None if ccfg["gbar_group"] == T["s2_no_group"] else ccfg["gbar_group"]

    # ── 4. 折线图 ─────────────────────────────────────────────────────────────
    if "折线图" in chart_types_sel:
        st.markdown(f'<div style="font-size:0.72rem;color:#a8b4d0;font-weight:600;margin:0.8rem 0 0.3rem 0;">🟢 {T["ct_line"]}</div>', unsafe_allow_html=True)
        x_line_pool = time_c + [c for c in cat_c if c not in time_c] + [c for c in ac if c not in time_c and c not in cat_c]
        x_line_pool = x_line_pool or ac
        c1, c2 = st.columns(2)
        with c1:
            if not time_c:
                st.caption(T["chart_warn_notime"])
            ccfg["line_x"] = st.selectbox(T["chart_x_time"], x_line_pool, index=0, key="chart_line_x")
        with c2:
            y_pool_ln = [v for v in nc if v != ccfg.get("line_x")]
            ccfg["line_y"] = st.multiselect(T["chart_y_multi"],
                y_pool_ln or nc,
                default=[_default_y] if _default_y in (y_pool_ln or nc) else (y_pool_ln or nc)[:1],
                key="chart_line_y")
            if len(ccfg["line_y"]) > 4:
                st.warning(T["chart_warn_line4"])
                ccfg["line_y"] = ccfg["line_y"][:4]

    # ── 5. 面积图 ─────────────────────────────────────────────────────────────
    if "面积图" in chart_types_sel:
        st.markdown(f'<div style="font-size:0.72rem;color:#a8b4d0;font-weight:600;margin:0.8rem 0 0.3rem 0;">🌊 {T["ct_area"]}</div>', unsafe_allow_html=True)
        x_area_pool = time_c + [c for c in ac if c not in time_c]
        c1, c2 = st.columns(2)
        with c1:
            ccfg["area_x"] = st.selectbox(T["chart_x_time"], x_area_pool or ac, index=0, key="chart_area_x")
        with c2:
            y_pool_ar = [v for v in nc if v != ccfg.get("area_x")]
            ccfg["area_y"] = st.multiselect(T["chart_y_multi"],
                y_pool_ar or nc,
                default=[_default_y] if _default_y in (y_pool_ar or nc) else (y_pool_ar or nc)[:1],
                key="chart_area_y")

    # ── 6. 散点图 ─────────────────────────────────────────────────────────────
    if "散点图" in chart_types_sel:
        st.markdown(f'<div style="font-size:0.72rem;color:#a8b4d0;font-weight:600;margin:0.8rem 0 0.3rem 0;">🟡 {T["ct_scatter"]}</div>', unsafe_allow_html=True)
        if len(nc) < 2:
            st.warning(T["chart_warn_sc2"])
        else:
            c1, c2 = st.columns(2)
            with c1:
                ccfg["scatter_x"] = st.selectbox(T["chart_x_num2"], nc,
                    index=nc.index(_default_x_num) if _default_x_num in nc else 0,
                    key="chart_scatter_x")
            with c2:
                y_pool_sc = [v for v in nc if v != ccfg["scatter_x"]]
                ccfg["scatter_y"] = st.selectbox(T["chart_y_num2"], y_pool_sc,
                    index=y_pool_sc.index(_default_y) if _default_y in y_pool_sc else 0,
                    key="chart_scatter_y")

    # ── 7. 气泡图 ─────────────────────────────────────────────────────────────
    if "气泡图" in chart_types_sel:
        st.markdown(f'<div style="font-size:0.72rem;color:#a8b4d0;font-weight:600;margin:0.8rem 0 0.3rem 0;">🫧 {T["ct_bubble"]}</div>', unsafe_allow_html=True)
        if len(nc) < 3:
            st.warning(T["chart_warn_bub3"])
        else:
            c1, c2, c3 = st.columns(3)
            with c1:
                ccfg["bubble_x"] = st.selectbox(T["chart_x_num2"], nc, index=0, key="chart_bubble_x")
            with c2:
                nc_bub_y = [v for v in nc if v != ccfg["bubble_x"]]
                ccfg["bubble_y"] = st.selectbox(T["chart_bubble_y"], nc_bub_y, index=0, key="chart_bubble_y")
            with c3:
                nc_bub_s = [v for v in nc if v not in [ccfg["bubble_x"], ccfg["bubble_y"]]]
                ccfg["bubble_size"] = st.selectbox(T["chart_bubble_size"],
                    nc_bub_s if nc_bub_s else nc, index=0, key="chart_bubble_size")

    # ── 8. 饼图 ───────────────────────────────────────────────────────────────
    if "饼图" in chart_types_sel:
        st.markdown(f'<div style="font-size:0.72rem;color:#a8b4d0;font-weight:600;margin:0.8rem 0 0.3rem 0;">🥧 {T["ct_pie"]}</div>', unsafe_allow_html=True)
        c1, c2 = st.columns(2)
        with c1:
            ccfg["pie_label"] = st.selectbox(T["chart_pie_label"],
                cat_or_low or ac, index=0, key="chart_pie_label")
        with c2:
            ccfg["pie_value"] = st.selectbox(T["chart_pie_value"],
                nc, index=nc.index(_default_y) if _default_y in nc else 0,
                key="chart_pie_value")

    # ── 9. 圆环图 ─────────────────────────────────────────────────────────────
    if "圆环图" in chart_types_sel:
        st.markdown(f'<div style="font-size:0.72rem;color:#a8b4d0;font-weight:600;margin:0.8rem 0 0.3rem 0;">⭕ {T["ct_donut"]}</div>', unsafe_allow_html=True)
        c1, c2 = st.columns(2)
        with c1:
            ccfg["donut_label"] = st.selectbox(T["chart_donut_label"], cat_or_low or ac, index=0, key="chart_donut_label")
        with c2:
            ccfg["donut_value"] = st.selectbox(T["chart_donut_value"], nc,
                index=nc.index(_default_y) if _default_y in nc else 0, key="chart_donut_value")

    # ── 10. 箱线图 ────────────────────────────────────────────────────────────
    if "箱线图" in chart_types_sel:
        st.markdown(f'<div style="font-size:0.72rem;color:#a8b4d0;font-weight:600;margin:0.8rem 0 0.3rem 0;">🔴 {T["ct_box"]}</div>', unsafe_allow_html=True)
        c1, c2 = st.columns(2)
        with c1:
            ccfg["box_var"] = st.selectbox(T["chart_box_num"], nc,
                index=nc.index(_default_y) if _default_y in nc else 0, key="chart_box_var")
        with c2:
            box_grp = [v for v in cat_or_low if v != ccfg.get("box_var")]
            ccfg["box_group"] = st.selectbox(T["chart_group_opt"], [T["s2_no_group"]] + box_grp, key="chart_box_group")
            ccfg["box_group"] = None if ccfg["box_group"] == T["s2_no_group"] else ccfg["box_group"]

    # ── 11. 小提琴图 ──────────────────────────────────────────────────────────
    if "小提琴图" in chart_types_sel:
        st.markdown(f'<div style="font-size:0.72rem;color:#a8b4d0;font-weight:600;margin:0.8rem 0 0.3rem 0;">🎻 {T["ct_violin"]}</div>', unsafe_allow_html=True)
        c1, c2 = st.columns(2)
        with c1:
            ccfg["violin_var"] = st.selectbox(T["chart_box_num"], nc,
                index=nc.index(_default_y) if _default_y in nc else 0, key="chart_violin_var")
        with c2:
            vio_grp = [v for v in cat_or_low if v != ccfg.get("violin_var")]
            ccfg["violin_group"] = st.selectbox(T["chart_group_opt"], [T["s2_no_group"]] + vio_grp, key="chart_violin_group")
            ccfg["violin_group"] = None if ccfg["violin_group"] == T["s2_no_group"] else ccfg["violin_group"]

    # ── 12. 雷达图 ────────────────────────────────────────────────────────────
    if "雷达图" in chart_types_sel:
        st.markdown(f'<div style="font-size:0.72rem;color:#a8b4d0;font-weight:600;margin:0.8rem 0 0.3rem 0;">🕸 {T["ct_radar"]}</div>', unsafe_allow_html=True)
        ccfg["radar_vars"] = st.multiselect(T["chart_radar_vars"], nc,
            default=nc[:min(5, len(nc))], key="chart_radar_vars",
            help="每个变量对应雷达图的一个维度")
        if len(ccfg.get("radar_vars", [])) < 3:
            st.warning(T["chart_warn_radar3"])
        vio_grp_r = [v for v in cat_or_low if v not in ccfg.get("radar_vars", [])]
        ccfg["radar_group"] = st.selectbox(T["chart_radar_grp"],
            [T["s2_no_group"]] + vio_grp_r, key="chart_radar_group")
        ccfg["radar_group"] = None if ccfg["radar_group"] == T["s2_no_group"] else ccfg["radar_group"]

    # ── 13. 热力图 ────────────────────────────────────────────────────────────
    if "热力图" in chart_types_sel:
        st.markdown(f'<div style="font-size:0.72rem;color:#a8b4d0;font-weight:600;margin:0.8rem 0 0.3rem 0;">🌡 {T["ct_heatmap"]}</div>', unsafe_allow_html=True)
        st.caption(T["chart_heat_caption"])
        hm_opts = ["自动（使用相关矩阵）"] + nc
        ccfg["heatmap_vars"] = st.multiselect(T["chart_heat_sel"],
            nc, default=[], key="chart_heatmap_vars")

    # ── 14. 变化排序图 ────────────────────────────────────────────────────────
    if "变化排序图" in chart_types_sel:
        st.markdown(f'<div style="font-size:0.72rem;color:#a8b4d0;font-weight:600;margin:0.8rem 0 0.3rem 0;">🟠 {T["ct_rank"]}</div>', unsafe_allow_html=True)
        ccfg["rank_var"] = st.selectbox(T["chart_rank_var"], nc,
            index=nc.index(_default_y) if _default_y in nc else 0, key="chart_rank_var")

    # ── 15. 组合图（柱+线）────────────────────────────────────────────────────
    if "组合图（柱+线）" in chart_types_sel:
        st.markdown(f'<div style="font-size:0.72rem;color:#a8b4d0;font-weight:600;margin:0.8rem 0 0.3rem 0;">📊+📈 {T["ct_combo"]}</div>', unsafe_allow_html=True)
        c1, c2, c3 = st.columns(3)
        with c1:
            ccfg["combo_x"] = st.selectbox(T["chart_combo_x"],
                (time_c + cat_or_low) or ac, index=0, key="chart_combo_x")
        with c2:
            nc_combo = [v for v in nc if v != ccfg.get("combo_x")]
            ccfg["combo_bar_y"] = st.selectbox(T["chart_combo_bar"],
                nc_combo or nc, index=0, key="chart_combo_bar_y")
        with c3:
            nc_line = [v for v in nc if v not in [ccfg.get("combo_x"), ccfg.get("combo_bar_y")]]
            ccfg["combo_line_y"] = st.selectbox(T["chart_combo_line"],
                nc_line or nc, index=0, key="chart_combo_line_y")

    # ── 校验 ─────────────────────────────────────────────────────────────────
    if "折线图" in chart_types_sel and not ccfg.get("line_y"):
        st.error(T["chart_warn_min1"])
    if "散点图" in chart_types_sel and len(nc) >= 2:
        if not ccfg.get("scatter_x") or not ccfg.get("scatter_y"):
            st.error(T["chart_warn_sc_xy"])
    if "雷达图" in chart_types_sel and len(ccfg.get("radar_vars", [])) < 3:
        st.error(T["chart_warn_radar_e"])

    st.markdown("</div>", unsafe_allow_html=True)

    # ── 导航按钮 ──────────────────────────────────────────────────────────────
    col_back, col_regen, col_confirm = st.columns([1, 1, 1])
    with col_back:
        if st.button(T["s2_back"]):
            st.session_state.step = 1
            st.rerun()
    with col_regen:
        if st.button(T["s2_regen"]):
            st.session_state.plan = None
            st.rerun()
    with col_confirm:
        if st.button(T["s2_confirm"], width='stretch'):
            if not chosen_types:
                st.error(T["s2_err_type"])
            else:
                # 写入 session
                # 始终以中文存储，保证后续判断逻辑一致
                st.session_state.confirmed_analysis_types = [_en2zh.get(t,t) for t in chosen_types]
                st.session_state.confirmed_desc_vars      = cfg.get("desc_vars", [])
                st.session_state.confirmed_target_vars    = cfg.get("target_vars", [])
                st.session_state.confirmed_group_var      = cfg.get("group_var")
                st.session_state.confirmed_time_var       = cfg.get("time_var")
                st.session_state.confirmed_corr_vars      = cfg.get("corr_vars", [])
                st.session_state.confirmed_y              = cfg.get("reg_y")
                st.session_state.confirmed_x              = cfg.get("reg_x", [])
                st.session_state.confirmed_scatter        = cfg.get("scatter_x")
                # 图表类型始终存中文
                st.session_state.chart_types = [_ct_label2zh.get(t, t) for t in chart_types_sel]
                st.session_state.chart_config             = ccfg
                st.session_state.analysis_done            = False
                st.session_state.reg_results              = None
                st.session_state.step                     = 3
                st.rerun()


# ══════════════════════════════════════════════════════════════════════════════
# STEP 3 · 分析结果与解释
# ══════════════════════════════════════════════════════════════════════════════
elif st.session_state.step == 3:
    T = _get_T()
    phase_header("03", T["s3_title"])

    question      = st.session_state.question
    df            = st.session_state.df
    chosen_types_raw = st.session_state.confirmed_analysis_types
    # 统一转成中文类型名用于逻辑判断（兼容中英文切换后的存储值）
    _type_map_en2zh = {"Descriptive Stats":"描述性统计","Comparison":"对比分析","Trend Analysis":"趋势分析","Correlation":"相关性分析","Regression":"回归分析"}
    chosen_types = [_type_map_en2zh.get(t,t) for t in chosen_types_raw]
    miss_pct_all  = df.isnull().sum().sum() / df.size * 100

    st.markdown(
        f'<div class="question-display">{T["s3_q_prefix"]}{question}</div>',
        unsafe_allow_html=True,
    )

    # 用于最终结论传递
    sig_vars_list   = []
    insig_vars_list = []
    desc_summary    = ""
    reg_r2          = 0.0
    _chart_images   = []  # 收集本次分析所有图表路径，用于报告生成

    def sub_header(label: str):
        st.markdown(
            f'<div style="font-size:1.05rem;font-weight:700;color:#dde2f0;'
            f'margin:1.6rem 0 0.8rem 0;padding-bottom:0.4rem;'
            f'border-bottom:2px solid #2a3a5a;display:flex;align-items:center;gap:0.5rem;">'
            f'<span style="color:#5b8dee;font-size:0.9rem;">▪</span> {label}</div>',
            unsafe_allow_html=True,
        )

    # ── A · 描述性统计 ────────────────────────────────────────────────────────
    if "描述性统计" in chosen_types:
        sub_header(T["s3_sec_a"])
        desc_vars = st.session_state.confirmed_desc_vars or \
                    df.select_dtypes(include=np.number).columns.tolist()[:4]
        num_df  = df[desc_vars].select_dtypes(include=np.number)
        # 内部始终用固定英文键，display用翻译列名
        desc    = num_df.describe().T
        desc["_skew"] = num_df.skew()
        desc["_kurt"] = num_df.kurt()
        desc.columns = ["_n","_mean","_std","_min","_p25","_med","_p75","_max","_skew","_kurt"]
        desc = desc.round(4)
        # 展示时用翻译列名
        desc_display = desc.copy()
        desc_display.columns = T["s3_desc_cols"]

        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.dataframe(desc_display, width='stretch')
        st.markdown("</div>", unsafe_allow_html=True)

        focus_var    = desc_vars[0] if desc_vars else ""
        stats_snip   = {}
        for c in desc_vars:
            if c in desc.index:
                stats_snip[c] = {
                    "均值": desc.loc[c,"_mean"], "标准差": desc.loc[c,"_std"],
                    "中位数": desc.loc[c,"_med"], "偏度": desc.loc[c,"_skew"],
                    "峰度": desc.loc[c,"_kurt"],
                }
        fb_desc = fallback_desc(stats_snip, focus_var, lang=st.session_state.get("lang","zh"))
        desc_summary = "; ".join([
            f"{c} mean={desc.loc[c,'_mean']:.2f} skew={desc.loc[c,'_skew']:.2f}"
            for c in desc_vars if c in desc.index
        ])

        with st.spinner(T["s3_spin_desc"]):
            text, is_ai = ai_desc_interp(
                question,
                json.dumps(stats_snip, ensure_ascii=False, default=str),
                focus_var,
                fb_desc,
                lang=st.session_state.get("lang","zh"),
            )
        render_ai(T["s3_lbl_desc"], text, is_ai)
        st.markdown("<hr class='divider'>", unsafe_allow_html=True)

    # ── B · 对比分析 / 趋势分析 ───────────────────────────────────────────────
    if "对比分析" in chosen_types or "趋势分析" in chosen_types:
        label_bt = T["s3_sec_b3"]
        if "对比分析" in chosen_types and "趋势分析" not in chosen_types:
            label_bt = T["s3_sec_b"]
        elif "趋势分析" in chosen_types and "对比分析" not in chosen_types:
            label_bt = T["s3_sec_b2"]
        sub_header(label_bt)

        target_vars = st.session_state.confirmed_target_vars
        group_var   = st.session_state.confirmed_group_var
        time_var    = st.session_state.get("confirmed_time_var")
        chart_types_raw = st.session_state.get("chart_types", [])
        # 图表类型名映射为中文用于判断（支持英文模式下的存储值）
        _ct_en2zh = {
            "Bar Chart":"柱状图","Horizontal Bar":"条形图","Grouped Bar":"分组柱状图",
            "Line Chart":"折线图","Area Chart":"面积图","Scatter Plot":"散点图",
            "Bubble Chart":"气泡图","Pie Chart":"饼图","Donut Chart":"圆环图",
            "Box Plot":"箱线图","Violin Plot":"小提琴图","Radar Chart":"雷达图",
            "Heatmap":"热力图","Ranked Change Chart":"变化排序图","Combined Chart":"组合图（柱+线）",
        }
        chart_types = [_ct_en2zh.get(t, t) for t in chart_types_raw]
        ccfg        = st.session_state.get("chart_config", {})

        if not target_vars:
            st.info(T["s3_no_target"])
        else:
            any_chart_drawn = False

            # ── 柱状图 ────────────────────────────────────────────────────────
            if "柱状图" in chart_types and ccfg.get("bar_x") and ccfg.get("bar_y"):
                bx, by = ccfg["bar_x"], ccfg["bar_y"]
                if bx in df.columns and by in df.columns:
                    st.markdown(f"**{T['ct_bar']}：{by} {T['ct_by']} {bx}**")
                    try:
                        grp = df.groupby(bx)[by].mean().reset_index()
                        fig, ax = plt.subplots(figsize=(8, 3.8))
                        apply_plot_style(fig, ax)
                        bars = ax.bar(grp[bx].astype(str), grp[by], color=ACCENT, alpha=0.85)
                        for bar, val in zip(bars, grp[by]):
                            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() * 1.01,
                                    f"{val:.2f}", ha="center", va="bottom", fontsize=8, color=PLOT_FG)
                        ax.set_xlabel(bx); ax.set_ylabel(by); ax.set_title(f"{by} {T["chart_mean"]}（{T["ct_by"]} {bx}）")
                        plt.xticks(rotation=20, ha="right")
                        plt.tight_layout(pad=0.3)
                        _img_path = os.path.join(_TMPDIR, "_chart_1.png")
                        fig.savefig(_img_path, bbox_inches="tight", dpi=150,
                                   facecolor=fig.get_facecolor())
                        _chart_images.append(_img_path)
                        st.pyplot(fig, width='stretch')
                        plt.close()
                        any_chart_drawn = True
                    except Exception as e:
                        st.warning(T["ct_bar"]+T["ct_fail"]+str(e))

            # ── 分组柱状图 ────────────────────────────────────────────────────
            if "分组柱状图" in chart_types and ccfg.get("gbar_x") and ccfg.get("gbar_y"):
                gx, gy, gg = ccfg["gbar_x"], ccfg["gbar_y"], ccfg.get("gbar_group")
                if gx in df.columns and gy in df.columns:
                    st.markdown(f"**{T['ct_gbar']}：{gy} {T['ct_by']} {gx}{(T['ct_grouped']+gg+')') if gg else ''}**")
                    try:
                        import matplotlib
                        if gg and gg in df.columns:
                            pivot = df.groupby([gx, gg])[gy].mean().unstack(fill_value=0)
                            fig, ax = plt.subplots(figsize=(9, 3.8))
                            apply_plot_style(fig, ax)
                            colors_g = [ACCENT, ACCENT2, "#f06b6b", "#f0b86b", "#a56dde"]
                            x_idx = range(len(pivot.index))
                            n_groups = len(pivot.columns)
                            w = 0.8 / n_groups
                            for i, col in enumerate(pivot.columns):
                                offsets = [xi + (i - n_groups/2 + 0.5) * w for xi in x_idx]
                                ax.bar(offsets, pivot[col], width=w * 0.9,
                                       color=colors_g[i % len(colors_g)], alpha=0.85, label=str(col))
                            ax.set_xticks(list(x_idx))
                            ax.set_xticklabels(pivot.index.astype(str), rotation=20, ha="right",
                                               fontsize=8, color=PLOT_FG)
                            ax.legend(facecolor=PLOT_BG, edgecolor=PLOT_GRID,
                                      labelcolor=PLOT_FG, fontsize=8)
                        else:
                            grp = df.groupby(gx)[gy].mean().reset_index()
                            fig, ax = plt.subplots(figsize=(8, 3.8))
                            apply_plot_style(fig, ax)
                            ax.bar(grp[gx].astype(str), grp[gy], color=ACCENT, alpha=0.85)
                        ax.set_xlabel(gx); ax.set_ylabel(gy); ax.set_title(f"{gy} {T["ct_gbar"]}")
                        plt.tight_layout(pad=0.3)
                        _img_path = os.path.join(_TMPDIR, "_chart_2.png")
                        fig.savefig(_img_path, bbox_inches="tight", dpi=150,
                                   facecolor=fig.get_facecolor())
                        _chart_images.append(_img_path)
                        st.pyplot(fig, width='stretch')
                        plt.close()
                        any_chart_drawn = True
                    except Exception as e:
                        st.warning(T["ct_gbar"]+T["ct_fail"]+str(e))

            # ── 折线图 ────────────────────────────────────────────────────────
            if "折线图" in chart_types and ccfg.get("line_x") and ccfg.get("line_y"):
                lx, ly_list = ccfg["line_x"], ccfg["line_y"]
                if lx in df.columns and ly_list:
                    st.markdown(f"**{T['ct_line']}：{'、'.join(ly_list[:4])} {T['ct_trend']} {lx}**")
                    try:
                        df_l = df[[lx] + ly_list[:4]].dropna().sort_values(lx)
                        fig, ax = plt.subplots(figsize=(9, 3.8))
                        apply_plot_style(fig, ax)
                        line_colors = [ACCENT, ACCENT2, "#f06b6b", "#a56dde"]
                        for i, lv in enumerate(ly_list[:4]):
                            ax.plot(df_l[lx], df_l[lv], color=line_colors[i],
                                    linewidth=2, label=lv, alpha=0.9)
                        ax.set_xlabel(lx)
                        ax.legend(facecolor=PLOT_BG, edgecolor=PLOT_GRID,
                                  labelcolor=PLOT_FG, fontsize=8)
                        ax.set_title(f"{'、'.join(ly_list[:4])} {T['ct_line']}")
                        plt.tight_layout(pad=0.3)
                        _img_path = os.path.join(_TMPDIR, "_chart_3.png")
                        fig.savefig(_img_path, bbox_inches="tight", dpi=150,
                                   facecolor=fig.get_facecolor())
                        _chart_images.append(_img_path)
                        st.pyplot(fig, width='stretch')
                        plt.close()
                        any_chart_drawn = True
                    except Exception as e:
                        st.warning(T["ct_line"]+T["ct_fail"]+str(e))

            # ── 箱线图 ────────────────────────────────────────────────────────
            if "箱线图" in chart_types and ccfg.get("box_var"):
                bv, bg = ccfg["box_var"], ccfg.get("box_group")
                if bv in df.columns:
                    st.markdown(f"**{T['ct_box']}：{bv}{'（'+T['ct_by']+' '+bg+' 分组）' if bg else ''}**")
                    try:
                        fig, ax = plt.subplots(figsize=(7, 3.8))
                        apply_plot_style(fig, ax)
                        if bg and bg in df.columns:
                            groups_data = [df[df[bg] == g][bv].dropna().values
                                           for g in df[bg].unique()]
                            group_labels = df[bg].unique().tolist()
                            bp = ax.boxplot(groups_data, patch_artist=True,
                                            medianprops=dict(color=ACCENT2, linewidth=2))
                            colors_b = [ACCENT, "#a56dde", "#3dd68c", "#f06b6b", "#f0b86b"]
                            for patch, color in zip(bp["boxes"], colors_b * 10):
                                patch.set_facecolor(color); patch.set_alpha(0.7)
                            ax.set_xticklabels(group_labels, rotation=20, ha="right",
                                               fontsize=8, color=PLOT_FG)
                            ax.set_title(f"{bv} {T["ct_by"]} {bg} {T["ct_box"]}")
                        else:
                            bp = ax.boxplot(df[bv].dropna().values, patch_artist=True,
                                            medianprops=dict(color=ACCENT2, linewidth=2))
                            bp["boxes"][0].set_facecolor(ACCENT); bp["boxes"][0].set_alpha(0.7)
                            ax.set_xticklabels([bv], color=PLOT_FG)
                            ax.set_title(f"{bv} {T["ct_box"]}")
                        ax.set_ylabel(bv)
                        plt.tight_layout(pad=0.3)
                        _img_path = os.path.join(_TMPDIR, "_chart_4.png")
                        fig.savefig(_img_path, bbox_inches="tight", dpi=150,
                                   facecolor=fig.get_facecolor())
                        _chart_images.append(_img_path)
                        st.pyplot(fig, width='stretch')
                        plt.close()
                        any_chart_drawn = True
                    except Exception as e:
                        st.warning(T["ct_box"]+T["ct_fail"]+str(e))

            # ── 变化排序图 ────────────────────────────────────────────────────
            if "变化排序图" in chart_types and ccfg.get("rank_var"):
                rv = ccfg["rank_var"]
                if rv in df.columns:
                    st.markdown(f"**{T['ct_rank']}：{rv}**")
                    try:
                        sorted_s = df[rv].dropna().sort_values().reset_index(drop=True)
                        fig, ax = plt.subplots(figsize=(9, 3.6))
                        apply_plot_style(fig, ax)
                        colors_rank = [ACCENT2 if v >= sorted_s.median() else "#f06b6b"
                                       for v in sorted_s]
                        ax.bar(range(len(sorted_s)), sorted_s.values, color=colors_rank, alpha=0.8)
                        ax.axhline(sorted_s.median(), color=ACCENT, linewidth=1.5,
                                   linestyle="--", label=f"{T.get('chart_median','中位数')}={sorted_s.median():.2f}")
                        ax.set_xlabel(T.get("chart_rank_x","样本排序（从小到大）")); ax.set_ylabel(rv)
                        ax.set_title(f"{rv} {T["ct_rank"]}")
                        ax.legend(facecolor=PLOT_BG, edgecolor=PLOT_GRID,
                                  labelcolor=PLOT_FG, fontsize=8)
                        plt.tight_layout(pad=0.3)
                        _img_path = os.path.join(_TMPDIR, "_chart_5.png")
                        fig.savefig(_img_path, bbox_inches="tight", dpi=150,
                                   facecolor=fig.get_facecolor())
                        _chart_images.append(_img_path)
                        st.pyplot(fig, width='stretch')
                        plt.close()
                        any_chart_drawn = True
                    except Exception as e:
                        st.warning(T["ct_rank"]+T["ct_fail"]+str(e))

            # ── 条形图（横向）────────────────────────────────────────────────
            if "条形图" in chart_types and ccfg.get("hbar_y") and ccfg.get("hbar_x"):
                hy, hx = ccfg["hbar_y"], ccfg["hbar_x"]
                if hy in df.columns and hx in df.columns:
                    st.markdown(f"**{T['ct_hbar']}：{hx} {T['ct_by']} {hy}**")
                    try:
                        grp = df.groupby(hy)[hx].mean().reset_index().sort_values(hx)
                        fig, ax = plt.subplots(figsize=(8, max(3.5, len(grp) * 0.4)))
                        apply_plot_style(fig, ax)
                        colors_hb = [ACCENT if v >= grp[hx].median() else "#a56dde" for v in grp[hx]]
                        ax.barh(grp[hy].astype(str), grp[hx], color=colors_hb, alpha=0.85)
                        for i, val in enumerate(grp[hx]):
                            ax.text(val * 1.01, i, f"{val:.2f}", va="center", fontsize=8, color=PLOT_FG)
                        ax.set_xlabel(hx); ax.set_title(f"{hx} {T["chart_mean"]}（{T["ct_by"]} {hy}）")
                        plt.tight_layout(pad=0.3)
                        _img_path = os.path.join(_TMPDIR, "_chart_6.png")
                        fig.savefig(_img_path, bbox_inches="tight", dpi=150,
                                   facecolor=fig.get_facecolor())
                        _chart_images.append(_img_path)
                        st.pyplot(fig, width='stretch')
                        plt.close()
                        any_chart_drawn = True
                    except Exception as e:
                        st.warning(T["ct_hbar"]+T["ct_fail"]+str(e))

            # ── 面积图 ────────────────────────────────────────────────────────
            if "面积图" in chart_types and ccfg.get("area_x") and ccfg.get("area_y"):
                ax_col, ay_list = ccfg["area_x"], ccfg["area_y"]
                if ax_col in df.columns and ay_list:
                    st.markdown(f"**{T['ct_area']}：{'、'.join(ay_list[:4])} {T['ct_trend']} {ax_col}**")
                    try:
                        df_a = df[[ax_col] + ay_list[:4]].dropna().sort_values(ax_col)
                        fig, ax = plt.subplots(figsize=(9, 3.8))
                        apply_plot_style(fig, ax)
                        area_colors = [ACCENT, ACCENT2, "#f06b6b", "#a56dde"]
                        for i, av in enumerate(ay_list[:4]):
                            ax.fill_between(df_a[ax_col], df_a[av],
                                            alpha=0.35, color=area_colors[i], label=av)
                            ax.plot(df_a[ax_col], df_a[av], color=area_colors[i], linewidth=1.5)
                        ax.set_xlabel(ax_col)
                        ax.legend(facecolor=PLOT_BG, edgecolor=PLOT_GRID, labelcolor=PLOT_FG, fontsize=8)
                        ax.set_title(f"{T['ct_area']}：{'、'.join(ay_list[:4])}")
                        plt.tight_layout(pad=0.3)
                        _img_path = os.path.join(_TMPDIR, "_chart_7.png")
                        fig.savefig(_img_path, bbox_inches="tight", dpi=150,
                                   facecolor=fig.get_facecolor())
                        _chart_images.append(_img_path)
                        st.pyplot(fig, width='stretch')
                        plt.close()
                        any_chart_drawn = True
                    except Exception as e:
                        st.warning(T["ct_area"]+T["ct_fail"]+str(e))

            # ── 气泡图 ────────────────────────────────────────────────────────
            if "气泡图" in chart_types and ccfg.get("bubble_x") and ccfg.get("bubble_y"):
                bx2 = ccfg["bubble_x"]; by2 = ccfg["bubble_y"]
                bs  = ccfg.get("bubble_size")
                if bx2 in df.columns and by2 in df.columns:
                    st.markdown(f"**{T['ct_bubble']}：{bx2} {T['ct_vs']} {by2}（{bs or T['ct_unified']}）**")
                    try:
                        valid = df[[bx2, by2] + ([bs] if bs and bs in df.columns else [])].dropna()
                        sizes = (valid[bs] / valid[bs].max() * 800 + 50).values if bs and bs in df.columns else 80
                        fig, ax = plt.subplots(figsize=(7, 4.5))
                        apply_plot_style(fig, ax)
                        sc = ax.scatter(valid[bx2], valid[by2], s=sizes,
                                        c=valid[bx2], cmap="Blues", alpha=0.7, edgecolors=PLOT_BG, linewidth=0.5)
                        plt.colorbar(sc, ax=ax, label=bx2, fraction=0.03, pad=0.02)
                        ax.set_xlabel(bx2); ax.set_ylabel(by2)
                        ax.set_title(f"{T["ct_bubble"]}：{by2} vs {bx2}")
                        plt.tight_layout(pad=0.3)
                        _img_path = os.path.join(_TMPDIR, "_chart_8.png")
                        fig.savefig(_img_path, bbox_inches="tight", dpi=150,
                                   facecolor=fig.get_facecolor())
                        _chart_images.append(_img_path)
                        st.pyplot(fig, width='stretch')
                        plt.close()
                        any_chart_drawn = True
                    except Exception as e:
                        st.warning(T["ct_bubble"]+T["ct_fail"]+str(e))

            # ── 饼图 ──────────────────────────────────────────────────────────
            if "饼图" in chart_types and ccfg.get("pie_label") and ccfg.get("pie_value"):
                pl, pv = ccfg["pie_label"], ccfg["pie_value"]
                if pl in df.columns and pv in df.columns:
                    st.markdown(f"**{T['ct_pie']}：{pv} {T['ct_by']} {pl} {T['ct_share']}**")
                    try:
                        pie_data = df.groupby(pl)[pv].sum().reset_index()
                        pie_data = pie_data[pie_data[pv] > 0].nlargest(10, pv)
                        fig, ax = plt.subplots(figsize=(6, 5))
                        fig.patch.set_facecolor(PLOT_BG); ax.set_facecolor(PLOT_BG)
                        pie_colors = [ACCENT, ACCENT2, "#f06b6b", "#f0b86b", "#a56dde",
                                      "#3dd68c", "#5b8dee", "#e07b6b", "#7b9dee", "#6bded0"]
                        wedges, texts, autotexts = ax.pie(
                            pie_data[pv], labels=pie_data[pl].astype(str),
                            colors=pie_colors[:len(pie_data)],
                            autopct="%1.1f%%", startangle=90,
                            textprops={"color": PLOT_FG, "fontsize": 8},
                            wedgeprops={"edgecolor": PLOT_BG, "linewidth": 1.5}
                        )
                        for at in autotexts:
                            at.set_color(PLOT_BG); at.set_fontsize(7)
                        ax.set_title(f"{pv} {T["ct_share"]}（{T["ct_by"]} {pl}）", color=PLOT_FG)
                        plt.tight_layout(pad=0.3)
                        _img_path = os.path.join(_TMPDIR, "_chart_9.png")
                        fig.savefig(_img_path, bbox_inches="tight", dpi=150,
                                   facecolor=fig.get_facecolor())
                        _chart_images.append(_img_path)
                        st.pyplot(fig, width='stretch')
                        plt.close()
                        any_chart_drawn = True
                    except Exception as e:
                        st.warning(T["ct_pie"]+T["ct_fail"]+str(e))

            # ── 圆环图 ────────────────────────────────────────────────────────
            if "圆环图" in chart_types and ccfg.get("donut_label") and ccfg.get("donut_value"):
                dl, dv = ccfg["donut_label"], ccfg["donut_value"]
                if dl in df.columns and dv in df.columns:
                    st.markdown(f"**{T['ct_donut']}：{dv} {T['ct_by']} {dl} {T['ct_share']}**")
                    try:
                        donut_data = df.groupby(dl)[dv].sum().reset_index()
                        donut_data = donut_data[donut_data[dv] > 0].nlargest(10, dv)
                        fig, ax = plt.subplots(figsize=(6, 5))
                        fig.patch.set_facecolor(PLOT_BG); ax.set_facecolor(PLOT_BG)
                        pie_colors = [ACCENT, ACCENT2, "#f06b6b", "#f0b86b", "#a56dde",
                                      "#3dd68c", "#5b8dee", "#e07b6b", "#7b9dee", "#6bded0"]
                        wedges, texts, autotexts = ax.pie(
                            donut_data[dv], labels=donut_data[dl].astype(str),
                            colors=pie_colors[:len(donut_data)],
                            autopct="%1.1f%%", startangle=90, pctdistance=0.75,
                            textprops={"color": PLOT_FG, "fontsize": 8},
                            wedgeprops={"edgecolor": PLOT_BG, "linewidth": 1.5, "width": 0.5}
                        )
                        for at in autotexts:
                            at.set_color(PLOT_BG); at.set_fontsize(7)
                        ax.set_title(f"{dv} {T["ct_donut"]}（{T["ct_by"]} {dl}）", color=PLOT_FG)
                        plt.tight_layout(pad=0.3)
                        _img_path = os.path.join(_TMPDIR, "_chart_10.png")
                        fig.savefig(_img_path, bbox_inches="tight", dpi=150,
                                   facecolor=fig.get_facecolor())
                        _chart_images.append(_img_path)
                        st.pyplot(fig, width='stretch')
                        plt.close()
                        any_chart_drawn = True
                    except Exception as e:
                        st.warning(T["ct_donut"]+T["ct_fail"]+str(e))

            # ── 小提琴图 ──────────────────────────────────────────────────────
            if "小提琴图" in chart_types and ccfg.get("violin_var"):
                vv = ccfg["violin_var"]; vg = ccfg.get("violin_group")
                if vv in df.columns:
                    st.markdown(f"**{T['ct_violin']}：{vv}{'（'+T['ct_by']+' '+vg+' 分组）' if vg else ''}**")
                    try:
                        import matplotlib.patches as mpatches
                        fig, ax = plt.subplots(figsize=(8, 4))
                        apply_plot_style(fig, ax)
                        if vg and vg in df.columns:
                            groups  = [g for g in df[vg].unique()]
                            gdata   = [df[df[vg] == g][vv].dropna().values for g in groups]
                            vparts  = ax.violinplot(gdata, positions=range(len(groups)),
                                                    showmedians=True, showmeans=False)
                            vcolors = [ACCENT, ACCENT2, "#f06b6b", "#f0b86b", "#a56dde"]
                            for i, pc in enumerate(vparts["bodies"]):
                                pc.set_facecolor(vcolors[i % len(vcolors)])
                                pc.set_alpha(0.7)
                            vparts["cmedians"].set_color(ACCENT2); vparts["cmedians"].set_linewidth(2)
                            for part in ["cbars","cmins","cmaxes"]:
                                if part in vparts:
                                    vparts[part].set_color(PLOT_FG); vparts[part].set_linewidth(0.8)
                            ax.set_xticks(range(len(groups)))
                            ax.set_xticklabels([str(g) for g in groups], rotation=20,
                                               ha="right", fontsize=8, color=PLOT_FG)
                            ax.set_title(f"{vv} {T["ct_by"]} {vg} {T["ct_violin"]}")
                        else:
                            vparts = ax.violinplot([df[vv].dropna().values],
                                                   showmedians=True, showmeans=False)
                            for pc in vparts["bodies"]:
                                pc.set_facecolor(ACCENT); pc.set_alpha(0.7)
                            vparts["cmedians"].set_color(ACCENT2); vparts["cmedians"].set_linewidth(2)
                            ax.set_xticks([1]); ax.set_xticklabels([vv], color=PLOT_FG)
                            ax.set_title(f"{vv} {T["ct_violin"]}")
                        ax.set_ylabel(vv)
                        plt.tight_layout(pad=0.3)
                        _img_path = os.path.join(_TMPDIR, "_chart_11.png")
                        fig.savefig(_img_path, bbox_inches="tight", dpi=150,
                                   facecolor=fig.get_facecolor())
                        _chart_images.append(_img_path)
                        st.pyplot(fig, width='stretch')
                        plt.close()
                        any_chart_drawn = True
                    except Exception as e:
                        st.warning(T["ct_violin"]+T["ct_fail"]+str(e))

            # ── 雷达图 ────────────────────────────────────────────────────────
            if "雷达图" in chart_types and ccfg.get("radar_vars") and len(ccfg["radar_vars"]) >= 3:
                rvars = ccfg["radar_vars"]; rgrp = ccfg.get("radar_group")
                st.markdown(f"**{T['ct_radar']}：{'、'.join(rvars)}{'（'+T['ct_by']+' '+rgrp+' 分组）' if rgrp else ''}**")
                try:
                    angles = np.linspace(0, 2 * np.pi, len(rvars), endpoint=False).tolist()
                    angles += angles[:1]
                    fig, ax = plt.subplots(figsize=(5.5, 5.5), subplot_kw=dict(polar=True))
                    fig.patch.set_facecolor(PLOT_BG); ax.set_facecolor(PLOT_BG)
                    ax.set_thetagrids(np.degrees(angles[:-1]), rvars,
                                      color=PLOT_FG, fontsize=8)
                    ax.tick_params(colors=PLOT_FG, labelsize=7)
                    ax.spines["polar"].set_color(PLOT_GRID)
                    ax.yaxis.set_tick_params(labelcolor=PLOT_FG)
                    radar_colors = [ACCENT, ACCENT2, "#f06b6b", "#a56dde", "#f0b86b"]
                    # 标准化各维度到 0-1
                    norm_df = df[rvars].copy()
                    for rv_col in rvars:
                        mn, mx = norm_df[rv_col].min(), norm_df[rv_col].max()
                        norm_df[rv_col] = (norm_df[rv_col] - mn) / (mx - mn + 1e-9)
                    if rgrp and rgrp in df.columns:
                        for gi, gv in enumerate(df[rgrp].unique()[:5]):
                            grp_vals = norm_df[df[rgrp] == gv][rvars].mean().tolist()
                            grp_vals += grp_vals[:1]
                            ax.plot(angles, grp_vals, color=radar_colors[gi % len(radar_colors)],
                                    linewidth=2, label=str(gv))
                            ax.fill(angles, grp_vals, alpha=0.12,
                                    color=radar_colors[gi % len(radar_colors)])
                        ax.legend(facecolor=PLOT_BG, edgecolor=PLOT_GRID,
                                  labelcolor=PLOT_FG, fontsize=8,
                                  loc="upper right", bbox_to_anchor=(1.3, 1.1))
                    else:
                        vals = norm_df[rvars].mean().tolist() + [norm_df[rvars].mean().tolist()[0]]
                        ax.plot(angles, vals, color=ACCENT, linewidth=2)
                        ax.fill(angles, vals, alpha=0.2, color=ACCENT)
                    ax.set_title(f"{T["ct_radar"]}（{T.get("chart_mean","均值")}）", color=PLOT_FG, pad=20)
                    plt.tight_layout(pad=0.3)
                    _img_path = os.path.join(_TMPDIR, "_chart_12.png")
                    fig.savefig(_img_path, bbox_inches="tight", dpi=150,
                               facecolor=fig.get_facecolor())
                    _chart_images.append(_img_path)
                    st.pyplot(fig, width='stretch')
                    plt.close()
                    any_chart_drawn = True
                except Exception as e:
                    st.warning(T["ct_radar"]+T["ct_fail"]+str(e))

            # ── 热力图 ────────────────────────────────────────────────────────
            if "热力图" in chart_types:
                hm_vars = ccfg.get("heatmap_vars") or target_vars
                hm_vars = [v for v in hm_vars if v in df.columns and v in df.select_dtypes(include=np.number).columns]
                if len(hm_vars) >= 2:
                    st.markdown(f"**{T['ct_heatmap']}：{' · '.join(hm_vars[:10])} {T['chart_corr_title'] if False else T.get('chart_heat_sub','相关矩阵')}**")
                    try:
                        hm_df = df[hm_vars[:10]].corr().round(2)
                        fig, ax = plt.subplots(figsize=(max(5, len(hm_vars)*1.1), max(4, len(hm_vars)*0.9)))
                        apply_plot_style(fig, ax)
                        im = ax.imshow(hm_df.values, cmap="RdYlBu_r", vmin=-1, vmax=1, aspect="auto")
                        plt.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
                        ax.set_xticks(range(len(hm_vars[:10]))); ax.set_yticks(range(len(hm_vars[:10])))
                        ax.set_xticklabels(hm_vars[:10], rotation=30, ha="right", fontsize=8, color=PLOT_FG)
                        ax.set_yticklabels(hm_vars[:10], fontsize=8, color=PLOT_FG)
                        for i in range(len(hm_vars[:10])):
                            for j in range(len(hm_vars[:10])):
                                val = hm_df.values[i, j]
                                ax.text(j, i, f"{val:.2f}", ha="center", va="center",
                                        fontsize=7, color="white" if abs(val) > 0.5 else PLOT_FG)
                        ax.set_title(T["corr_heatmap"], color=PLOT_FG)
                        plt.tight_layout(pad=0.3)
                        _img_path = os.path.join(_TMPDIR, "_chart_13.png")
                        fig.savefig(_img_path, bbox_inches="tight", dpi=150,
                                   facecolor=fig.get_facecolor())
                        _chart_images.append(_img_path)
                        st.pyplot(fig, width='stretch')
                        plt.close()
                        any_chart_drawn = True
                    except Exception as e:
                        st.warning(T["ct_heatmap"]+T["ct_fail"]+str(e))

            # ── 组合图（柱+线）───────────────────────────────────────────────
            if "组合图（柱+线）" in chart_types and ccfg.get("combo_x") and ccfg.get("combo_bar_y") and ccfg.get("combo_line_y"):
                cx = ccfg["combo_x"]; cby = ccfg["combo_bar_y"]; cly = ccfg["combo_line_y"]
                if cx in df.columns and cby in df.columns and cly in df.columns:
                    st.markdown(f"**{T['ct_combo']}：{cby}（{T['ct_bar']}）+ {cly}（{T['ct_line']}）{T['ct_by']} {cx}**")
                    try:
                        grp_c = df.groupby(cx)[[cby, cly]].mean().reset_index()
                        fig, ax1 = plt.subplots(figsize=(9, 4))
                        apply_plot_style(fig, ax1)
                        x_idx = np.arange(len(grp_c))
                        ax1.bar(x_idx, grp_c[cby], color=ACCENT, alpha=0.75, label=cby, width=0.6)
                        ax1.set_ylabel(cby, color=ACCENT)
                        ax1.tick_params(axis="y", labelcolor=ACCENT)
                        ax2 = ax1.twinx()
                        ax2.set_facecolor(PLOT_BG)
                        ax2.plot(x_idx, grp_c[cly], color=ACCENT2, linewidth=2.5,
                                 marker="o", markersize=5, label=cly)
                        ax2.set_ylabel(cly, color=ACCENT2)
                        ax2.tick_params(axis="y", labelcolor=ACCENT2)
                        for sp in ax2.spines.values():
                            sp.set_visible(False)
                        ax1.set_xticks(x_idx)
                        ax1.set_xticklabels(grp_c[cx].astype(str), rotation=20,
                                            ha="right", fontsize=8, color=PLOT_FG)
                        lines1, labels1 = ax1.get_legend_handles_labels()
                        lines2, labels2 = ax2.get_legend_handles_labels()
                        ax1.legend(lines1 + lines2, labels1 + labels2,
                                   facecolor=PLOT_BG, edgecolor=PLOT_GRID,
                                   labelcolor=PLOT_FG, fontsize=8, loc="upper left")
                        ax1.set_title(f"{T['ct_combo']}：{cby} {T['ct_vs']} {cly}")
                        plt.tight_layout(pad=0.3)
                        _img_path = os.path.join(_TMPDIR, "_chart_14.png")
                        fig.savefig(_img_path, bbox_inches="tight", dpi=150,
                                   facecolor=fig.get_facecolor())
                        _chart_images.append(_img_path)
                        st.pyplot(fig, width='stretch')
                        plt.close()
                        any_chart_drawn = True
                    except Exception as e:
                        st.warning(T["ct_combo"]+T["ct_fail"]+str(e))
            if not any_chart_drawn:
                plot_vars = target_vars[:8]
                if len(plot_vars) >= 2:
                    st.markdown(f"**{T['ct_mean_cmp']}：{' · '.join(plot_vars)}**")
                    fig, ax = plt.subplots(figsize=(max(8, len(plot_vars) * 1.2), 4.2))
                    apply_plot_style(fig, ax)
                    x     = np.arange(len(plot_vars))
                    means = [df[v].mean() for v in plot_vars]
                    stds  = [df[v].std()  for v in plot_vars]
                    bar_colors = [ACCENT, ACCENT2, "#f06b6b", "#f0b86b",
                                  "#a56dde", "#3dd68c", "#f06b6b", "#5b8dee"]
                    bars = ax.bar(x, means, color=[bar_colors[i % len(bar_colors)]
                                                   for i in range(len(plot_vars))],
                                  alpha=0.85, width=0.6,
                                  yerr=stds, capsize=4,
                                  error_kw=dict(ecolor="#8890a8", elinewidth=1.2, capthick=1.2))
                    for bar, m in zip(bars, means):
                        ax.text(bar.get_x() + bar.get_width() / 2,
                                bar.get_height() * 1.01 + max(stds) * 0.05,
                                f"{m:.2f}", ha="center", va="bottom",
                                fontsize=8, color=PLOT_FG)
                    ax.set_xticks(x)
                    ax.set_xticklabels(plot_vars, rotation=20, ha="right",
                                       fontsize=8, color=PLOT_FG)
                    ax.set_ylabel(T["ct_mean_y"])
                    ax.set_title(T["ct_mean_title"])
                    plt.tight_layout(pad=0.3)
                    _img_path = os.path.join(_TMPDIR, "_chart_15.png")
                    fig.savefig(_img_path, bbox_inches="tight", dpi=150,
                               facecolor=fig.get_facecolor())
                    _chart_images.append(_img_path)
                    st.pyplot(fig, width='stretch')
                    plt.close()
                else:
                    st.info(T["ct_min2"])

        st.markdown("<hr class='divider'>", unsafe_allow_html=True)

    # ── C · 相关性分析 ────────────────────────────────────────────────────────
    if "相关性分析" in chosen_types:
        sub_header(T["s3_sec_c"])
        corr_vars = st.session_state.confirmed_corr_vars
        if len(corr_vars) < 2:
            st.info(T["s3_no_corr"])
        else:
            corr_matrix = df[corr_vars].corr().round(3)
            # 相关系数热力图
            fig, ax = plt.subplots(figsize=(max(5, len(corr_vars) * 1.2),
                                            max(4, len(corr_vars) * 1.0)))
            apply_plot_style(fig, ax)
            import matplotlib.colors as mcolors
            cmap = plt.cm.RdYlBu_r
            im = ax.imshow(corr_matrix.values, cmap=cmap, vmin=-1, vmax=1, aspect="auto")
            plt.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
            ax.set_xticks(range(len(corr_vars))); ax.set_yticks(range(len(corr_vars)))
            ax.set_xticklabels(corr_vars, rotation=30, ha="right", fontsize=8, color=PLOT_FG)
            ax.set_yticklabels(corr_vars, fontsize=8, color=PLOT_FG)
            for i in range(len(corr_vars)):
                for j in range(len(corr_vars)):
                    val = corr_matrix.values[i, j]
                    ax.text(j, i, f"{val:.2f}", ha="center", va="center",
                            fontsize=7, color="white" if abs(val) > 0.5 else PLOT_FG)
            ax.set_title(T["corr_heatmap"], color=PLOT_FG)
            plt.tight_layout(pad=0.3)
            _img_path = os.path.join(_TMPDIR, "_chart_16.png")
            fig.savefig(_img_path, bbox_inches="tight", dpi=150,
                       facecolor=fig.get_facecolor())
            _chart_images.append(_img_path)
            st.pyplot(fig, width='stretch')
            plt.close()

            # 展示相关系数表
            st.markdown('<div class="card">', unsafe_allow_html=True)
            st.dataframe(corr_matrix, width='stretch')
            st.markdown("</div>", unsafe_allow_html=True)

            # 提取显著变量对
            var_pairs = []
            for i, v1 in enumerate(corr_vars):
                for j, v2 in enumerate(corr_vars):
                    if i < j:
                        var_pairs.append((v1, v2, float(corr_matrix.loc[v1, v2])))
            var_pairs.sort(key=lambda x: abs(x[2]), reverse=True)
            top_pairs = var_pairs[:6]

            fb_corr = fallback_corr(top_pairs)
            corr_summary = "；".join([f"{v1}~{v2}(r={r:.3f})" for v1, v2, r in top_pairs])

            with st.spinner(T["s3_spin_corr"]):
                text, is_ai = ai_corr_interp(question, corr_summary, fb_corr,
                                             lang=st.session_state.get("lang","zh"))
            render_ai(T["s3_lbl_corr"], text, is_ai)

        st.markdown("<hr class='divider'>", unsafe_allow_html=True)

    # ── D · 回归分析 ──────────────────────────────────────────────────────────
    if "回归分析" in chosen_types:
        sub_header(T["s3_sec_d"])
        y_var  = st.session_state.confirmed_y
        x_vars = st.session_state.confirmed_x
        scatter_x = st.session_state.confirmed_scatter

        if not y_var or not x_vars:
            st.info(T["s3_no_reg"])
        else:
            series = df[y_var].dropna()
            chart_types = st.session_state.get("chart_types", [])
            ccfg        = st.session_state.get("chart_config", {})

            # ── 用户所选图表（在回归块内渲染）─────────────────────────────────
            corr_val = None
            reg_charts_drawn = False

            # 散点图
            if "散点图" in chart_types and ccfg.get("scatter_x") and ccfg.get("scatter_y"):
                sx, sy = ccfg["scatter_x"], ccfg["scatter_y"]
                if sx in df.columns and sy in df.columns:
                    st.markdown(f"**{T['ct_scatter_vs']}：{sy} vs {sx}**")
                    try:
                        fig, ax = plt.subplots(figsize=(6, 3.8))
                        apply_plot_style(fig, ax)
                        valid = df[[sx, sy]].dropna()
                        ax.scatter(valid[sx], valid[sy], color=ACCENT, alpha=0.5, s=16, edgecolors="none")
                        if len(valid) > 2:
                            z = np.polyfit(valid[sx], valid[sy], 1)
                            xs_arr = np.linspace(valid[sx].min(), valid[sx].max(), 200)
                            ax.plot(xs_arr, np.poly1d(z)(xs_arr), color=ACCENT2,
                                    linewidth=2, linestyle="--", label=T.get("chart_trend","趋势线"))
                            ax.legend(facecolor=PLOT_BG, edgecolor=PLOT_GRID, labelcolor=PLOT_FG, fontsize=8)
                            corr_val = float(valid[sx].corr(valid[sy]))
                        ax.set_xlabel(sx); ax.set_ylabel(sy)
                        ax.set_title(f"{sy} ~ {sx}（r={corr_val:.3f}）" if corr_val is not None else f"{sy} ~ {sx}")
                        plt.tight_layout(pad=0.3)
                        _img_path = os.path.join(_TMPDIR, "_chart_17.png")
                        fig.savefig(_img_path, bbox_inches="tight", dpi=150,
                                   facecolor=fig.get_facecolor())
                        _chart_images.append(_img_path)
                        st.pyplot(fig, width='stretch')
                        plt.close()
                        reg_charts_drawn = True
                    except Exception as e:
                        st.warning(T["ct_scatter"]+T["ct_fail"]+str(e))

            # 折线图
            if "折线图" in chart_types and ccfg.get("line_x") and ccfg.get("line_y"):
                lx, ly_list = ccfg["line_x"], ccfg["line_y"]
                if lx in df.columns and ly_list:
                    st.markdown(f"**{T['ct_line']}：{'、'.join(ly_list[:4])} {T['ct_trend']} {lx}**")
                    try:
                        df_l = df[[lx] + ly_list[:4]].dropna().sort_values(lx)
                        fig, ax = plt.subplots(figsize=(9, 3.8))
                        apply_plot_style(fig, ax)
                        lc = [ACCENT, ACCENT2, "#f06b6b", "#a56dde"]
                        for i, lv in enumerate(ly_list[:4]):
                            ax.plot(df_l[lx], df_l[lv], color=lc[i], linewidth=2, label=lv, alpha=0.9)
                        ax.set_xlabel(lx)
                        ax.legend(facecolor=PLOT_BG, edgecolor=PLOT_GRID, labelcolor=PLOT_FG, fontsize=8)
                        ax.set_title(T["ct_line"])
                        plt.tight_layout(pad=0.3)
                        _img_path = os.path.join(_TMPDIR, "_chart_18.png")
                        fig.savefig(_img_path, bbox_inches="tight", dpi=150,
                                   facecolor=fig.get_facecolor())
                        _chart_images.append(_img_path)
                        st.pyplot(fig, width='stretch')
                        plt.close()
                        reg_charts_drawn = True
                    except Exception as e:
                        st.warning(T["ct_line"]+T["ct_fail"]+str(e))

            # 柱状图
            if "柱状图" in chart_types and ccfg.get("bar_x") and ccfg.get("bar_y"):
                bx, by = ccfg["bar_x"], ccfg["bar_y"]
                if bx in df.columns and by in df.columns:
                    st.markdown(f"**{T['ct_bar']}：{by} {T['ct_by']} {bx}**")
                    try:
                        grp = df.groupby(bx)[by].mean().reset_index()
                        fig, ax = plt.subplots(figsize=(8, 3.8))
                        apply_plot_style(fig, ax)
                        bars_b = ax.bar(grp[bx].astype(str), grp[by], color=ACCENT, alpha=0.85)
                        for bar, val in zip(bars_b, grp[by]):
                            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() * 1.01,
                                    f"{val:.2f}", ha="center", va="bottom", fontsize=8, color=PLOT_FG)
                        ax.set_xlabel(bx); ax.set_ylabel(by)
                        ax.set_title(f"{by} {T['chart_mean']}（{T['ct_by']} {bx}）")
                        plt.xticks(rotation=20, ha="right")
                        plt.tight_layout(pad=0.3)
                        _img_path = os.path.join(_TMPDIR, "_chart_19.png")
                        fig.savefig(_img_path, bbox_inches="tight", dpi=150,
                                   facecolor=fig.get_facecolor())
                        _chart_images.append(_img_path)
                        st.pyplot(fig, width='stretch')
                        plt.close()
                        reg_charts_drawn = True
                    except Exception as e:
                        st.warning(T["ct_bar"]+T["ct_fail"]+str(e))

            # 变化排序图
            if "变化排序图" in chart_types and ccfg.get("rank_var"):
                rv = ccfg["rank_var"]
                if rv in df.columns:
                    st.markdown(f"**{T['ct_rank']}：{rv}**")
                    try:
                        sorted_s = df[rv].dropna().sort_values().reset_index(drop=True)
                        fig, ax = plt.subplots(figsize=(9, 3.6))
                        apply_plot_style(fig, ax)
                        cr = [ACCENT2 if v >= sorted_s.median() else "#f06b6b" for v in sorted_s]
                        ax.bar(range(len(sorted_s)), sorted_s.values, color=cr, alpha=0.8)
                        ax.axhline(sorted_s.median(), color=ACCENT, linewidth=1.5, linestyle="--",
                                   label=f"{T.get('chart_median','中位数')}={sorted_s.median():.2f}")
                        ax.set_xlabel(T.get("chart_rank_x","样本排序（从小到大）")); ax.set_ylabel(rv)
                        ax.set_title(f"{rv} {T["ct_rank"]}")
                        ax.legend(facecolor=PLOT_BG, edgecolor=PLOT_GRID, labelcolor=PLOT_FG, fontsize=8)
                        plt.tight_layout(pad=0.3)
                        _img_path = os.path.join(_TMPDIR, "_chart_20.png")
                        fig.savefig(_img_path, bbox_inches="tight", dpi=150,
                                   facecolor=fig.get_facecolor())
                        _chart_images.append(_img_path)
                        st.pyplot(fig, width='stretch')
                        plt.close()
                        reg_charts_drawn = True
                    except Exception as e:
                        st.warning(T["ct_rank"]+T["ct_fail"]+str(e))

            # 箱线图
            if "箱线图" in chart_types and ccfg.get("box_var"):
                bv, bg = ccfg["box_var"], ccfg.get("box_group")
                if bv in df.columns:
                    st.markdown(f"**{T['ct_box']}：{bv}{'（'+T['ct_by']+' '+bg+' 分组）' if bg else ''}**")
                    try:
                        fig, ax = plt.subplots(figsize=(7, 3.8))
                        apply_plot_style(fig, ax)
                        if bg and bg in df.columns:
                            gd = [df[df[bg] == g][bv].dropna().values for g in df[bg].unique()]
                            gl = df[bg].unique().tolist()
                            bp = ax.boxplot(gd, patch_artist=True, medianprops=dict(color=ACCENT2, linewidth=2))
                            cb = [ACCENT, "#a56dde", "#3dd68c", "#f06b6b", "#f0b86b"]
                            for patch, color in zip(bp["boxes"], cb * 10):
                                patch.set_facecolor(color); patch.set_alpha(0.7)
                            ax.set_xticklabels(gl, rotation=20, ha="right", fontsize=8, color=PLOT_FG)
                            ax.set_title(f"{bv} {T["ct_by"]} {bg} {T["ct_box"]}")
                        else:
                            bp = ax.boxplot(df[bv].dropna().values, patch_artist=True,
                                            medianprops=dict(color=ACCENT2, linewidth=2))
                            bp["boxes"][0].set_facecolor(ACCENT); bp["boxes"][0].set_alpha(0.7)
                            ax.set_xticklabels([bv], color=PLOT_FG)
                            ax.set_title(f"{bv} {T["ct_box"]}")
                        ax.set_ylabel(bv)
                        plt.tight_layout(pad=0.3)
                        _img_path = os.path.join(_TMPDIR, "_chart_21.png")
                        fig.savefig(_img_path, bbox_inches="tight", dpi=150,
                                   facecolor=fig.get_facecolor())
                        _chart_images.append(_img_path)
                        st.pyplot(fig, width='stretch')
                        plt.close()
                        reg_charts_drawn = True
                    except Exception as e:
                        st.warning(T["ct_box"]+T["ct_fail"]+str(e))

            # 兜底：若用户未选任何图表，保留原分布直方图 + 散点图
            if not reg_charts_drawn:
                col_l, col_r = st.columns(2)
                with col_l:
                    st.markdown(f"**{y_var} {T['chart_dist_sub']}**")
                    fig, ax = plt.subplots(figsize=(5.5, 3.8))
                    apply_plot_style(fig, ax)
                    ax.hist(series, bins=30, color=ACCENT, alpha=0.75, edgecolor=PLOT_BG, linewidth=0.5)
                    if len(series) > 10:
                        kde = gaussian_kde(series)
                        xs  = np.linspace(series.min(), series.max(), 300)
                        ax2 = ax.twinx(); ax2.set_facecolor(PLOT_BG)
                        ax2.plot(xs, kde(xs), color=ACCENT2, linewidth=2)
                        ax2.set_yticks([])
                        for sp in ax2.spines.values(): sp.set_edgecolor(PLOT_GRID)
                    ax.set_xlabel(y_var); ax.set_ylabel(T["chart_freq"]); ax.set_title(f"{y_var} {T["chart_dist"]}")
                    plt.tight_layout(pad=0.3)
                    _img_path = os.path.join(_TMPDIR, "_chart_22.png")
                    fig.savefig(_img_path, bbox_inches="tight", dpi=150,
                               facecolor=fig.get_facecolor())
                    _chart_images.append(_img_path)
                    st.pyplot(fig, width='stretch')
                    plt.close()

                scatter_x = st.session_state.confirmed_scatter
                with col_r:
                    if scatter_x and scatter_x in df.columns:
                        st.markdown(f"**{y_var} vs {scatter_x}**")
                        fig, ax = plt.subplots(figsize=(5.5, 3.8))
                        apply_plot_style(fig, ax)
                        valid = df[[scatter_x, y_var]].dropna()
                        ax.scatter(valid[scatter_x], valid[y_var], color=ACCENT, alpha=0.5, s=16, edgecolors="none")
                        if len(valid) > 2:
                            z = np.polyfit(valid[scatter_x], valid[y_var], 1)
                            xs_arr = np.linspace(valid[scatter_x].min(), valid[scatter_x].max(), 200)
                            ax.plot(xs_arr, np.poly1d(z)(xs_arr), color=ACCENT2, linewidth=2, linestyle="--", label=T.get("chart_trend","趋势线"))
                            ax.legend(facecolor=PLOT_BG, edgecolor=PLOT_GRID, labelcolor=PLOT_FG, fontsize=8)
                            corr_val = float(valid[scatter_x].corr(valid[y_var]))
                        ax.set_xlabel(scatter_x); ax.set_ylabel(y_var); ax.set_title(f"{y_var} ~ {scatter_x}")
                        plt.tight_layout(pad=0.3)
                        _img_path = os.path.join(_TMPDIR, "_chart_23.png")
                        fig.savefig(_img_path, bbox_inches="tight", dpi=150,
                                   facecolor=fig.get_facecolor())
                        _chart_images.append(_img_path)
                        st.pyplot(fig, width='stretch')
                        plt.close()

            if corr_val is not None:
                _sx = ccfg.get("scatter_x") or st.session_state.confirmed_scatter or ""
                _sy = ccfg.get("scatter_y") or y_var
                fb_chart = fallback_chart(_sy, float(df[_sy].dropna().mean()),
                                          float(df[_sy].dropna().median()),
                                          float(df[_sy].dropna().skew()), _sx, corr_val,
                                          lang=st.session_state.get("lang","zh"))
                with st.spinner(T["s3_spin_chart"]):
                    text, is_ai = ai_chart_interp(
                        question, _sy, float(df[_sy].dropna().mean()),
                        float(df[_sy].dropna().median()), float(df[_sy].dropna().skew()),
                        _sx, corr_val, fb_chart,
                        lang=st.session_state.get("lang","zh"),
                    )
                render_ai(T["s3_lbl_chart"], text, is_ai)

            # 执行回归
            if not st.session_state.analysis_done:
                reg_df = df[[y_var] + x_vars].dropna()
                if len(reg_df) < len(x_vars) + 2:
                    st.error(T["s3_no_sample"])
                    st.stop()
                X = reg_df[x_vars].values
                y = reg_df[y_var].values
                model = LinearRegression(); model.fit(X, y)
                y_pred = model.predict(X)
                n, k   = len(y), len(x_vars)
                r2     = r2_score(y, y_pred)
                adj_r2 = 1 - (1 - r2) * (n - 1) / (n - k - 1)
                rmse   = np.sqrt(mean_squared_error(y, y_pred))
                resid  = y - y_pred
                X_aug  = np.column_stack([np.ones(n), X])
                s2     = np.sum(resid**2) / (n - k - 1)
                se_all = np.sqrt(np.diag(np.linalg.pinv(X_aug.T @ X_aug)) * s2)
                coefs  = np.concatenate([[model.intercept_], model.coef_])
                t_vals = coefs / se_all
                p_vals = [2 * (1 - stats.t.cdf(abs(t), df=n-k-1)) for t in t_vals]
                ci95   = se_all * stats.t.ppf(0.975, df=n-k-1)
                terms  = " + ".join([f"({model.coef_[i]:.4f}×{x_vars[i]})" for i in range(k)])
                eq     = f"{y_var} = {model.intercept_:.4f} + {terms}"
                st.session_state.reg_results = dict(
                    model=model, y_pred=y_pred, resid=resid, n=n, k=k,
                    r2=r2, adj_r2=adj_r2, rmse=rmse, coefs=coefs,
                    se_all=se_all, t_vals=t_vals, p_vals=p_vals, ci95=ci95, equation=eq,
                )
                st.session_state.analysis_done = True

            R = st.session_state.reg_results
            reg_r2 = R["r2"]

            # 模型摘要
            st.markdown('<div class="card">', unsafe_allow_html=True)
            st.markdown(
                '<div style="font-size:0.68rem;font-weight:700;letter-spacing:0.1em;'
                'text-transform:uppercase;color:#5b8dee;margin-bottom:0.8rem;">模型摘要</div>',
                unsafe_allow_html=True,
            )
            mc1, mc2, mc3, mc4 = st.columns(4)
            for col, val, lbl in zip(
                [mc1, mc2, mc3, mc4],
                [f"{R['r2']:.4f}", f"{R['adj_r2']:.4f}", f"{R['rmse']:.4f}", R['n']],
                [T["s3_r2"], T["s3_adj_r2"], T["s3_rmse"], T["s3_n_eff"]],
            ):
                col.markdown(
                    f'<div class="stat-box"><div class="stat-number">{val}</div>'
                    f'<div class="stat-label">{lbl}</div></div>',
                    unsafe_allow_html=True,
                )
            st.markdown("</div>", unsafe_allow_html=True)

            # 系数表
            names = ["(截距)"] + x_vars; rows_html = ""
            for i, nm in enumerate(names):
                p = R["p_vals"][i]
                if p < 0.001: star = "***"
                elif p < 0.01: star = "**"
                elif p < 0.05: star = "*"
                elif p < 0.1:  star = "."
                else:          star = ""
                if nm != "(截距)":
                    d = "正向" if R["coefs"][i] > 0 else "负向"
                    (sig_vars_list if p < 0.05 else insig_vars_list).append(
                        f"{nm}（{d}，β={R['coefs'][i]:.4f}，p={p:.4f}）"
                    )
                p_str  = f"<span class='sig'>{p:.4f} {star}</span>" if star else f"{p:.4f}"
                ci_str = f"[{R['coefs'][i]-R['ci95'][i]:.4f}, {R['coefs'][i]+R['ci95'][i]:.4f}]"
                rows_html += (
                    f"<tr><td><b>{nm}</b></td><td>{R['coefs'][i]:.4f}</td>"
                    f"<td>{R['se_all'][i]:.4f}</td><td>{R['t_vals'][i]:.4f}</td>"
                    f"<td>{p_str}</td><td>{ci_str}</td></tr>"
                )

            st.markdown('<div class="card">', unsafe_allow_html=True)
            st.markdown(
                '<div style="font-size:0.68rem;font-weight:700;letter-spacing:0.1em;'
                'text-transform:uppercase;color:#5b8dee;margin-bottom:0.8rem;">回归系数</div>',
                unsafe_allow_html=True,
            )
            st.markdown(
                f"<table class='reg-table'><thead><tr>"
                f"<th>" + "</th><th>".join(T["s3_reg_cols"]) + "</th>"
                f"</tr></thead><tbody>{rows_html}</tbody></table>"
                f"<div style='font-size:0.7rem;color:#3d4560;margin-top:0.7rem;'>"
                f"*** p&lt;0.001 &nbsp;** p&lt;0.01 &nbsp;* p&lt;0.05 &nbsp;. p&lt;0.1</div>",
                unsafe_allow_html=True,
            )
            st.markdown("</div>", unsafe_allow_html=True)

            # 残差图
            rd1, rd2 = st.columns(2)
            with rd1:
                fig, ax = plt.subplots(figsize=(5.5, 3.6))
                apply_plot_style(fig, ax)
                ax.scatter(R["y_pred"], R["resid"], color=ACCENT, alpha=0.5, s=15, edgecolors="none")
                ax.axhline(0, color=ACCENT2, linewidth=1.5, linestyle="--")
                ax.set_xlabel(T["s3_resid_x"]); ax.set_ylabel(T["s3_resid_y"]); ax.set_title(T["s3_resid_ttl"])
                plt.tight_layout(pad=0.3)
                _img_path = os.path.join(_TMPDIR, "_chart_24.png")
                fig.savefig(_img_path, bbox_inches="tight", dpi=150,
                           facecolor=fig.get_facecolor())
                _chart_images.append(_img_path)
                st.pyplot(fig, width='stretch')
                plt.close()
            with rd2:
                fig, ax = plt.subplots(figsize=(5.5, 3.6))
                apply_plot_style(fig, ax)
                (osm, osr), (slope_q, intercept_q, _) = stats.probplot(R["resid"])
                ax.scatter(osm, osr, color=ACCENT, alpha=0.55, s=15, edgecolors="none")
                ax.plot([min(osm), max(osm)],
                        [slope_q*min(osm)+intercept_q, slope_q*max(osm)+intercept_q],
                        color=ACCENT2, linewidth=2, linestyle="--")
                ax.set_xlabel(T["s3_qq_x"]); ax.set_ylabel(T["s3_qq_y"]); ax.set_title(T["s3_qq_ttl"])
                plt.tight_layout(pad=0.3)
                _img_path = os.path.join(_TMPDIR, "_chart_25.png")
                fig.savefig(_img_path, bbox_inches="tight", dpi=150,
                           facecolor=fig.get_facecolor())
                _chart_images.append(_img_path)
                st.pyplot(fig, width='stretch')
                plt.close()

            # 回归方程
            st.markdown(
                f'<div class="card" style="font-family:\'JetBrains Mono\',monospace;'
                f'font-size:0.83rem;color:#b8c2d8;">📐 {R["equation"]}</div>',
                unsafe_allow_html=True,
            )

            # AI 回归解读
            fb_reg = fallback_regression(
                y_var, "、".join(x_vars), R["r2"], R["adj_r2"], R["rmse"],
                "；".join(sig_vars_list), "；".join(insig_vars_list), R["equation"],
                lang=st.session_state.get("lang","zh")
            )
            with st.spinner(T["s3_spin_reg"]):
                text, is_ai = ai_reg_interp(
                    question=question, y_var=y_var, x_vars_str="、".join(x_vars),
                    r2=R["r2"], adj_r2=R["adj_r2"], rmse=R["rmse"],
                    sig_vars="；".join(sig_vars_list), insig_vars="；".join(insig_vars_list),
                    equation=R["equation"], fallback_text=fb_reg,
                    lang=st.session_state.get("lang","zh"),
                )
            render_ai(T["s3_lbl_reg"], text, is_ai)

    # 导航
    st.markdown("<hr class='divider'>", unsafe_allow_html=True)
    col_back, col_next = st.columns([1, 1])
    with col_back:
        if st.button(T["s3_back"]):
            st.session_state.step = 2
            st.rerun()
    with col_next:
        if st.button(T["s3_next"], width='stretch'):
            st.session_state["desc_summary"]    = desc_summary
            st.session_state["sig_vars_list"]   = sig_vars_list
            st.session_state["insig_vars_list"] = insig_vars_list
            st.session_state["reg_r2_cache"]    = reg_r2
            st.session_state["chart_images"]    = _chart_images
            st.session_state.step = 4
            st.rerun()


# ══════════════════════════════════════════════════════════════════════════════
# STEP 4 · 总体结论与建议
# ══════════════════════════════════════════════════════════════════════════════
elif st.session_state.step == 4:
    T = _get_T()
    phase_header("04", T["s4_title"])

    question      = st.session_state.question
    df            = st.session_state.df
    chosen_types_raw4 = st.session_state.confirmed_analysis_types
    _type_map_en2zh4 = {"Descriptive Stats":"描述性统计","Comparison":"对比分析","Trend Analysis":"趋势分析","Correlation":"相关性分析","Regression":"回归分析"}
    chosen_types = [_type_map_en2zh4.get(t,t) for t in chosen_types_raw4]
    R             = st.session_state.reg_results
    sig_vars      = "；".join(st.session_state.get("sig_vars_list", []))
    miss_pct      = df.isnull().sum().sum() / df.size * 100
    desc_summary  = st.session_state.get("desc_summary", "")
    use_reg       = "回归分析" in chosen_types

    st.markdown(
        f'<div class="question-display">{T["s3_q_prefix"]}{question}</div>',
        unsafe_allow_html=True,
    )

    # 摘要卡片
    _type_zh2display = {"描述性统计":T["s2_all_types"][0] if len(T["s2_all_types"])>0 else "描述性统计",
                        "对比分析":T["s2_all_types"][1] if len(T["s2_all_types"])>1 else "对比分析",
                        "趋势分析":T["s2_all_types"][2] if len(T["s2_all_types"])>2 else "趋势分析",
                        "相关性分析":T["s2_all_types"][3] if len(T["s2_all_types"])>3 else "相关性分析",
                        "回归分析":T["s2_all_types"][4] if len(T["s2_all_types"])>4 else "回归分析"}
    types_display = "  ·  ".join([_type_zh2display.get(t,t) for t in chosen_types]) if chosen_types else "—"
    r2_display    = f"{R['r2']:.4f}" if (use_reg and R) else T["s4_no_reg"]
    sig_count     = len(sig_vars.split("；")) if sig_vars else 0
    sig_display   = (f"{sig_count} " + T.get("sig_label","个显著")) if use_reg else "—"
    n_display     = R["n"] if (use_reg and R) else df.shape[0]

    c1, c2, c3, c4 = st.columns(4)
    for col, val, lbl in zip(
        [c1, c2, c3, c4],
        [types_display, str(n_display), r2_display, sig_display],
        [T["s4_stat_types"], T["s4_stat_n"], T["s4_stat_r2"], T["s4_stat_sig"]],
    ):
        col.markdown(
            f'<div class="stat-box">'
            f'<div class="stat-number" style="font-size:0.88rem;padding:0.4rem 0;line-height:1.4;">'
            f'{val}</div><div class="stat-label">{lbl}</div></div>',
            unsafe_allow_html=True,
        )

    st.markdown("<div style='height:1rem'></div>", unsafe_allow_html=True)

    # AI 综合结论
    analysis_types_str = "、".join(chosen_types)
    r2_for_ai  = st.session_state.get("reg_r2_cache", 0.0)
    n_for_ai   = R["n"] if R else int(df.shape[0])
    y_for_ai   = st.session_state.confirmed_y or (
        (st.session_state.confirmed_desc_vars or
         st.session_state.confirmed_target_vars or ["—"])[0]
    )
    x_for_ai   = "、".join(st.session_state.confirmed_x) if st.session_state.confirmed_x else "—"

    fb_conc = fallback_conclusion(
        question, y_for_ai, x_for_ai, chosen_types,
        r2_for_ai, sig_vars, miss_pct, n_for_ai, desc_summary,
        lang=st.session_state.get("lang","zh")
    )
    with st.spinner(T["s4_spinner"]):
        conclusion, is_ai = ai_conclusion(
            question=question, y_var=y_for_ai, x_vars_str=x_for_ai,
            analysis_types_str=analysis_types_str,
            r2=r2_for_ai, sig_vars=sig_vars,
            miss_pct=miss_pct, n=n_for_ai,
            desc_summary=desc_summary, fallback_text=fb_conc,
            lang=st.session_state.get("lang","zh"),
        )

    conclusion_label = "AI 总体结论" if is_ai else "规则总体结论"

    # 格式化结论文字：与 render_ai 相同的分点渲染逻辑
    conc_lines = conclusion.split("\n")
    conc_html = ""
    for line in conc_lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith("【") and "】" in line:
            conc_html += (
                f'<div style="color:#5b8dee;font-weight:700;font-size:0.88rem;'
                f'margin-top:1.1rem;margin-bottom:0.4rem;">{line}</div>'
            )
        elif len(line) >= 2 and line[0].isdigit() and line[1] == '.':
            num  = line[0]
            rest = line[2:].strip()
            conc_html += (
                f'<div style="display:flex;gap:0.6rem;margin-bottom:0.5rem;align-items:flex-start;">'
                f'<span style="color:#5b8dee;font-weight:700;font-size:0.88rem;'
                f'min-width:1.2rem;flex-shrink:0;">{num}.</span>'
                f'<span style="color:#c0cce0;font-size:0.9rem;line-height:1.75;">{rest}</span>'
                f'</div>'
            )
        else:
            conc_html += f'<div style="color:#a8b4d0;font-size:0.9rem;line-height:1.75;margin-bottom:0.3rem;">{line}</div>'

    st.markdown(
        f'<div class="conclusion-card">'
        f'<div class="conclusion-title">{T["s4_conc_title"]}'
        f'<span style="font-size:0.68rem;font-weight:400;color:#3d4d6a;margin-left:0.8rem;">'
        f'{"（AI 生成）" if is_ai else "（本地规则生成）"}</span></div>'
        f'{conc_html}'
        f'</div>',
        unsafe_allow_html=True,
    )

    st.markdown("<hr class='divider'>", unsafe_allow_html=True)

    # ══════════════════════════════════════════════════════════════════════════
    # 报告生成与下载模块
    # ══════════════════════════════════════════════════════════════════════════
    st.markdown(
        '<div style="font-size:1.05rem;font-weight:700;color:#dde2f0;'
        'margin:0.5rem 0 1rem 0;padding-bottom:0.4rem;'
        'border-bottom:2px solid #2a3a5a;">'
        f"{T['s4_rpt_title']}</div>",
        unsafe_allow_html=True,
    )
    st.markdown(
        f'<div style="font-size:0.82rem;color:#6b7a9a;margin-bottom:1rem;">{T["s4_rpt_hint"]}</div>',
        unsafe_allow_html=True,
    )

    # ── AI 报告生成函数 ────────────────────────────────────────────────────────
    @st.cache_data(show_spinner=False)
    def generate_report_text(question: str, analysis_types_str: str,
                             desc_summary: str, sig_vars: str, conclusion: str,
                             r2: float, n: int, miss_pct: float,
                             y_var: str, x_vars_str: str,
                             lang: str = "zh") -> str:
        _rpt_lang = lang
        if _rpt_lang == "en":
            _intro = "Please generate a complete data analysis report based on the following information."
            _struct = (
                "Please strictly follow this structure, with each section marked by [Section Name]:\n\n"
                "[1. Research Question]\n1. State the core question\n2. Research background and purpose\n\n"
                "[2. Data Overview]\n1. Sample size, variables, missing rate (with values)\n2. Key variable stats\n\n"
                "[3. Methods]\n1. Analysis types used and rationale\n2. Variable selection rationale\n\n"
                "[4. Key Findings]\n1. Most important finding (with values) — significance\n"
                "2. Second finding (with values) — significance\n3. Third finding (with values) — significance\n\n"
                "[5. Conclusion]\n1. Direct answer to research question (with data)\n2. Overall judgment\n\n"
                "[6. Recommendations]\n1. First actionable recommendation\n2. Second recommendation\n3. Third recommendation\n\n"
                "[7. Limitations]\n1. Data limitations\n2. Method limitations\n3. Suggested improvements\n\n"
                "Requirements: each point includes specific values + explanation. Professional report style, no ** bold or # headers."
            )
        else:
            _intro = "请根据以下分析信息，生成一份完整的数据分析报告。"
            _struct = (
                "请严格按以下结构生成报告，每个章节用【章节名】标记：\n\n"
                "【1. 研究问题】\n1. 明确说明本次分析要回答的核心问题\n2. 说明研究背景与分析目的\n\n"
                "【2. 数据概览】\n1. 样本量、变量数量、缺失率等基本信息（含具体数值）\n2. 核心变量描述（含均值、标准差等）\n\n"
                "【3. 分析方法】\n1. 本次使用的分析类型及其适用理由\n2. 变量选择依据\n\n"
                "【4. 关键发现】\n1. 最重要的发现（含具体数值）— 说明意义\n"
                "2. 第二重要发现（含具体数值）— 说明意义\n3. 第三重要发现（含具体数值）— 说明意义\n\n"
                "【5. 结论】\n1. 直接回答研究问题（含关键数据）\n2. 综合判断\n\n"
                "【6. 建议】\n1. 第一条可执行建议\n2. 第二条可执行建议\n3. 第三条可执行建议\n\n"
                "【7. 局限性】\n1. 数据局限性\n2. 方法局限性\n3. 建议改进方向\n\n"
                "要求：每点包含具体数值+分析解释，风格为专业分析报告，不使用 ** 加粗或 # 标题。"
            )
        prompt = (
            f"{_intro}\n\n"
            f"研究问题：{question}\n"
            f"分析方法：{analysis_types_str}\n"
            f"样本量：{n}，缺失率：{miss_pct:.1f}%\n"
            f"核心变量：{y_var}（自变量：{x_vars_str}）\n"
            f"回归 R²：{r2:.4f}\n"
            f"显著自变量：{sig_vars or '无'}\n"
            f"描述性统计摘要：{desc_summary}\n"
            f"分析结论：{conclusion}\n\n"
        ) + _struct
        result = call_claude(_get_analyst_system(), prompt, max_tokens=2000)
        if _is_api_error(result):
            desc_snip = desc_summary[:200] if desc_summary else "见描述性统计表格"
            conc_snip = conclusion[:300] if conclusion else "见上方结论部分"
            quality   = "数据质量良好" if miss_pct < 5 else "需关注缺失值影响"
            sample_ok = "较为充足" if n >= 100 else "相对有限，结论需谨慎外推"
            sig_text  = sig_vars if sig_vars else "暂无显著自变量"
            parts = [
                "【1. 研究问题】",
                f"1. 本次分析围绕 {question} 展开，旨在通过数据探索回答上述问题。",
                "2. 分析目的：理解核心变量的分布规律与变量间的关联关系。",
                "",
                "【2. 数据概览】",
                f"1. 样本量 {n}，数据缺失率 {miss_pct:.1f}%。",
                f"2. 核心变量：{y_var}；{desc_snip}。",
                "",
                "【3. 分析方法】",
                f"1. 本次执行了 {analysis_types_str} 分析，涵盖主要变量的统计特征与关系探索。",
                f"2. 核心变量选择：{y_var} 为目标变量，{x_vars_str} 为解释变量。",
                "",
                "【4. 关键发现】",
                f"1. 回归模型 R²={r2:.4f}，模型解释了因变量约 {r2*100:.1f}% 的变异。",
                f"2. 显著影响变量：{sig_text}。",
                f"3. 数据缺失率 {miss_pct:.1f}%，{quality}。",
                "",
                "【5. 结论】",
                f"1. {conc_snip}",
                "",
                "【6. 建议】",
                "1. 结合专业知识验证统计结论，避免纯数据驱动的误判。",
                "2. 考虑增加控制变量或扩大样本量以提升模型稳健性。",
                "3. 如存在非线性关系，可尝试多项式回归或机器学习方法。",
                "",
                "【7. 局限性】",
                f"1. 样本量 {n}，{sample_ok}。",
                "2. 线性回归假设变量间为线性关系，若存在非线性关系则模型精度有限。",
                "3. 建议在更大样本或跨时间段数据上重复验证本次结论。",
            ]
            return "\n".join(parts)
        return result

    # ── 生成报告按钮 ──────────────────────────────────────────────────────────
    col_gen, col_dl_word, col_nav_back, col_nav_restart = st.columns([2, 2, 1, 1])

    with col_gen:
        if st.button(T["s4_gen_btn"], width='stretch'):
            with st.spinner(T["s4_gen_spin"]):
                # 分别生成中文结论和英文结论
                _fb_zh = fallback_conclusion(question, y_for_ai, x_for_ai, chosen_types,
                                             r2_for_ai, sig_vars, miss_pct, n_for_ai,
                                             desc_summary, lang="zh")
                _conc_zh, _ = ai_conclusion(
                    question=question, y_var=y_for_ai, x_vars_str=x_for_ai,
                    analysis_types_str=analysis_types_str,
                    r2=r2_for_ai, sig_vars=sig_vars,
                    miss_pct=miss_pct, n=n_for_ai,
                    desc_summary=desc_summary, fallback_text=_fb_zh,
                    lang="zh",
                )
                _fb_en = fallback_conclusion(question, y_for_ai, x_for_ai, chosen_types,
                                             r2_for_ai, sig_vars, miss_pct, n_for_ai,
                                             desc_summary, lang="en")
                _conc_en, _ = ai_conclusion(
                    question=question, y_var=y_for_ai, x_vars_str=x_for_ai,
                    analysis_types_str=analysis_types_str,
                    r2=r2_for_ai, sig_vars=sig_vars,
                    miss_pct=miss_pct, n=n_for_ai,
                    desc_summary=desc_summary, fallback_text=_fb_en,
                    lang="en",
                )
                # 分析方法字符串也要双语
                _types_zh = "、".join(chosen_types)
                _types_en = "、".join([{"描述性统计":"Descriptive Stats","对比分析":"Comparison",
                    "趋势分析":"Trend Analysis","相关性分析":"Correlation","回归分析":"Regression"
                }.get(t, t) for t in chosen_types])
                # 生成中文版报告（用中文结论）
                report_zh = generate_report_text(
                    question=question, analysis_types_str=_types_zh,
                    desc_summary=desc_summary, sig_vars=sig_vars, conclusion=_conc_zh,
                    r2=r2_for_ai, n=n_for_ai, miss_pct=miss_pct,
                    y_var=y_for_ai, x_vars_str=x_for_ai,
                    lang="zh",
                )
                # 生成英文版报告（用英文结论）
                report_en = generate_report_text(
                    question=question, analysis_types_str=_types_en,
                    desc_summary=desc_summary, sig_vars=sig_vars, conclusion=_conc_en,
                    r2=r2_for_ai, n=n_for_ai, miss_pct=miss_pct,
                    y_var=y_for_ai, x_vars_str=x_for_ai,
                    lang="en",
                )
                st.session_state["report_text"] = report_zh
                st.session_state["report_text_en"] = report_en
            st.success(T["s4_gen_done"])

    # ── 报告预览 ──────────────────────────────────────────────────────────────
    report_text = st.session_state.get("report_text", "")
    if report_text:
        with st.expander(T["s4_preview"], expanded=False):
            preview_lines = report_text.split("\n")
            preview_html  = ""
            for line in preview_lines:
                line = line.strip()
                if not line:
                    preview_html += "<div style='height:0.4rem'></div>"
                elif line.startswith("【") and "】" in line:
                    preview_html += (
                        f'<div style="color:#5b8dee;font-weight:700;font-size:0.88rem;'
                        f'margin-top:1rem;margin-bottom:0.3rem;">{line}</div>'
                    )
                elif len(line) >= 2 and line[0].isdigit() and line[1] == '.':
                    num  = line[0]; rest = line[2:].strip()
                    preview_html += (
                        f'<div style="display:flex;gap:0.5rem;margin-bottom:0.4rem;">'
                        f'<span style="color:#5b8dee;font-weight:700;min-width:1rem;">{num}.</span>'
                        f'<span style="color:#c0cce0;font-size:0.86rem;line-height:1.7;">{rest}</span>'
                        f'</div>'
                    )
                else:
                    preview_html += f'<div style="color:#a8b4d0;font-size:0.86rem;line-height:1.7;">{line}</div>'
            st.markdown(f'<div style="padding:0.5rem 0;">{preview_html}</div>', unsafe_allow_html=True)

        # ── Word 下载 ──────────────────────────────────────────────────────────
        with col_dl_word:
            try:
                from docx import Document as _DocxDoc
                from docx.shared import Pt, RGBColor, Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                import io as _io
                import os as _os
                from datetime import datetime as _dt

                def _build_docx(report_txt: str, question: str,
                                chart_paths: list, lang: str = "zh") -> bytes:
                    doc = _DocxDoc()
                    # 根据 lang 参数选择封面文字（独立于 UI 语言）
                    _T = _TRANSLATIONS[lang]

                    # ── 封面 ──────────────────────────────────────────────────
                    title_p = doc.add_paragraph()
                    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = title_p.add_run(_T["s4_rpt_cover"])
                    run.bold = True; run.font.size = Pt(22)
                    run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x6a)

                    sub_p = doc.add_paragraph()
                    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    sub_run = sub_p.add_run(_T["s4_rpt_q"] + question)
                    sub_run.font.size = Pt(12)
                    sub_run.font.color.rgb = RGBColor(0x5b, 0x6a, 0x90)

                    date_p = doc.add_paragraph()
                    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    date_run = date_p.add_run(_T["s4_rpt_date"] + _dt.now().strftime('%Y-%m-%d %H:%M'))
                    date_run.font.size = Pt(10)
                    date_run.font.color.rgb = RGBColor(0x8a, 0x92, 0xb0)

                    doc.add_paragraph()

                    # ── 正文 ──────────────────────────────────────────────────
                    for line in report_txt.split("\n"):
                        line = line.strip()
                        if not line:
                            doc.add_paragraph()
                            continue
                        _is_zh_header = line.startswith("【") and "】" in line
                        _is_en_header = (line.startswith("[") and ". " in line and "]" in line
                                         and len(line) > 3 and line[1].isdigit())
                        if _is_zh_header or _is_en_header:
                            h = doc.add_heading(line, level=2)
                            for run in h.runs:
                                run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x8a)
                        elif len(line) >= 2 and line[0].isdigit() and line[1] == '.':
                            p = doc.add_paragraph(style="List Number")
                            p.add_run(line[2:].strip()).font.size = Pt(11)
                        else:
                            p = doc.add_paragraph(line)
                            if p.runs: p.runs[0].font.size = Pt(11)

                    # ── 图表附录 ──────────────────────────────────────────────
                    valid_charts = [p for p in chart_paths if _os.path.exists(p)]
                    if valid_charts:
                        doc.add_page_break()
                        h_app = doc.add_heading(_T["s4_rpt_annex"], level=1)
                        for run in h_app.runs:
                            run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x8a)
                        for i, img_path in enumerate(valid_charts[:12], 1):
                            try:
                                cap = doc.add_paragraph(_T["s4_rpt_fig"] + str(i))
                                cap.runs[0].font.size = Pt(9)
                                cap.runs[0].font.color.rgb = RGBColor(0x8a, 0x92, 0xb0)
                                doc.add_picture(img_path, width=Inches(5.5))
                                doc.add_paragraph()
                            except Exception:
                                pass

                    buf = _io.BytesIO()
                    doc.save(buf)
                    return buf.getvalue()

                _charts = st.session_state.get("chart_images", [])
                _q_slug = question[:20].replace(' ','_')
                # 中文 Word
                docx_zh = _build_docx(report_text, question, _charts, lang="zh")
                st.download_button(
                    label="⬇ 下载中文分析报告 (.docx)",
                    data=docx_zh,
                    file_name=f"分析报告_{_q_slug}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    width='stretch',
                )
                # 英文 Word
                report_text_en = st.session_state.get("report_text_en","")
                if report_text_en:
                    try:
                        docx_en = _build_docx(report_text_en, question, _charts, lang="en")
                        st.download_button(
                            label="⬇ Download English Report (.docx)",
                            data=docx_en,
                            file_name=f"report_{_q_slug}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            width='stretch',
                        )
                    except Exception as _e2:
                        st.warning(f"English Word failed: {_e2}")
            except ImportError:
                st.warning(T["s4_no_docx"])
            except Exception as e:
                st.error(T["s4_word_fail"] + str(e))

    # ── 纯文本下载 ────────────────────────────────────────────────────────────
    if report_text:
        _q_slug2 = question[:20].replace(' ','_')
        st.download_button(
            label="⬇ 下载中文报告 (.txt)",
            data=report_text.encode("utf-8"),
            file_name=f"分析报告_{_q_slug2}.txt",
            mime="text/plain",
        )
        report_text_en = st.session_state.get("report_text_en","")
        if report_text_en:
            st.download_button(
                label="⬇ Download English Report (.txt)",
                data=report_text_en.encode("utf-8"),
                file_name=f"report_{_q_slug2}.txt",
                mime="text/plain",
            )

    st.markdown("<hr class='divider'>", unsafe_allow_html=True)

    col_restart, col_back = st.columns([1, 1])
    with col_back:
        if st.button(T["s4_back"]):
            st.session_state.step = 3
            st.rerun()
    with col_restart:
        if st.button(T["s4_restart"], width='stretch'):
            for k in list(st.session_state.keys()):
                del st.session_state[k]
            st.rerun()
